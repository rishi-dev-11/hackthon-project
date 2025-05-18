# output_generator.py
import logging
import os
import tempfile
import io
import base64 # For image data handling if needed by add_figure_to_docx
import docx
from docx.shared import Pt, Inches # RGBColor not used directly here but good to have
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PIL_Image # Renamed to avoid clash with docx.image.Image
import cv2 # For image conversion fallback
import numpy as np # For image conversion fallback
import fitz # For PDF text extraction fallback

logger = logging.getLogger(__name__)

# Import enhanced functions from documorph_fixes if available
try:
    from documorph_fixes import (
        add_table_to_docx as enhanced_add_table_to_docx,
        add_figure_to_docx as enhanced_add_figure_to_docx
    )
    FIXES_MODULE_AVAILABLE_OUTPUT = True
    logger.info("Using enhanced table/figure adding functions from documorph_fixes.")
except ImportError:
    FIXES_MODULE_AVAILABLE_OUTPUT = False
    logger.warning("documorph_fixes not available. Using internal table/figure adding functions.")
    # Define fallback internal functions if fixes are not available
    def enhanced_add_table_to_docx(doc, table_data, caption):
        return _internal_add_table_to_docx(doc, table_data, caption)
    def enhanced_add_figure_to_docx(doc, figure_data, caption):
        return _internal_add_figure_to_docx(doc, figure_data, caption)

def _internal_add_figure_to_docx(doc, figure, caption):
    """Internal fallback: Add a figure to a DOCX document."""
    try:
        if not figure.get("image_bytes"):
            logger.warning(f"No image bytes for figure {figure.get('figure_id', 'unknown')}")
            p_placeholder = doc.add_paragraph()
            p_placeholder.add_run(f"[Figure placeholder: {caption}]").italic = True
            p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(caption, style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True

        p_img = doc.add_paragraph()
        r_img = p_img.add_run()
        
        tmp_path = None # Define tmp_path before try block
        try:
            pil_image = PIL_Image.open(io.BytesIO(figure["image_bytes"]))
            image_format = pil_image.format.lower() if pil_image.format else figure.get('image_type', 'png')
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{image_format}") as tmp:
                tmp_path = tmp.name
                pil_image.save(tmp_path, format=image_format.upper())
            
            width_inches = min(6, figure.get("width", 500) / 96.0) # Convert pixels to inches (assuming 96 DPI)
            r_img.add_picture(tmp_path, width=Inches(width_inches))
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e_add_pic:
            logger.error(f"Error adding picture {figure.get('figure_id')} directly: {e_add_pic}. Trying conversion.")
            # Fallback: Try to convert to PNG if direct adding fails
            try:
                if tmp_path: os.unlink(tmp_path) # Clean up previous attempt
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_png:
                    tmp_path = tmp_png.name
                    # Convert using OpenCV if PIL failed or format was tricky
                    img_array = np.frombuffer(figure["image_bytes"], np.uint8)
                    cv_img = cv2.imdecode(img_array, cv2.IMREAD_COLOR) # Try as color
                    if cv_img is None: # Try as unchanged if color fails (e.g. for alpha channel)
                        cv_img = cv2.imdecode(img_array, cv2.IMREAD_UNCHANGED)
                    if cv_img is None: raise ValueError("OpenCV could not decode image")
                    cv2.imwrite(tmp_path, cv_img)

                width_inches = min(6, figure.get("width", 500) / 96.0)
                r_img.add_picture(tmp_path, width=Inches(width_inches))
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e_conv_pic:
                logger.error(f"Error adding converted picture {figure.get('figure_id')}: {e_conv_pic}")
                p_img.runs.clear() # Clear any failed attempts from run
                r_img = p_img.add_run(f"[Image display error: {figure.get('figure_id')}]")
                r_img.italic = True
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER


        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
        
        doc.add_paragraph(caption, style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception as e:
        logger.error(f"Critical error adding figure {figure.get('figure_id')} to DOCX: {e}", exc_info=True)
        p_err = doc.add_paragraph()
        p_err.add_run(f"[Figure error: {caption}]").italic = True
        return False

def _internal_add_table_to_docx(doc, table_data_dict, caption):
    """Internal fallback: Add a table to a DOCX document."""
    try:
        doc.add_paragraph(caption, style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        rows_to_add = table_data_dict.get("raw_data")
        df_to_add = table_data_dict.get("dataframe")

        if df_to_add is not None:
            # Prefer DataFrame if available
            num_rows = len(df_to_add) + 1 # +1 for header
            num_cols = len(df_to_add.columns)
            if num_rows == 1 and num_cols == 0: # Empty DataFrame
                 logger.warning(f"Empty DataFrame for table {caption}. Adding placeholder.")
                 doc.add_paragraph(f"[Empty Table: {caption}]").italic = True
                 return True


            table_obj = doc.add_table(rows=num_rows, cols=num_cols)
            table_obj.style = 'Table Grid'
            # Add headers
            for j, col_name in enumerate(df_to_add.columns):
                table_obj.cell(0, j).text = str(col_name)
                run = table_obj.cell(0,j).paragraphs[0].runs[0]
                run.bold = True
            # Add data
            for i, row_series in df_to_add.iterrows():
                for j, col_name in enumerate(df_to_add.columns):
                    table_obj.cell(i + 1, j).text = str(row_series[col_name])
            return True
        elif rows_to_add:
            if not any(rows_to_add): # All sublists are empty
                logger.warning(f"Empty raw_data for table {caption}. Adding placeholder.")
                doc.add_paragraph(f"[Empty Table: {caption}]").italic = True
                return True

            num_rows_raw = len(rows_to_add)
            num_cols_raw = max(len(r) for r in rows_to_add) if num_rows_raw > 0 else 0
            if num_rows_raw == 0 or num_cols_raw == 0:
                logger.warning(f"Effectively empty raw_data for table {caption}. Adding placeholder.")
                doc.add_paragraph(f"[Empty Table: {caption}]").italic = True
                return True

            table_obj = doc.add_table(rows=num_rows_raw, cols=num_cols_raw)
            table_obj.style = 'Table Grid'
            for i, row_list in enumerate(rows_to_add):
                for j, cell_text in enumerate(row_list):
                    table_obj.cell(i, j).text = str(cell_text)
                if i == 0: # Bold header for raw data too
                    for cell_obj in table_obj.rows[i].cells:
                        if cell_obj.paragraphs and cell_obj.paragraphs[0].runs:
                            cell_obj.paragraphs[0].runs[0].bold = True
            return True
        else:
            logger.warning(f"No data found for table {caption}. Adding placeholder.")
            doc.add_paragraph(f"[Table data not available: {caption}]").italic = True
            return False
    except Exception as e:
        logger.error(f"Error adding table {caption} to DOCX: {e}", exc_info=True)
        doc.add_paragraph(f"[Table error: {caption}]").italic = True
        return False


def apply_template_to_document(doc_path, template_settings, output_path, tables_list=None, figures_list=None, chapters_list=None):
    try:
        # Determine source document type and load/create initial docx.Document
        if doc_path.endswith('.docx'):
            source_doc = docx.Document(doc_path)
            # Create a new document to avoid modifying the original in-place during templating
            # and to ensure a clean slate for applying styles. Content will be copied.
            doc = docx.Document()
            for para in source_doc.paragraphs: # Copy paragraphs
                doc.add_paragraph(para.text, style=para.style if para.style.name != "No Spacing" else None)
            # Note: Copying tables and images from source docx to new docx is complex if they need to be preserved
            # The current logic focuses on adding *newly extracted* tables/figures.
            # If preserving original docx tables/figures is needed, this part requires more advanced docx manipulation.
        else: # PDF or TXT
            doc = docx.Document()
            if doc_path.endswith('.txt'):
                with open(doc_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    for para_text in content.split('\n\n'): # Assume double newline is paragraph
                        if para_text.strip(): doc.add_paragraph(para_text.strip())
            elif doc_path.endswith('.pdf'):
                try:
                    # Use fitz to extract text from PDF for the new DOCX
                    pdf_doc_obj = fitz.open(doc_path) # Using fitz directly
                    for page_num in range(len(pdf_doc_obj)):
                        page_text = pdf_doc_obj[page_num].get_text("text")
                        for para_text in page_text.split('\n\n'): # Heuristic paragraph split
                             if para_text.strip(): doc.add_paragraph(para_text.strip())
                except Exception as e_pdf_extract:
                    logger.error(f"Error extracting text from PDF {doc_path} for DOCX conversion: {e_pdf_extract}")
                    doc.add_paragraph(f"[Error extracting text from PDF: {os.path.basename(doc_path)}]")

        if not doc.paragraphs: # Ensure doc is not empty
            doc.add_paragraph("Content could not be loaded or document was empty.")

        # Apply template settings
        style = doc.styles['Normal']
        font = style.font
        font.name = template_settings.get('body_font', 'Calibri')
        font.size = Pt(template_settings.get('body_font_size', 11))
        
        # Set paragraph format for Normal style (line spacing, paragraph spacing)
        para_format = style.paragraph_format
        para_format.line_spacing = template_settings.get('line_spacing', 1.15)
        para_format.space_after = Pt(template_settings.get('paragraph_spacing', 0)) # Default to 0 if not specified

        for i in range(1, 4):
            heading_style_name = f'Heading {i}'
            if heading_style_name in doc.styles:
                heading_style = doc.styles[heading_style_name]
                h_font = heading_style.font
                h_font.name = template_settings.get(f'heading{i}_font', template_settings.get('heading_font', 'Calibri'))
                h_font.size = Pt(template_settings.get(f'heading{i}_font_size', 16 - i*2)) # Default size
                h_font.bold = template_settings.get(f'heading{i}_bold', True)
                
                h_para_format = heading_style.paragraph_format
                h_para_format.line_spacing = template_settings.get(f'heading{i}_line_spacing', 1.0)
                h_para_format.space_before = Pt(template_settings.get(f'heading{i}_space_before', 12))
                h_para_format.space_after = Pt(template_settings.get(f'heading{i}_space_after', 6))


        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(template_settings.get('margin_top', 1.0))
            section.bottom_margin = Inches(template_settings.get('margin_bottom', 1.0))
            section.left_margin = Inches(template_settings.get('margin_left', 1.0))
            section.right_margin = Inches(template_settings.get('margin_right', 1.0))

        if template_settings.get('header_text'):
            header = doc.sections[0].header
            header.paragraphs[0].text = template_settings['header_text']
        
        if template_settings.get('footer_text'):
            footer = doc.sections[0].footer
            footer.paragraphs[0].text = template_settings['footer_text']
            # Example: Add page number to footer if placeholder exists
            if "[Page]" in template_settings['footer_text']:
                 # This requires more complex field insertion for dynamic page numbers
                 logger.info("Page number placeholder found in footer_text. Manual field insertion needed for dynamic page numbers in python-docx.")


        if template_settings.get('include_tables_figures', True):
            # Use enhanced_add_table_to_docx and enhanced_add_figure_to_docx (which point to internal or fixed versions)
            if tables_list:
                doc.add_heading('Tables', level=1) # Add a section for tables
                for table_item in tables_list:
                    enhanced_add_table_to_docx(doc, table_item, table_item.get("caption", "Table"))
                    doc.add_paragraph() # Spacing
            
            if figures_list:
                doc.add_heading('Figures', level=1) # Add a section for figures
                for figure_item in figures_list:
                    enhanced_add_figure_to_docx(doc, figure_item, figure_item.get("caption", "Figure"))
                    doc.add_paragraph() # Spacing
        
        doc.save(output_path)
        logger.info(f"Templated document saved to {output_path}")
        return True, output_path
    except Exception as e:
        logger.error(f"Error applying template: {e}", exc_info=True)
        # Attempt to save a minimal error document
        try:
            error_doc = docx.Document()
            error_doc.add_paragraph(f"Error applying template to {os.path.basename(doc_path)}.")
            error_doc.add_paragraph(str(e))
            error_doc.save(output_path) # Save with error content
        except Exception as e_save:
            logger.error(f"Could not even save error document: {e_save}")
        return False, f"Error applying template: {str(e)}"
