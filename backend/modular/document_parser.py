# document_parser.py
import logging
import os
import tempfile
import io
import base64
import pandas as pd
import cv2 # OpenCV for image processing in fallbacks
import numpy as np # For image processing
from PIL import Image # For image processing

import fitz # PyMuPDF

from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument # Rename to avoid conflict with docx.Document

# unstructured imports will be conditional
try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.text import partition_text
    from unstructured.partition.image import partition_image
    from unstructured.documents.elements import Table as UnstructuredTable, Image as UnstructuredImageElement, Title
    UNSTRUCTURED_AVAILABLE_PARSER = True
except ImportError:
    UNSTRUCTURED_AVAILABLE_PARSER = False
    UnstructuredTable, UnstructuredImageElement, Title = None, None, None # Define for type hinting
    def partition_pdf(*args, **kwargs): raise ImportError("unstructured.partition.pdf not available")
    def partition_docx(*args, **kwargs): raise ImportError("unstructured.partition.docx not available")
    def partition_text(*args, **kwargs): raise ImportError("unstructured.partition.text not available")
    def partition_image(*args, **kwargs): raise ImportError("unstructured.partition.image not available")


logger = logging.getLogger(__name__)

# Import open_pdf from documorph_fixes if available, else use a local one
try:
    from documorph_fixes import open_pdf
    logger.info("Using open_pdf from documorph_fixes.")
except ImportError:
    logger.warning("documorph_fixes.open_pdf not found. Using internal open_pdf_internal.")
    from utils import open_pdf_internal as open_pdf


# Fallback for table_extraction and chart_extraction if top-level import fails
# These specific imports are for the version1.0.0 fallback logic
_table_extractor_v1 = None
_chart_extractor_v1 = None

# Attempt to import table_extraction from version1.0.0
try:
    import sys
    # Assuming 'backend/version1.0.0' is relative to where documorph_ai.py (and thus this module) runs
    v1_extraction_path = os.path.join(os.getcwd(), "backend", "version1.0.0")
    if os.path.exists(os.path.join(v1_extraction_path, "table_extraction.py")):
        sys.path.insert(0, v1_extraction_path) # Prepend to path
        from table_extraction import extract_tables as extract_tables_v1_func
        _table_extractor_v1 = extract_tables_v1_func
        logger.info("Successfully imported table_extraction from version1.0.0 for fallback.")
        sys.path.pop(0) # Clean up path
    else:
        logger.warning("backend/version1.0.0/table_extraction.py not found.")
except ImportError:
    logger.warning("Could not import table_extraction from version1.0.0.")
except Exception as e:
    logger.error(f"Error importing table_extraction from version1.0.0: {e}")


# Attempt to import chart_extraction from version1.0.0
try:
    import sys
    v1_extraction_path = os.path.join(os.getcwd(), "backend", "version1.0.0") # Redefine or pass as arg
    if os.path.exists(os.path.join(v1_extraction_path, "chart_extraction.py")):
        sys.path.insert(0, v1_extraction_path)
        from chart_extraction import identify_and_extract_charts as identify_and_extract_charts_v1_func
        _chart_extractor_v1 = identify_and_extract_charts_v1_func
        logger.info("Successfully imported chart_extraction from version1.0.0 for fallback.")
        sys.path.pop(0)
    else:
        logger.warning("backend/version1.0.0/chart_extraction.py not found.")
except ImportError:
    logger.warning("Could not import chart_extraction from version1.0.0.")
except Exception as e:
    logger.error(f"Error importing chart_extraction from version1.0.0: {e}")


def process_with_unstructured(file_path, file_type, user_tier_ocr_languages):
    if not UNSTRUCTURED_AVAILABLE_PARSER:
        logger.warning("Unstructured module not available for parsing. Cannot use process_with_unstructured.")
        return [], [], [], "Unstructured module not available for parsing."
    
    try:
        logger.info(f"Processing document with unstructured.io: {file_path}")
        elements = []
        extracted_tables = []
        extracted_figures = []
        
        extra_params = {"languages": user_tier_ocr_languages, "strategy": "hi_res" if "hi" in user_tier_ocr_languages else "fast"}
        if "hi" not in user_tier_ocr_languages: # Assuming "hi_res" strategy for premium, "fast" for free
             extra_params["strategy"] = "fast"


        if file_type == "pdf":
            elements = partition_pdf(file_path, **extra_params)
        elif file_type == "docx":
            # For docx, unstructured might not need OCR languages in the same way
            elements = partition_docx(file_path, strategy=extra_params.get("strategy", "fast"))
        elif file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as f: text_content = f.read()
            elements = partition_text(text_content)
        elif file_type in ["png", "jpg", "jpeg", "tiff"]:
            elements = partition_image(file_path, languages=user_tier_ocr_languages, strategy=extra_params.get("strategy", "fast"))
        else:
            return [], [], [], f"Unsupported file type for unstructured: {file_type}"
        
        texts_for_chunks = []
        for i, element in enumerate(elements):
            if isinstance(element, UnstructuredTable) and hasattr(element, "text"):
                table_text = element.text
                # Try to get HTML for better structure if available
                if hasattr(element, "metadata") and hasattr(element.metadata, "text_as_html"):
                    table_html = element.metadata.text_as_html
                    try:
                        df = pd.read_html(io.StringIO(table_html))[0]
                        raw_data_from_html = [df.columns.tolist()] + df.values.tolist()
                    except Exception:
                        df = None
                        raw_data_from_html = None
                else: # Fallback to parsing text
                    df = None
                    raw_data_from_html = None

                rows = []
                if raw_data_from_html:
                    rows = raw_data_from_html
                elif table_text: # Fallback to parsing plain text
                    lines = table_text.split('\n')
                    for line in lines:
                        cells = [cell.strip() for cell in line.split('|') if cell.strip()] # Basic pipe separation
                        if cells: rows.append(cells)
                
                # Create DataFrame if not already created and rows exist
                if df is None and rows:
                    try:
                        if len(rows) > 1 and rows[0]:
                           df = pd.DataFrame(rows[1:], columns=rows[0])
                        elif rows: # Single row table or no header
                           df = pd.DataFrame(rows)
                    except Exception as e_df:
                        logger.warning(f"Could not create DataFrame for unstructured table {i}: {e_df}")

                extracted_tables.append({
                    "table_id": f"table_unstructured_{i}",
                    "raw_data": rows,
                    "dataframe": df,
                    "page": getattr(element.metadata, 'page_number', 1),
                    "html": getattr(element.metadata, 'text_as_html', None)
                })
                texts_for_chunks.append(f"[TABLE table_unstructured_{i}]\n{table_text}\n[/TABLE]")
            
            elif isinstance(element, UnstructuredImageElement):
                fig_id = f"figure_unstructured_{i}"
                img_bytes = getattr(element.metadata, 'image_bytes', None) # Newer unstructured
                if not img_bytes and hasattr(element.metadata, 'image_data'): # Older unstructured
                    img_bytes = element.metadata.image_data

                extracted_figures.append({
                    "figure_id": fig_id,
                    "page": getattr(element.metadata, 'page_number', 1),
                    "image_bytes": img_bytes,
                    "image_type": getattr(element.metadata, 'filetype', 'png'), # or 'image_format'
                    "caption": getattr(element, 'text', f"Figure {fig_id}")
                })
                texts_for_chunks.append(f"[FIGURE {fig_id}]\n{getattr(element, 'text', '')}\n[/FIGURE]")
            
            elif isinstance(element, Title) and hasattr(element, "text"):
                texts_for_chunks.append(f"[TITLE] {element.text} [/TITLE]")
            
            elif hasattr(element, "text"):
                texts_for_chunks.append(element.text)
        
        doc_chunks = [LangchainDocument(page_content=text, metadata={"source": file_path, "chunk_idx": i}) for i, text in enumerate(texts_for_chunks)]
        logger.info(f"Unstructured extraction complete: {len(doc_chunks)} chunks, {len(extracted_tables)} tables, {len(extracted_figures)} figures")
        return doc_chunks, extracted_tables, extracted_figures, None
        
    except ImportError as e: # Specific catch for unstructured dependency issues
        logger.error(f"Unstructured or its dependency missing: {e}", exc_info=True)
        return [], [], [], f"Unstructured parsing failed due to missing dependency: {e}"
    except Exception as e:
        logger.error(f"Error processing with unstructured: {e}", exc_info=True)
        return [], [], [], f"Error processing with unstructured: {str(e)}"


def detect_tables_from_pdf(pdf_path):
    tables = []
    try:
        doc = open_pdf(pdf_path) # Uses open_pdf from fixes or internal
        for page_num, page in enumerate(doc):
            try:
                page_tables = page.find_tables()
                if page_tables.tables: # Check if find_tables found anything
                    for i, table_obj in enumerate(page_tables.tables):
                        raw_data = table_obj.extract()
                        df = None
                        if raw_data: # Ensure data is not empty
                            try: # Attempt to create DataFrame
                                if len(raw_data) > 1 and raw_data[0]: # Header + data
                                    df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
                                elif raw_data: # Single row or no header
                                    df = pd.DataFrame(raw_data)
                            except Exception as e_df:
                                logger.warning(f"Could not create DataFrame from PyMuPDF table on page {page_num+1}, table {i}: {e_df}")

                        if raw_data or df is not None: # Add if we have raw_data or a df
                            table_entry = {
                                "page": page_num + 1,
                                "table_id": f"table_pymupdf_{page_num + 1}_{i + 1}",
                                "dataframe": df,
                                "raw_data": raw_data,
                                "rect": list(table_obj.bbox), # Bbox as list
                                "html": df.to_html(index=False) if df is not None else pd.DataFrame(raw_data).to_html(index=False, header=False) if raw_data else ""
                            }
                            tables.append(table_entry)
            except Exception as e_page_table:
                 logger.error(f"Error processing tables on page {page_num+1} with PyMuPDF: {e_page_table}")


        if not tables and _table_extractor_v1: # Fallback to version1.0.0 table_extraction
            logger.info(f"PyMuPDF found no tables. Trying fallback table_extraction from version1.0.0 for {pdf_path}")
            doc = open_pdf(pdf_path) # Re-open doc if necessary
            for page_num, page in enumerate(doc):
                try:
                    table_chunks_v1 = _table_extractor_v1(page, page_num) # Call the imported v1 function
                    for i, chunk in enumerate(table_chunks_v1):
                        if "Table:" in chunk and not chunk.startswith("[No tables") and not chunk.startswith("[Error"):
                            table_text = chunk.split("Table:", 1)[1].strip()
                            rows_v1 = [list(map(str.strip, line.split("|"))) for line in table_text.split("\n") if "|" in line]
                            df_v1 = None
                            if rows_v1:
                                try:
                                    if len(rows_v1) > 1 and rows_v1[0]:
                                        df_v1 = pd.DataFrame(rows_v1[1:], columns=rows_v1[0])
                                    elif rows_v1:
                                        df_v1 = pd.DataFrame(rows_v1)
                                except Exception as e_df_v1:
                                    logger.warning(f"Could not create DataFrame from v1 extracted table: {e_df_v1}")
                            
                            if rows_v1 or df_v1 is not None:
                                tables.append({
                                    "page": page_num + 1,
                                    "table_id": f"table_v1_{page_num + 1}_{i + 1}",
                                    "dataframe": df_v1,
                                    "raw_data": rows_v1,
                                    "html": df_v1.to_html(index=False) if df_v1 is not None else pd.DataFrame(rows_v1).to_html(index=False,header=False) if rows_v1 else "",
                                    "extraction_method": "v1_fallback"
                                })
                except Exception as e_v1_page:
                    logger.error(f"Error with v1 table extraction on page {page_num+1}: {e_v1_page}")
        
        logger.info(f"Extracted {len(tables)} tables from PDF: {pdf_path}")
        return tables
    except Exception as e:
        logger.error(f"Error extracting tables from PDF {pdf_path}: {e}", exc_info=True)
        return []


def detect_tables_from_docx(docx_path):
    tables = []
    try:
        import docx # Import here to keep it local to this function
        doc = docx.Document(docx_path)
        for i, table_obj in enumerate(doc.tables):
            rows_data = []
            for row in table_obj.rows:
                rows_data.append([cell.text for cell in row.cells])
            
            df = None
            if rows_data:
                try:
                    if len(rows_data) > 1 and rows_data[0]:
                        df = pd.DataFrame(rows_data[1:], columns=rows_data[0])
                    elif rows_data:
                         df = pd.DataFrame(rows_data)
                except Exception as e_df:
                    logger.warning(f"Could not create DataFrame for DOCX table {i}: {e_df}")

            if rows_data or df is not None:
                tables.append({
                    "table_id": f"table_docx_{i + 1}",
                    "dataframe": df,
                    "raw_data": rows_data,
                    "html": df.to_html(index=False) if df is not None else pd.DataFrame(rows_data).to_html(index=False, header=False) if rows_data else ""
                })
        logger.info(f"Detected {len(tables)} tables in DOCX: {docx_path}")
        return tables
    except Exception as e:
        logger.error(f"Error detecting tables from DOCX {docx_path}: {e}", exc_info=True)
        return []


def detect_figures_from_pdf(pdf_path):
    figures = []
    try:
        doc = open_pdf(pdf_path) # Uses open_pdf from fixes or internal
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image or not base_image["image"]: continue # Skip if no image data

                    image_bytes = base_image["image"]
                    image_type = base_image["ext"].lower()
                    
                    pil_img = Image.open(io.BytesIO(image_bytes))
                    width, height = pil_img.size
                    img_b64 = base64.b64encode(image_bytes).decode('utf-8')
                    
                    figures.append({
                        "page": page_num + 1,
                        "figure_id": f"figure_pymupdf_{page_num + 1}_{img_index + 1}",
                        "image_bytes": image_bytes,
                        "rect": list(page.get_image_bbox(img_info).irect) if page.get_image_bbox(img_info) else None, # Useirect for int coords
                        "width": width, "height": height, "image_type": image_type,
                        "preview_src": f"data:image/{image_type};base64,{img_b64}",
                        "type": "image"
                    })
                except Exception as e_img:
                    logger.warning(f"Error extracting image {xref} on page {page_num+1}: {e_img}")

        if _chart_extractor_v1: # Fallback to version1.0.0 chart_extraction
            logger.info(f"Trying fallback chart_extraction from version1.0.0 for {pdf_path}")
            doc = open_pdf(pdf_path) # Re-open
            for page_num, page in enumerate(doc):
                try:
                    chart_texts_v1 = _chart_extractor_v1(page, page_num) # Call imported v1 function
                    for i, chart_text_entry in enumerate(chart_texts_v1):
                        if "Chart:" in chart_text_entry and not "[No charts detected]" in chart_text_entry and not "[Error:" in chart_text_entry:
                            chart_content = chart_text_entry.split("Chart:",1)[1].strip()
                            try: # Try to get page image for chart visual
                                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # 2x zoom
                                chart_img_bytes = pix.tobytes("png")
                                chart_pil_img = Image.open(io.BytesIO(chart_img_bytes))
                                chart_width, chart_height = chart_pil_img.size
                                chart_b64 = base64.b64encode(chart_img_bytes).decode('utf-8')
                                figures.append({
                                    "page": page_num + 1,
                                    "figure_id": f"chart_v1_{page_num + 1}_{i + 1}",
                                    "image_bytes": chart_img_bytes, "rect": None,
                                    "width": chart_width, "height": chart_height, "image_type": "png",
                                    "preview_src": f"data:image/png;base64,{chart_b64}",
                                    "type": "chart", "chart_text": chart_content,
                                    "extraction_method": "v1_fallback"
                                })
                            except Exception as e_chart_img:
                                logger.warning(f"Failed to create chart image from page {page_num+1} for v1 chart: {e_chart_img}")
                                # Add chart text even if image fails
                                figures.append({
                                    "page": page_num + 1,
                                    "figure_id": f"chart_v1_text_{page_num + 1}_{i + 1}",
                                    "image_bytes": None, "type": "chart_text_only", 
                                    "chart_text": chart_content, "extraction_method": "v1_fallback"
                                })

                except Exception as e_v1_chart_page:
                     logger.error(f"Error with v1 chart extraction on page {page_num+1}: {e_v1_chart_page}")

        logger.info(f"Extracted {len(figures)} figures/charts from PDF: {pdf_path}")
        return figures
    except Exception as e:
        logger.error(f"Error extracting figures from PDF {pdf_path}: {e}", exc_info=True)
        return []


def detect_figures_from_docx(docx_path):
    figures = []
    # python-docx does not robustly extract images with original bytes and type easily.
    # A common workaround is to convert DOCX to PDF (e.g., using LibreOffice CLI or MS Office COM)
    # or to unzip the .docx and parse relationships, which is complex.
    # For simplicity, this placeholder will be limited.
    # The original code used PyMuPDF on the .docx, which is unconventional but might work if PyMuPDF's DOCX filter is good.
    logger.warning("Robust figure extraction from DOCX is complex. This is a simplified version.")
    try:
        # Attempting the PyMuPDF method from original code
        doc_pdf_equivalent = open_pdf(docx_path) # PyMuPDF can try to open DOCX
        for page_num, page in enumerate(doc_pdf_equivalent):
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc_pdf_equivalent.extract_image(xref)
                    if base_image and base_image["image"]:
                        image_bytes = base_image["image"]
                        image_type = base_image["ext"].lower()
                        pil_img = Image.open(io.BytesIO(image_bytes))
                        width, height = pil_img.size
                        img_b64 = base64.b64encode(image_bytes).decode('utf-8')

                        figures.append({
                            "page": page_num + 1, # Page num might be less relevant for DOCX source
                            "figure_id": f"figure_docx_{page_num + 1}_{img_index + 1}",
                            "image_bytes": image_bytes, "image_type": image_type,
                            "width": width, "height": height,
                            "preview_src": f"data:image/{image_type};base64,{img_b64}",
                            "type": "image"
                        })
                except Exception as e_img:
                    logger.warning(f"Error extracting image {xref} from DOCX (via PyMuPDF) on page {page_num+1}: {e_img}")
        logger.info(f"Detected {len(figures)} figures in DOCX (via PyMuPDF attempt): {docx_path}")
    except Exception as e:
        logger.error(f"Error detecting figures from DOCX {docx_path}: {e}", exc_info=True)
    return figures


def process_document_master(file_path, file_type, user_tier_ocr_languages, unstructured_enabled_flag):
    """Master function to process document, deciding between unstructured and standard."""
    # Try using unstructured.io first if available and enabled
    if unstructured_enabled_flag and UNSTRUCTURED_AVAILABLE_PARSER:
        logger.info(f"Attempting to process {file_path} with unstructured.io")
        # Assuming UserTier.get_tier_features()[user_tier]["ocr_languages"] is passed as user_tier_ocr_languages
        doc_chunks, tables, figures, error = process_with_unstructured(file_path, file_type, user_tier_ocr_languages)
        if error is None and doc_chunks:
            return doc_chunks, tables, figures, None # error is the 4th element
        else:
            logger.warning(f"Unstructured processing failed: {error}. Falling back to standard processing.")
    
    # Fall back to standard processing
    logger.info(f"Processing {file_path} with standard methods.")
    tables, figures = [], []
    loader = None
    if file_type == "pdf":
        loader = PyMuPDFLoader(file_path)
        tables = detect_tables_from_pdf(file_path)
        figures = detect_figures_from_pdf(file_path)
    elif file_type == "docx":
        loader = Docx2txtLoader(file_path)
        tables = detect_tables_from_docx(file_path)
        figures = detect_figures_from_docx(file_path) # Simplified
    elif file_type == "txt":
        loader = TextLoader(file_path)
        # Text files typically don't have structured tables/figures in the same way
    elif file_type in ["png", "jpg", "jpeg", "tiff"]: # Handle images as documents
        # For images, if unstructured failed, we might try Tesseract directly or just return empty text
        # This part needs more thought for a non-unstructured image pathway
        logger.warning(f"Standard processing for image type '{file_type}' without unstructured is basic.")
        # Create a single LangchainDocument with a placeholder or basic OCR if implemented
        # For now, let's assume an image processed this way yields no text chunks if unstructured fails.
        return [LangchainDocument(page_content=f"[Image file: {os.path.basename(file_path)} - content not extracted by standard means]", metadata={"source": file_path})], [], [], None

    else:
        return [], [], [], "Unsupported file type for standard processing"

    if loader is None: # Should not happen if file_type is supported
         return [], [], [], "Loader could not be initialized."

    try:
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        doc_chunks = text_splitter.split_documents(documents)
        
        # Add table/figure text to chunks or as metadata if desired - current setup extracts separately
        # For simplicity, keeping table/figure data separate from main text chunks from PyMuPDFLoader etc.
        
        logger.info(f"Standard processing of {file_path} resulted in {len(doc_chunks)} chunks, {len(tables)} tables, {len(figures)} figures.")
        return doc_chunks, tables, figures, None
    except Exception as e:
        logger.error(f"Error in standard document processing for {file_path}: {e}", exc_info=True)
        return [], [], [], f"Error in standard document processing: {str(e)}"

def assign_figure_table_numbers(figures, tables, chapters):
    """Assign figure and table numbers based on chapter structure."""
    # This is a simplified numbering based on chapter and order of appearance.
    # A more robust system would involve analyzing text context or element positions.
    if not chapters: chapters = [{"number": 1, "title": "Document", "subsections": []}]
    
    # Initialize chapter-based counters
    figure_counters = {ch["number"]: 0 for ch in chapters}
    table_counters = {ch["number"]: 0 for ch in chapters}

    numbered_figures = []
    for fig in sorted(figures, key=lambda x: (x.get("page", 0), x.get("rect")[1] if x.get("rect") else 0)): # Sort by page, then Y position
        page = fig.get("page", 1)
        # Simplistic: assign to the first chapter whose page range MIGHT cover this.
        # This needs a better way to map content to chapters.
        # For now, we'll assume figures are implicitly associated with chapters by their order or page.
        # A more advanced method: Find the chapter whose text content is nearest to the figure.
        # Here, we will default to chapter 1 if no specific logic.
        assigned_chapter_num = fig.get("chapter_number") # If already assigned by some other logic
        if not assigned_chapter_num:
            # Simplistic: Find chapter by page. This is a rough heuristic.
            # It's better if figures/tables could be associated with sections during parsing.
            # For now, let's just number sequentially within a default chapter or by page.
            # A better approach: if chapters have page ranges, use that.
            # If not, assign to chapter 1 or use simple sequential.
            current_doc_chapter = 1 # Default
            # This logic is very basic and might not correctly assign chapters.
            # It would be better to associate figures/tables with specific text chunks/sections during parsing.
            for chapter_info in chapters:
                # This needs actual page ranges for chapters if we want to use page numbers.
                # Placeholder: if fig page > some_threshold, assign to next chapter.
                pass # No good heuristic here without more info on chapter spans.
            assigned_chapter_num = current_doc_chapter # Default for now
        
        if assigned_chapter_num not in figure_counters: # Ensure counter exists
            figure_counters[assigned_chapter_num] = 0
            
        figure_counters[assigned_chapter_num] += 1
        fig_num = figure_counters[assigned_chapter_num]
        
        fig_copy = fig.copy()
        fig_copy["chapter_number"] = assigned_chapter_num
        fig_copy["sequence_number"] = fig_num
        fig_copy["full_number"] = f"{assigned_chapter_num}.{fig_num}"
        fig_copy["caption"] = fig.get("caption_prefix", "Figure") + f" {assigned_chapter_num}.{fig_num}" + (f": {fig.get('caption_text', '')}" if fig.get('caption_text') else "")
        numbered_figures.append(fig_copy)

    numbered_tables = []
    for tbl in sorted(tables, key=lambda x: (x.get("page", 0), x.get("rect")[1] if x.get("rect") else 0)):
        page = tbl.get("page", 1)
        assigned_chapter_num = tbl.get("chapter_number")
        if not assigned_chapter_num:
            current_doc_chapter = 1 # Default
            assigned_chapter_num = current_doc_chapter

        if assigned_chapter_num not in table_counters:
             table_counters[assigned_chapter_num] = 0

        table_counters[assigned_chapter_num] += 1
        tbl_num = table_counters[assigned_chapter_num]
        
        tbl_copy = tbl.copy()
        tbl_copy["chapter_number"] = assigned_chapter_num
        tbl_copy["sequence_number"] = tbl_num
        tbl_copy["full_number"] = f"{assigned_chapter_num}.{tbl_num}"
        tbl_copy["caption"] = tbl.get("caption_prefix", "Table") + f" {assigned_chapter_num}.{tbl_num}" + (f": {tbl.get('caption_text', '')}" if tbl.get('caption_text') else "")
        numbered_tables.append(tbl_copy)
        
    return numbered_figures, numbered_tables