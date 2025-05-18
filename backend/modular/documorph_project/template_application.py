import os
import logging
import docx # python-docx
from docx.shared import Pt, Inches, RGBColor # Added RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING # Added WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE # For creating styles
from docx.oxml import OxmlElement # For TOC
from docx.oxml.ns import qn # For TOC

import tempfile
from PIL import Image as PIL_Image # Aliased to avoid conflict if Image is used elsewhere
import io
import cv2 # OpenCV for image conversion fallback
import numpy as np
import pandas as pd

# ReportLab imports (conditional for PDF generation)
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, PageBreak, Table as ReportLabTable, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.lib.units import inch as reportlab_inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    # logger.warning("ReportLab not installed. PDF generation will be basic or unavailable.") # Logger not defined here yet

# WeasyPrint imports (conditional for advanced PDF generation)
WEASYPRINT_AVAILABLE = False
try:
    from weasyprint import HTML, CSS
    from weasyprint.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    # logger.warning("WeasyPrint not installed. Advanced HTML to PDF generation will be unavailable.")
except Exception as e:
    WEASYPRINT_AVAILABLE = False
    # logger.warning(f"Error loading WeasyPrint: {e}. Advanced HTML to PDF generation will be unavailable.")


# Initialize logger for this module
logger = logging.getLogger("DocuMorphAI")


def get_color_from_scheme(template_config, element_type="body_text"):
    """
    Returns an RGBColor object based on the template's color scheme.
    element_type can be 'body_text', 'heading1', 'accent', etc.
    """
    scheme = template_config.get("color_scheme", "Default (Black/White)")
    # Define color mappings (hex strings)
    color_map = {
        "Default (Black/White)": {"body_text": "000000", "heading1": "000000", "accent": "444444"},
        "Blue Professional": {"body_text": "333333", "heading1": "004080", "accent": "0059b3"},
        "Modern Gray": {"body_text": "2c3e50", "heading1": "34495e", "accent": "7f8c8d"},
        "Earthy Tones": {"body_text": "5D4037", "heading1": "3E2723", "accent": "A1887F"},
        "Academic": {"body_text": "212121", "heading1": "000000", "accent": "555555"},
    }
    hex_color = color_map.get(scheme, color_map["Default (Black/White)"]).get(element_type, "000000")
    return RGBColor.from_string(hex_color)


def apply_template_to_document(
    source_doc_obj: docx.document.Document, # Expects a python-docx Document object
    template_config: dict,
    output_docx_path: str,
    tables_data: list = None,
    figures_data: list = None,
    chapters_data: list = None,
    document_title_override: str = None, # For dynamic title from Persona/LLM
    abstract_text: str = None # For inserting LLM-generated abstract
):
    """
    Applies a formatting template to a python-docx Document object and saves it.
    This function now primarily targets DOCX output. PDF generation will be a separate step.
    """
    try:
        # Start with a new document to ensure clean application of template styles
        # Content will be transferred from source_doc_obj
        output_doc = docx.Document()

        # 0. Clear default paragraphs in the new document if any
        for para in output_doc.paragraphs[:]:
            p_element = para._element
            p_element.getparent().remove(p_element)

        # --- Apply Document-Wide Settings from Template ---
        section = output_doc.sections[0] # Get the first section
        section.page_height = Inches(float(template_config.get('page_height', 11)))
        section.page_width = Inches(float(template_config.get('page_width', 8.5)))
        section.left_margin = Inches(float(template_config.get('margin_left', 1)))
        section.right_margin = Inches(float(template_config.get('margin_right', 1)))
        section.top_margin = Inches(float(template_config.get('margin_top', 1)))
        section.bottom_margin = Inches(float(template_config.get('margin_bottom', 1)))
        # Add orientation if specified
        if template_config.get('orientation', 'portrait').lower() == 'landscape':
            section.orientation = WD_ORIENTATION.LANDSCAPE
            # Swap page height/width if orientation changes, as python-docx expects these for the new orientation
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT


        # --- Define Styles Based on Template ---
        styles = output_doc.styles

        # Normal Style (Body Text)
        try:
            normal_style = styles['Normal']
        except KeyError: # Should not happen for 'Normal' but good practice
            normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = template_config.get('body_font', 'Calibri')
        normal_font.size = Pt(int(template_config.get('body_font_size', 11)))
        normal_font.color.rgb = get_color_from_scheme(template_config, "body_text")
        normal_pp = normal_style.paragraph_format
        normal_pp.line_spacing = float(template_config.get('line_spacing', 1.15))
        normal_pp.space_after = Pt(int(template_config.get('paragraph_spacing_after', 8))) # Default paragraph spacing
        if template_config.get('first_line_indent', 0.0) > 0:
            normal_pp.first_line_indent = Inches(float(template_config.get('first_line_indent')))


        # Heading Styles (H1-H3)
        for i in range(1, 4):
            style_name = f'Heading {i}'
            try:
                heading_style = styles[style_name]
            except KeyError:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_style.base_style = normal_style # Base it on Normal

            h_font = heading_style.font
            h_font.name = template_config.get(f'heading{i}_font', normal_font.name)
            h_font.size = Pt(int(template_config.get(f'heading{i}_font_size', 16 - (i * 2))))
            h_font.bold = template_config.get(f'heading{i}_bold', True if i == 1 else False)
            h_font.italic = template_config.get(f'heading{i}_italic', False)
            h_font.color.rgb = get_color_from_scheme(template_config, f"heading{i}")

            h_pp = heading_style.paragraph_format
            h_pp.line_spacing = float(template_config.get(f'heading{i}_line_spacing', normal_pp.line_spacing if normal_pp.line_spacing else 1.0)) # Default to normal line spacing or 1.0
            h_pp.space_before = Pt(int(template_config.get(f'heading{i}_space_before', 12 if i==1 else 6)))
            h_pp.space_after = Pt(int(template_config.get(f'heading{i}_space_after', 6 if i==1 else 3)))
            align_str = template_config.get(f'heading{i}_align', 'left').upper()
            if align_str == "CENTER": h_pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_str == "RIGHT": h_pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: h_pp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Caption Style
        try:
            caption_style = styles['Caption']
        except KeyError:
            caption_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.base_style = normal_style
        caption_font = caption_style.font
        caption_font.size = Pt(int(template_config.get('caption_font_size', 9)))
        caption_font.italic = template_config.get('caption_italic', True)
        caption_font.color.rgb = get_color_from_scheme(template_config, "accent")
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_before = Pt(3)
        caption_style.paragraph_format.space_after = Pt(6)


        # --- Add Content Elements ---

        # Title Page (if enabled)
        if template_config.get('title_page_enabled', False):
            title_text = document_title_override if document_title_override else template_config.get('title_page_title', "Document Title")
            title_para = output_doc.add_paragraph()
            title_run = title_para.add_run(title_text)
            # Apply specific title font from template, or default to a larger heading style
            title_font_name = template_config.get('title_font', h_font.name) # Use H1 font if title_font not set
            title_font_size = Pt(int(template_config.get('title_font_size', 24)))
            title_run.font.name = title_font_name
            title_run.font.size = title_font_size
            title_run.font.bold = template_config.get('title_font_bold', True)
            title_run.font.color.rgb = get_color_from_scheme(template_config, "heading1") # Main title color
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(int(template_config.get('title_space_after', 24)))

            author = template_config.get('title_page_author', '')
            if author:
                author_para = output_doc.add_paragraph(author, style='Normal') # Apply normal style then customize
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_para.runs[0].font.italic = True # Example customization
                author_para.paragraph_format.space_after = Pt(12)

            affiliation = template_config.get('title_page_affiliation', '')
            if affiliation:
                 output_doc.add_paragraph(affiliation, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date if configured
            # Add other title page elements as per template_config

            output_doc.add_page_break()

        # Abstract (if provided by LLM)
        if abstract_text:
            output_doc.add_heading("Abstract", level=1) # Or a specific "Abstract Heading" style
            # Split abstract into paragraphs if it contains double newlines
            for abstract_para in abstract_text.split('\n\n'):
                output_doc.add_paragraph(abstract_para.strip(), style='Normal') # Apply normal style to abstract text
            output_doc.add_paragraph() # Extra space after abstract

        # Table of Contents (if enabled)
        if template_config.get('toc_enabled', False):
            output_doc.add_heading(template_config.get('toc_title', "Table of Contents"), level=1)
            # Adding a proper TOC field requires complex OXML manipulation.
            # This is a placeholder that Word can update.
            toc_paragraph = output_doc.add_paragraph()
            run = toc_paragraph.add_run()
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            # TOC levels from template, default 1-3
            toc_levels_str = f"1-{template_config.get('toc_levels', 3)}"
            instrText.text = f'TOC \\o "{toc_levels_str}" \\h \\z \\u'
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            
            toc_paragraph._element.append(fldChar_begin)
            toc_paragraph._element.append(instrText)
            toc_paragraph._element.append(fldChar_end)
            output_doc.add_page_break()

        # --- Transfer and Style Content from source_doc_obj ---
        # This needs to be more intelligent if source_doc_obj has rich formatting.
        # For now, assuming source_doc_obj provides paragraphs that need styling.
        # If source_doc_obj is from PDF/TXT, it's likely just a sequence of paragraphs.
        # If from DOCX, it has existing (but unstyled by template) paragraphs.

        # A more robust way to transfer content while respecting some structure:
        for element in source_doc_obj.element.body:
            if element.tag.endswith('p'):
                # Reconstruct paragraph from element to handle its content correctly
                para = docx.text.paragraph.Paragraph(element, source_doc_obj)
                text = para.text
                
                # Heuristic to determine if it's a heading from source,
                # or if LLM structure analysis should override.
                # This part is complex: deciding original style vs. new template style vs. LLM structure.
                # Simplification: if text matches a known heading style name from source, map it.
                # Otherwise, apply 'Normal' or use LLM-derived structure if available.

                # For now, just copy text with 'Normal' style.
                # LLM-identified headings should ideally be inserted into this flow.
                if text.strip(): # Add non-empty paragraphs
                    output_doc.add_paragraph(text, style='Normal')
            
            elif element.tag.endswith('tbl'):
                # Table transfer is more complex; needs to reconstruct table with new styling.
                # For now, if tables_data is provided, those will be added later.
                # If not, this basic transfer won't preserve tables well.
                logger.debug("Skipping direct table element transfer; tables handled separately via tables_data.")
                pass


        # --- Add Processed Tables and Figures ---
        # (Assuming chapters_data has been populated, e.g., by LLM or rule-based)
        if not chapters_data:
            chapters_data = [{"number": 1, "title": "Main Document", "subsections": []}]

        numbered_figures, numbered_tables = assign_figure_table_numbers(
            figures_data if figures_data else [],
            tables_data if tables_data else [],
            chapters_data
        )

        # Option to place tables/figures in appendix or intersperse (complex)
        # For now, adding to appendix-like sections
        if template_config.get('include_tables_figures', True):
            if numbered_tables:
                output_doc.add_page_break()
                output_doc.add_heading(template_config.get('tables_appendix_title', "Tables"), level=1)
                for table_item in numbered_tables:
                    enhanced_add_table_to_docx(output_doc, table_item, table_item.get("caption", f"Table {table_item.get('full_number')}"))
                    output_doc.add_paragraph()

            if numbered_figures:
                output_doc.add_page_break()
                output_doc.add_heading(template_config.get('figures_appendix_title', "Figures"), level=1)
                for figure_item in numbered_figures:
                    enhanced_add_figure_to_docx(output_doc, figure_item, figure_item.get("caption", f"Figure {figure_item.get('full_number')}"))
                    output_doc.add_paragraph()


        # --- Add Headers/Footers (Final Pass) ---
        # (This was already done for the section, but re-iterating ensures all sections if more are added)
        header_text_template = template_config.get('header_text', '')
        footer_text_template = template_config.get('footer_text', '')
        header_align_str = template_config.get('header_align', 'center').upper()
        footer_align_str = template_config.get('footer_align', 'center').upper()

        for section in output_doc.sections:
            if header_text_template:
                header = section.header
                if not header.paragraphs: header.add_paragraph()
                ht_para = header.paragraphs[0]
                # Clear existing runs to avoid appending
                for run in ht_para.runs: run._r.getparent().remove(run._r)
                # Smart field replacement
                ht_para.text = _replace_doc_fields(header_text_template, output_doc)
                if header_align_str == "LEFT": ht_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif header_align_str == "RIGHT": ht_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: ht_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if footer_text_template:
                footer = section.footer
                if not footer.paragraphs: footer.add_paragraph()
                ft_para = footer.paragraphs[0]
                for run in ft_para.runs: run._r.getparent().remove(run._r)
                ft_para.text = _replace_doc_fields(footer_text_template, output_doc)
                if footer_align_str == "LEFT": ft_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif footer_align_str == "RIGHT": ft_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: ft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


        if not output_doc.paragraphs and not output_doc.tables:
             output_doc.add_paragraph("Document content appears to be empty after processing.")
        
        output_doc.save(output_docx_path)
        logger.info(f"Formatted DOCX saved to: {output_docx_path}")
        return True, output_docx_path

    except Exception as e:
        logger.error(f"Error applying template to generate DOCX: {e}", exc_info=True)
        return False, f"Error generating DOCX: {str(e)}"


def _replace_doc_fields(text, doc=None, current_page=None, total_pages=None):
    """Helper to replace placeholders like [PageNumber] in header/footer text."""
    # Actual page number and total pages are tricky to get accurately before final rendering
    # For DOCX, Word handles these fields. For PDF, ReportLab needs callbacks.
    text = text.replace("[PageNumber]", "PAGE_FIELD") # Placeholder for Word PAGE field
    text = text.replace("[TotalPages]", "NUMPAGES_FIELD") # Placeholder for Word NUMPAGES field
    text = text.replace("[Page]", "PAGE_FIELD")
    text = text.replace("[Pages]", "NUMPAGES_FIELD")
    text = text.replace("[Date]", java.time.LocalDate.now().format(java.time.format.DateTimeFormatter.ofLocalizedDate(java.time.format.FormatStyle.MEDIUM))) # Java for platform-agnostic date
    text = text.replace("[CurrentYear]", str(java.time.LocalDate.now().getYear()))
    # More fields can be added: [Last Name], [Company Name], [Document Title], [Chapter]
    # These would need to be passed into this function or accessed from a global context/doc properties.
    return text

def assign_figure_table_numbers(figures_data, tables_data, chapters_data):
    """Assigns hierarchical numbers (e.g., 1.1, 1.2, 2.1) to figures and tables based on chapters."""
    # (This is a refined version of the original logic)
    
    numbered_figures = []
    numbered_tables = []

    if not chapters_data: # If no chapters, number sequentially 1, 2, 3...
        for i, fig in enumerate(figures_data):
            fig['full_number'] = str(i + 1)
            fig['caption'] = fig.get('caption_prefix', 'Figure') + f" {fig['full_number']}: " + fig.get('text_content', fig.get('text_around_image', 'Untitled Figure'))
            numbered_figures.append(fig)
        for i, tbl in enumerate(tables_data):
            tbl['full_number'] = str(i + 1)
            tbl['caption'] = tbl.get('caption_prefix', 'Table') + f" {tbl['full_number']}: " + tbl.get('title', 'Untitled Table')
            numbered_tables.append(tbl)
        return numbered_figures, numbered_tables

    # If chapters exist, attempt chapter-based numbering
    # This requires figures/tables to have reliable 'page' numbers and chapters to have 'start_page' (approx)
    # For simplicity, if 'page' and 'start_page' are not reliably available, this will fall back to simpler logic.

    # Create a map of chapter number to a counter for figures and tables
    fig_counters = {ch['number']: 0 for ch in chapters_data}
    tbl_counters = {ch['number']: 0 for ch in chapters_data}

    # A very simplified way to assign chapter numbers if not already present.
    # Assumes figures/tables are somewhat ordered or have page numbers.
    # A more robust method would involve analyzing text context around figures/tables.
    def get_chapter_for_element(element_page, chapters_list):
        if not element_page: return chapters_list[0]['number'] # Default to first chapter
        # This requires chapters_list to have page range info, which is complex to get accurately.
        # For now, we'll assume if 'chapter_number' is already in element, use it. Else, default.
        return element_page # Placeholder logic, needs refinement if page ranges are used
    

    for fig in figures_data:
        # Attempt to determine chapter:
        # 1. If fig has 'chapter_number' already (e.g. from previous analysis)
        # 2. If fig has 'page' and chapters_data have page ranges (complex)
        # 3. Default or heuristic
        assigned_chapter_num = fig.get('chapter_number')
        if not assigned_chapter_num:
            # Placeholder: Assign based on order or a very rough page estimate if available
            # This part needs to be made much smarter for accurate chapter association
            assigned_chapter_num = chapters_data[0]['number'] # Fallback to first chapter

        fig_counters[assigned_chapter_num] = fig_counters.get(assigned_chapter_num, 0) + 1
        fig['full_number'] = f"{assigned_chapter_num}.{fig_counters[assigned_chapter_num]}"
        fig['caption'] = fig.get('caption_prefix', 'Figure') + f" {fig['full_number']}: " + fig.get('text_content', fig.get('text_around_image', 'Untitled Figure'))
        numbered_figures.append(fig)

    for tbl in tables_data:
        assigned_chapter_num = tbl.get('chapter_number')
        if not assigned_chapter_num:
            assigned_chapter_num = chapters_data[0]['number']

        tbl_counters[assigned_chapter_num] = tbl_counters.get(assigned_chapter_num, 0) + 1
        tbl['full_number'] = f"{assigned_chapter_num}.{tbl_counters[assigned_chapter_num]}"
        tbl['caption'] = tbl.get('caption_prefix', 'Table') + f" {tbl['full_number']}: " + tbl.get('title', 'Untitled Table')
        numbered_tables.append(tbl)
        
    return numbered_figures, numbered_tables


# --- PDF Generation Functions (using ReportLab, with attempts for better image handling) ---

def generate_formatted_pdf_reportlab(
    source_doc_obj: docx.document.Document, # Content source
    template_config: dict,
    output_pdf_path: str,
    tables_data: list = None,
    figures_data: list = None,
    chapters_data: list = None, # For potential chapter headings in PDF
    document_title_override: str = None,
    abstract_text: str = None
):
    if not REPORTLAB_AVAILABLE:
        logger.error("ReportLab is not installed. Cannot generate PDF.")
        return False, "ReportLab library not found."

    try:
        # Page setup for ReportLab
        page_width_inch = float(template_config.get('page_width', 8.5))
        page_height_inch = float(template_config.get('page_height', 11))
        
        # Handle orientation
        is_landscape = template_config.get('orientation', 'portrait').lower() == 'landscape'
        pagesize = (page_height_inch * reportlab_inch, page_width_inch * reportlab_inch) if is_landscape \
                   else (page_width_inch * reportlab_inch, page_height_inch * reportlab_inch)

        pdf_doc = SimpleDocTemplate(
            output_pdf_path,
            pagesize=pagesize,
            leftMargin=float(template_config.get('margin_left', 1)) * reportlab_inch,
            rightMargin=float(template_config.get('margin_right', 1)) * reportlab_inch,
            topMargin=float(template_config.get('margin_top', 1)) * reportlab_inch,
            bottomMargin=float(template_config.get('margin_bottom', 1)) * reportlab_inch,
            title=document_title_override or template_config.get('title_page_title', "Formatted Document")
        )
        
        story = []
        styles = getSampleStyleSheet()

        # --- Define ReportLab Styles from template_config ---
        # Body Text Style
        body_font_name_rl = template_config.get('body_font_rl', 'Helvetica') # Ensure RL compatible font name
        body_font_size_rl = int(template_config.get('body_font_size', 11))
        body_color_hex = get_color_from_scheme(template_config, "body_text").to_rgb_string()[1:] # Get hex without #
        
        body_style_rl = ParagraphStyle(
            'BodyRL', parent=styles['Normal'],
            fontName=body_font_name_rl, fontSize=body_font_size_rl,
            leading=body_font_size_rl * float(template_config.get('line_spacing_rl', 1.2)), # RL line spacing
            textColor=colors.HexColor(f"#{body_color_hex}"),
            spaceAfter=Pt(int(template_config.get('paragraph_spacing_after', 8))),
            firstLineIndent=Inches(float(template_config.get('first_line_indent', 0.0))) * 72 # Convert inches to points
        )

        # Heading Styles
        heading_styles_rl = {}
        for i in range(1, 4):
            h_font_name = template_config.get(f'heading{i}_font_rl', body_font_name_rl)
            h_font_size = int(template_config.get(f'heading{i}_font_size', 16 - (i*2)))
            h_color_hex = get_color_from_scheme(template_config, f"heading{i}").to_rgb_string()[1:]
            h_align_str = template_config.get(f'heading{i}_align', 'left').upper()
            alignment_rl = TA_LEFT
            if h_align_str == "CENTER": alignment_rl = TA_CENTER
            elif h_align_str == "RIGHT": alignment_rl = TA_RIGHT
            
            heading_styles_rl[f'H{i}'] = ParagraphStyle(
                f'Heading{i}RL', parent=body_style_rl,
                fontName=h_font_name + ('-Bold' if template_config.get(f'heading{i}_bold', True if i==1 else False) else ''),
                fontSize=h_font_size,
                textColor=colors.HexColor(f"#{h_color_hex}"),
                leading=h_font_size * 1.2,
                spaceBefore=Pt(int(template_config.get(f'heading{i}_space_before', 12 if i==1 else 6))),
                spaceAfter=Pt(int(template_config.get(f'heading{i}_space_after', 6 if i==1 else 3))),
                alignment=alignment_rl
            )
        
        # Caption Style
        caption_font_name_rl = template_config.get('caption_font_rl', 'Helvetica')
        caption_font_size_rl = int(template_config.get('caption_font_size', 9))
        caption_color_hex = get_color_from_scheme(template_config, "accent").to_rgb_string()[1:]
        caption_style_rl = ParagraphStyle(
            'CaptionRL', parent=styles['Italic'],
            fontName=caption_font_name_rl + ('-Oblique' if template_config.get('caption_italic', True) else ''),
            fontSize=caption_font_size_rl,
            textColor=colors.HexColor(f"#{caption_color_hex}"),
            alignment=TA_CENTER,
            spaceBefore=Pt(3), spaceAfter=Pt(6)
        )

        # --- Build PDF Story ---
        # Title Page (simplified for PDF)
        if template_config.get('title_page_enabled', False):
            title_text = document_title_override or template_config.get('title_page_title', "Document Title")
            title_style = ParagraphStyle('TitleRL', parent=heading_styles_rl['H1'], fontSize=Pt(int(template_config.get('title_font_size', 24))), alignment=TA_CENTER, spaceAfter=Pt(24))
            story.append(Paragraph(title_text, title_style))
            
            author = template_config.get('title_page_author', '')
            if author: story.append(Paragraph(author, body_style_rl)) # Add more styling
            story.append(PageBreak())

        # Abstract
        if abstract_text:
            story.append(Paragraph("Abstract", heading_styles_rl['H1']))
            for p_text in abstract_text.split('\n\n'):
                story.append(Paragraph(p_text.strip(), body_style_rl))
            story.append(Spacer(1, 0.25 * reportlab_inch))

        # TOC (placeholder for ReportLab - true TOC is complex)
        if template_config.get('toc_enabled', False):
            story.append(Paragraph(template_config.get('toc_title', "Table of Contents"), heading_styles_rl['H1']))
            story.append(Paragraph("[Table of Contents - Manual update may be required in PDF viewer]", body_style_rl))
            story.append(PageBreak())


        # Main Content from source_doc_obj
        for para_obj in source_doc_obj.paragraphs:
            text = para_obj.text.strip()
            if not text:
                story.append(Spacer(1, body_style_rl.fontSize * 0.5)) # Represent empty lines
                continue
            
            # Basic style mapping (can be improved by inspecting para_obj.style.name)
            # This is a simplification. A full DOCX to ReportLab style conversion is very complex.
            style_to_apply_rl = body_style_rl
            if para_obj.style.name.startswith('Heading 1'): style_to_apply_rl = heading_styles_rl['H1']
            elif para_obj.style.name.startswith('Heading 2'): style_to_apply_rl = heading_styles_rl['H2']
            elif para_obj.style.name.startswith('Heading 3'): style_to_apply_rl = heading_styles_rl['H3']
            # Add more style mappings if needed

            story.append(Paragraph(text, style_to_apply_rl))
        
        # Tables and Figures (numbered and captioned)
        if template_config.get('include_tables_figures', True):
            if not chapters_data: chapters_data = [{"number": 1, "title": "Main Document", "subsections": []}]
            numbered_figures, numbered_tables = assign_figure_table_numbers(figures_data or [], tables_data or [], chapters_data)

            if numbered_tables:
                story.append(PageBreak())
                story.append(Paragraph(template_config.get('tables_appendix_title', "Tables"), heading_styles_rl['H1']))
                for table_item in numbered_tables:
                    story.append(Paragraph(table_item.get("caption", "Table"), caption_style_rl))
                    df = table_item.get("dataframe")
                    if df is not None and not df.empty:
                        # Convert DataFrame to list of lists for ReportLabTable
                        data_for_rl_table = [df.columns.tolist()] + df.values.tolist()
                        try:
                            rl_table_obj = ReportLabTable(data_for_rl_table, repeatRows=1) # Repeat header row
                            # Basic table styling
                            rl_table_obj.setStyle(TableStyle([
                                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                                ('FONTNAME', (0,0), (-1,0), body_font_name_rl + '-Bold'), # Header font
                                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                                ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                                ('GRID', (0,0), (-1,-1), 1, colors.black)
                            ]))
                            story.append(rl_table_obj)
                        except Exception as e_rl_tbl:
                            logger.error(f"Error creating ReportLabTable for {table_item.get('caption')}: {e_rl_tbl}")
                            story.append(Paragraph(f"[Table data for {table_item.get('caption')} could not be rendered]", body_style_rl))

                    story.append(Spacer(1, 0.2 * reportlab_inch))


            if numbered_figures:
                story.append(PageBreak())
                story.append(Paragraph(template_config.get('figures_appendix_title', "Figures"), heading_styles_rl['H1']))
                for fig_item in numbered_figures:
                    story.append(Paragraph(fig_item.get("caption", "Figure"), caption_style_rl))
                    if fig_item.get("image_bytes"):
                        try:
                            img_bytesio = io.BytesIO(fig_item["image_bytes"])
                            # Calculate width for ReportLab image
                            img_width_px = fig_item.get("width", 500)
                            # Max width in PDF document (considering margins)
                            max_img_width_points = pdf_doc.width - pdf_doc.leftMargin - pdf_doc.rightMargin - (0.2 * reportlab_inch) # Small buffer
                            img_width_points = min(img_width_px * 0.75, max_img_width_points) # Convert pixels to points (approx) and cap

                            rl_img = ReportLabImage(img_bytesio, width=img_width_points, height=None) # Auto height
                            rl_img.hAlign = 'CENTER'
                            story.append(rl_img)
                        except Exception as e_rl_img:
                            logger.error(f"Error creating ReportLabImage for {fig_item.get('caption')}: {e_rl_img}")
                            story.append(Paragraph(f"[Image for {fig_item.get('caption')} could not be rendered]", body_style_rl))
                    story.append(Spacer(1, 0.2 * reportlab_inch))


        # --- Build PDF with Header/Footer ---
        def header_footer_canvas(canvas, current_doc):
            canvas.saveState()
            # Header
            header_text_rl = template_config.get('header_text', '')
            if header_text_rl:
                page_num_str = str(canvas.getPageNumber())
                formatted_header = _replace_doc_fields(header_text_rl, current_page=page_num_str) # Simple replacement
                canvas.setFont(template_config.get('header_font_rl', 'Helvetica'), int(template_config.get('header_font_size_pt', 9)))
                header_x = current_doc.leftMargin
                if template_config.get('header_align','center').lower() == 'right':
                    header_x = current_doc.width - current_doc.rightMargin - canvas.stringWidth(formatted_header, template_config.get('header_font_rl', 'Helvetica'), int(template_config.get('header_font_size_pt', 9)))
                elif template_config.get('header_align','center').lower() == 'center':
                    header_x = (current_doc.width - canvas.stringWidth(formatted_header, template_config.get('header_font_rl', 'Helvetica'), int(template_config.get('header_font_size_pt', 9)))) / 2.0

                canvas.drawString(header_x, current_doc.height + current_doc.topMargin - 0.5*reportlab_inch, formatted_header)
            # Footer
            footer_text_rl = template_config.get('footer_text', '')
            if footer_text_rl:
                page_num_str = str(canvas.getPageNumber())
                # total_pages_str = str(getattr(current_doc, "page_count_total", page_num_str)) # Getting total pages in RL is tricky
                formatted_footer = _replace_doc_fields(footer_text_rl, current_page=page_num_str) # total_pages=total_pages_str)
                canvas.setFont(template_config.get('footer_font_rl', 'Helvetica'), int(template_config.get('footer_font_size_pt', 9)))
                
                footer_x = current_doc.leftMargin
                if template_config.get('footer_align','center').lower() == 'right':
                    footer_x = current_doc.width - current_doc.rightMargin - canvas.stringWidth(formatted_footer, template_config.get('footer_font_rl', 'Helvetica'), int(template_config.get('footer_font_size_pt', 9)))
                elif template_config.get('footer_align','center').lower() == 'center':
                     footer_x = (current_doc.width - canvas.stringWidth(formatted_footer, template_config.get('footer_font_rl', 'Helvetica'), int(template_config.get('footer_font_size_pt', 9)))) / 2.0

                canvas.drawString(footer_x, current_doc.bottomMargin - 0.3*reportlab_inch, formatted_footer)
            canvas.restoreState()

        pdf_doc.build(story, onFirstPage=header_footer_canvas, onLaterPages=header_footer_canvas)
        logger.info(f"Formatted PDF (ReportLab) saved to: {output_pdf_path}")
        return True, output_pdf_path

    except Exception as e:
        logger.error(f"Error generating PDF with ReportLab: {e}", exc_info=True)
        return False, f"Error generating PDF (ReportLab): {str(e)}"


# Placeholder for WeasyPrint PDF generation (more advanced, requires CSS)
def generate_formatted_pdf_weasyprint(
    html_content: str, # Expects fully styled HTML content
    template_config: dict,
    output_pdf_path: str
):
    if not WEASYPRINT_AVAILABLE:
        logger.error("WeasyPrint is not installed. Cannot generate PDF with this method.")
        return False, "WeasyPrint library not found."
    try:
        # Basic CSS from template (can be expanded significantly)
        # This would ideally come from a CSS file or more detailed template config
        font_config = FontConfiguration()
        css_rules = f"""
            @page {{
                size: {template_config.get('page_width', 8.5)}in {template_config.get('page_height', 11)}in;
                margin-top: {template_config.get('margin_top', 1)}in;
                margin-bottom: {template_config.get('margin_bottom', 1)}in;
                margin-left: {template_config.get('margin_left', 1)}in;
                margin-right: {template_config.get('margin_right', 1)}in;
            }}
            body {{
                font-family: "{template_config.get('body_font', 'Arial')}", sans-serif;
                font-size: {template_config.get('body_font_size', 11)}pt;
                line-height: {template_config.get('line_spacing', 1.15)};
                color: #{get_color_from_scheme(template_config, "body_text").to_rgb_string()[1:]};
            }}
            h1 {{
                font-family: "{template_config.get('heading1_font', template_config.get('body_font', 'Arial'))}", sans-serif;
                font-size: {template_config.get('heading1_font_size', 16)}pt;
                color: #{get_color_from_scheme(template_config, "heading1").to_rgb_string()[1:]};
                font-weight: {'bold' if template_config.get('heading1_bold', True) else 'normal'};
            }}
            /* Add more styles for h2, h3, p, table, img, caption etc. */
        """
        css = CSS(string=css_rules, font_config=font_config)
        HTML(string=html_content).write_pdf(output_pdf_path, stylesheets=[css], font_config=font_config)
        logger.info(f"Formatted PDF (WeasyPrint) saved to: {output_pdf_path}")
        return True, output_pdf_path
    except Exception as e:
        logger.error(f"Error generating PDF with WeasyPrint: {e}", exc_info=True)
        return False, f"Error generating PDF (WeasyPrint): {str(e)}"

# Added functions for table and figure handling
def enhanced_add_table_to_docx(doc, table_item, caption_text):
    """
    Add a table to the DOCX document with proper formatting.
    
    Args:
        doc: python-docx Document object
        table_item: Dictionary containing table data
        caption_text: Caption to add below the table
    
    Returns:
        bool: Success or failure
    """
    try:
        # Get data from table_item
        df = table_item.get("dataframe")
        raw_data = table_item.get("raw_data")
        
        if df is not None and not df.empty:
            # Convert DataFrame to list of lists for docx table
            table_data = [df.columns.tolist()] + df.values.tolist()
        elif raw_data:
            table_data = raw_data
        else:
            # No data available
            doc.add_paragraph(f"[Table data unavailable: {caption_text}]", style='Caption')
            return False
            
        # Create table in docx
        rows = len(table_data)
        cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Fill table with data
        for i, row_data in enumerate(table_data):
            for j, cell_value in enumerate(row_data):
                if j < cols:  # Ensure we don't exceed table bounds
                    cell = table.cell(i, j)
                    cell.text = str(cell_value)
                    # Style header row
                    if i == 0:
                        run = cell.paragraphs[0].runs
                        if run:
                            run[0].bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption_para = doc.add_paragraph(caption_text, style='Caption')
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return True
    except Exception as e:
        logger.error(f"Error adding table to DOCX: {e}")
        # Add error message as a paragraph
        doc.add_paragraph(f"[Error adding table: {str(e)}]")
        return False

def enhanced_add_figure_to_docx(doc, figure_item, caption_text):
    """
    Add a figure/image to the DOCX document with proper formatting.
    
    Args:
        doc: python-docx Document object
        figure_item: Dictionary containing figure data
        caption_text: Caption to add below the figure
    
    Returns:
        bool: Success or failure
    """
    try:
        # Get image bytes
        image_bytes = figure_item.get("image_bytes")
        if not image_bytes:
            # No image data
            doc.add_paragraph(f"[Figure unavailable: {caption_text}]", style='Caption')
            return False
            
        # Create BytesIO object from image bytes
        image_stream = io.BytesIO(image_bytes)
        
        # Add image to document
        width = figure_item.get("width", 500)  # Default width in pixels
        # Convert to inches (approximate)
        width_inches = min(width / 96, 6.0)  # Cap at 6 inches, assuming 96 DPI
        
        # Add picture with width
        doc.add_picture(image_stream, width=Inches(width_inches))
        
        # Add caption
        caption_para = doc.add_paragraph(caption_text, style='Caption')
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return True
    except Exception as e:
        logger.error(f"Error adding figure to DOCX: {e}")
        # Add error message as a paragraph
        doc.add_paragraph(f"[Error adding figure: {str(e)}]")
        return False