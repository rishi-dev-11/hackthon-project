import streamlit as st
import base64
import os
import docx
try:
    import fitz # PyMuPDF
except ImportError:
    # Fallback if PyMuPDF/fitz is not available
    fitz = None
    st.error("PyMuPDF (fitz) module is not available. Document preview functionality may be limited.")
from config import logger, UserTier, PERSONA_CATEGORIES # Import UserTier and PERSONAS

# Function to display document preview (largely same, ensure robustness)
def display_document_preview(doc_path):
    # ... (implementation from previous ui_helpers.py, ensure it's robust) ...
    # Minor refinement: check file existence more explicitly at start
    if not doc_path or not os.path.exists(doc_path):
        st.error(f"Preview Error: Document not found or path is invalid ('{doc_path}').")
        return False
    try:
        # ... (rest of the preview logic)
        # Ensure proper HTML styling for DOCX preview as before
        if doc_path.endswith('.docx'):
            doc = docx.Document(doc_path)
            if not doc.paragraphs:
                st.warning("DOCX for preview appears to be empty.")
                # Provide download anyway
                with open(doc_path, "rb") as fp:
                    st.download_button("Download DOCX (empty)", fp, os.path.basename(doc_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_empty_{os.path.basename(doc_path)}")
                return False # Indicate preview wasn't fully successful

            content_html_parts = []
            # ... (rest of DOCX preview HTML generation logic) ...
            for para in doc.paragraphs:
                text = para.text.replace('\n', '<br>')
                # Basic style detection for preview emphasis
                style_name = para.style.name.lower()
                if 'heading 1' in style_name: content_html_parts.append(f"<h1>{text}</h1>")
                elif 'heading 2' in style_name: content_html_parts.append(f"<h2>{text}</h2>")
                elif 'heading 3' in style_name: content_html_parts.append(f"<h3>{text}</h3>")
                else: content_html_parts.append(f"<p>{text}</p>")
            
            try: preview_font = doc.styles['Normal'].font.name if doc.styles['Normal'].font.name else 'sans-serif'
            except: preview_font = 'sans-serif'

            preview_html = f"""<div style="border:1px solid #ddd; padding:15px; border-radius:5px; height:450px; overflow-y:scroll; background-color:#ffffff; font-family:'{preview_font}'; color: #333;">{''.join(content_html_parts)}</div>"""
            st.markdown("### Document Preview (DOCX)")
            st.markdown(preview_html, unsafe_allow_html=True)

        elif doc_path.endswith('.pdf'):
            # ... PDF preview logic ...
            st.markdown("### Document Preview (PDF)")
            with open(doc_path, "rb") as file:
                base64_pdf = base64.b64encode(file.read()).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600px" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        elif doc_path.endswith('.txt'):
            # ... TXT preview logic ...
            st.markdown("### Document Preview (TXT)")
            with open(doc_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()
            st.text_area("Content", content, height=400, disabled=True)
        else:
            st.warning(f"Preview not available for this file type: {os.path.splitext(doc_path)[1]}")
        
        # Download button for the previewed file
        with open(doc_path, "rb") as fp_dl:
            st.download_button(
                label=f"Download Previewed: {os.path.basename(doc_path)}",
                data=fp_dl.read(),
                file_name=os.path.basename(doc_path),
                key=f"dl_preview_{os.path.basename(doc_path).replace('.', '_')}"
            )
        return True

    except Exception as e:
        logger.error(f"Error displaying document preview for '{doc_path}': {e}", exc_info=True)
        st.error(f"Could not generate preview for {os.path.basename(doc_path)}.")
        return False


def get_persona_specific_prompt(persona_selected, task_type):
    """
    Generate persona-specific prompts for LLM tasks.
    'persona_selected' should be one of the PERSONA_CATEGORIES.
    'task_type' can be 'abstract', 'section_titles', 'style', 'structure'.
    """
    # (This function was already well-defined in your original script,
    # ensure PERSONA_CATEGORIES is consistent with what's in config.py)
    
    # Default prompts provide a fallback
    default_prompts = {
        "abstract": "Generate a concise, professional abstract summarizing the main points of this document. The abstract should be well-structured, clear, and approximately 150-250 words.",
        "section_titles": "Suggest professional section titles that would improve the organization of this document. Titles should be clear, descriptive, and help guide the reader through the content.",
        "style": "Professional and Clear", # This is a style description for the LLM
        "structure": "Analyze this document and suggest an improved structure. Focus on logical organization, clear progression of ideas, and appropriate section divisions."
    }
    
    # Persona-specific prompts (expand this dictionary with detailed prompts for each persona and task)
    persona_prompts_map = {
        "Student": {
            "abstract": "Generate a concise, academic abstract for this student document. Focus on clarity, proper academic terminology, and highlighting the main arguments or findings. The abstract should be well-structured and approximately 150-250 words. Use formal language appropriate for academic submission.",
            "section_titles": "Suggest academic section titles appropriate for a student paper (e.g., Introduction, Literature Review, Methodology, Results, Discussion, Conclusion). Titles should be clear and descriptive.",
            "style": "Academic Formal",
            "structure": "Analyze this student document and suggest an academic structure with clear sections like Introduction, Body Paragraphs with supporting evidence, and Conclusion. Ensure logical flow."
        },
        "Content Creator": {
            "abstract": "Craft an engaging summary (around 100-150 words) for this piece of content. It should hook the reader, highlight the most interesting points, and make them want to read more. Use a slightly informal and inviting tone.",
            "section_titles": "Suggest 3-4 catchy and engaging section titles for this content that will spark curiosity. Titles should be relatively short and hint at the value within each section.",
            "style": "Engaging, Conversational, and Clear",
            "structure": "Suggest a structure for this content that maximizes reader engagement. Think about a strong opening, logical flow of ideas with smooth transitions, and a satisfying conclusion or call to action."
        },
        "Researcher": {
            "abstract": "Generate a formal abstract (200-300 words) for this research paper. Include Background, Methods, Key Results, and Main Conclusions. Use precise scientific language.",
            "section_titles": "Suggest standard scientific section titles for a research paper (e.g., Abstract, Introduction, Materials and Methods, Results, Discussion, Conclusion, Acknowledgements, References).",
            "style": "Formal Scientific",
            "structure": "Propose a logical structure for this research paper, ensuring all standard scientific sections are present and ordered correctly for clear communication of research."
        },
        
        "Business Professional": {
             "abstract": "Generate a concise executive summary (150-200 words) for this business document. Focus on key objectives, findings, implications, and recommendations. Language should be professional and action-oriented.",
            "section_titles": "Suggest clear, professional section titles for a business report (e.g., Executive Summary, Problem Statement, Proposed Solution, Financial Projections, Conclusion).",
            "style": "Professional Business Formal",
            "structure": "Outline a standard structure for a business proposal or report based on this content, emphasizing clarity, logical flow from problem to solution, and actionable insights."
        },
        "General/Others": default_prompts # Use defaults if persona is "Others" or not specifically defined
    }
    
    # Get the specific set of prompts for the role, or fallback to defaults
    selected_persona_prompts = persona_prompts_map.get(persona_selected, default_prompts)
    # Get the specific task prompt, or fallback to the default for that task
    return selected_persona_prompts.get(task_type, default_prompts.get(task_type, "Provide general guidance for this task."))


def can_use_template_role(user_tier_val, template_role_val):
    """Checks if the user's tier allows using templates with a specific role/persona_category."""
    # (This was already in your original script, ensure UserTier is accessible)
    # This function might be better placed in config.py if UserTier is there,
    # or just called directly using UserTier.get_tier_features()
    tier_features = UserTier.get_tier_features()
    allowed_categories = tier_features.get(user_tier_val, {}).get("allowed_personas_for_llm", []) # Changed from template_categories
    return template_role_val in allowed_categories