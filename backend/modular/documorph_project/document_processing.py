import os
import logging
import tempfile
import fitz # PyMuPDF
import docx # python-docx
import pandas as pd
from PIL import Image
import io
import cv2 # OpenCV for image processing
import numpy as np
import base64 # For image previews
import zipfile
import re
import pytesseract # For OCR
import chardet # For text encoding detection

from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader, UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document # For creating Langchain Documents

from config import UserTier, UNSTRUCTURED_AVAILABLE, logger # Import logger from config
# Import table and chart extraction functions; these will now point to the user-provided scripts.
# If these files are in the same directory, direct import works.
# Otherwise, adjust sys.path or use relative imports if part of a package.
try:
    from table_extraction import extract_tables as extract_tables_from_pdf_page_unstructured # Alias to distinguish
    from table_extraction import clean_column_names # Import the cleaning function
except ImportError:
    logger.warning("Could not import custom 'table_extraction.py'. PDF table extraction might be limited.")
    def extract_tables_from_pdf_page_unstructured(page, page_num): return [] # Placeholder
    def clean_column_names(df): return df # Placeholder

try:
    from chart_extraction import identify_and_extract_charts
except ImportError:
    logger.warning("Could not import custom 'chart_extraction.py'. Chart extraction might be limited.")
    def identify_and_extract_charts(page, page_num): return [] # Placeholder


# Unstructured.io specific imports (conditional)
if UNSTRUCTURED_AVAILABLE:
    try:
        from unstructured.partition.pdf import partition_pdf
        from unstructured.partition.docx import partition_docx as unstructured_partition_docx
        from unstructured.partition.image import partition_image
        from unstructured.partition.text import partition_text
        from unstructured.documents.elements import Title, NarrativeText, Table as UnstructuredTableCls, Image as UnstructuredImageCls
    except ImportError:
        UNSTRUCTURED_AVAILABLE = False # Ensure flag is correct if sub-imports fail
        logger.error("Failed to import specific unstructured components even though UNSTRUCTURED_AVAILABLE was True.")


def open_pdf_fitz(pdf_path):
    """Safely open a PDF with PyMuPDF."""
    try:
        return fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"Error opening PDF with fitz: {pdf_path} - {e}")
        raise # Re-raise the exception to be caught by the caller

def extract_and_zip_images_from_pdf(pdf_path, output_zip_name="extracted_pdf_images.zip"):
    """
    Extracts unique images from a PDF and saves them into a ZIP archive.
    Returns the path to the ZIP file or None if no images found or error.
    """
    image_details_for_zip = [] # List of (filename_in_zip, image_bytes)
    
    try:
        doc = open_pdf_fitz(pdf_path) # Use our robust PDF opener
        if not doc:
            logger.error(f"Could not open PDF for image extraction: {pdf_path}")
            return None

        saved_xrefs = set()
        img_count = 0

        for page_number in range(len(doc)):
            page = doc[page_number]
            images_on_page = page.get_images(full=True)
            
            for img_index, img_info in enumerate(images_on_page):
                xref = img_info[0]
                if xref in saved_xrefs:
                    continue 
                saved_xrefs.add(xref)
                
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image: # Should not happen if xref is valid from get_images
                        logger.warning(f"Could not extract base_image for xref {xref} on page {page_number+1}")
                        continue
                        
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"].lower() # Ensure lowercase extension
                    
                    if not image_bytes:
                        logger.warning(f"Empty image bytes for xref {xref} on page {page_number+1}")
                        continue

                    # Validate and ensure correct extension (simple validation)
                    try:
                        pil_img_test = Image.open(io.BytesIO(image_bytes))
                        actual_format = pil_img_test.format
                        if actual_format:
                            image_ext = actual_format.lower()
                    except Exception:
                        logger.warning(f"Could not validate image with PIL for xref {xref}, using original extension: {image_ext}")
                        # If PIL fails, image might be corrupt or an unusual format not directly usable by PIL without conversion

                    img_filename_in_zip = f"page{page_number+1}_img{img_index+1}.{image_ext}"
                    image_details_for_zip.append((img_filename_in_zip, image_bytes))
                    img_count += 1
                except Exception as e_inner:
                    logger.error(f"Error extracting image xref {xref} on page {page_number+1}: {e_inner}")
        
        doc.close()

        if not image_details_for_zip:
            logger.info(f"No unique images found in PDF: {pdf_path}")
            return None

        # Create a ZIP file in memory
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for filename_in_zip, img_bytes_data in image_details_for_zip:
                zf.writestr(filename_in_zip, img_bytes_data)
        
        zip_buffer.seek(0)
        
        # Save the zip file to a temporary location to provide a path for download
        # Or, Streamlit's download_button can take bytes directly.
        # For consistency, let's save it temporarily if a path is preferred by calling functions.
        temp_zip_dir = os.path.join(tempfile.gettempdir(), "documorph_image_exports")
        os.makedirs(temp_zip_dir, exist_ok=True)
        temp_zip_path = os.path.join(temp_zip_dir, output_zip_name)

        with open(temp_zip_path, "wb") as f_zip:
            f_zip.write(zip_buffer.getvalue())
            
        logger.info(f"Successfully extracted {img_count} images and zipped to {temp_zip_path}")
        return temp_zip_path

    except Exception as e:
        logger.error(f"Error extracting images from PDF {pdf_path}: {e}", exc_info=True)
        return None
    
# --- Unstructured.io Processing ---
def process_with_unstructured(file_path, file_type, ocr_languages):
    """
    Process document with unstructured.io with improved error handling.
    """
    if not UNSTRUCTURED_AVAILABLE:
        logger.warning("Unstructured module not available for parsing.")
        return [], [], [], "Unstructured module not available for parsing."
    
    try:
        logger.info(f"Processing document with unstructured.io: {file_path}")
        elements = []
        extracted_tables = []
        extracted_figures = []
        
        # Prepare OCR parameters
        extra_params = {
            "languages": ocr_languages, 
            "strategy": "hi_res" if len(ocr_languages) > 1 else "fast",
            "ocr_languages": ocr_languages
        }
        
        # Verify file exists and is accessible
        if not os.path.exists(file_path):
            return [], [], [], f"File not found: {file_path}"
            
        # Handle different file types
        if file_type == "pdf":
            try:
                from unstructured.partition.pdf import partition_pdf
                elements = partition_pdf(file_path, **extra_params)
            except ImportError:
                return [], [], [], "unstructured.partition.pdf not available"
        elif file_type == "docx":
            try:
                from unstructured.partition.docx import partition_docx
                elements = partition_docx(file_path)
            except ImportError:
                return [], [], [], "unstructured.partition.docx not available"
        elif file_type == "txt":
            try:
                # Handle text encoding issues
                file_type, encoding, confidence = detect_file_type(file_path)
                if file_type != 'txt':
                    return [], [], [], f"File is not a valid text file: {file_path}"
                    
                with open(file_path, "r", encoding=encoding) as f:
                    text_content = f.read()
                
                from unstructured.partition.text import partition_text
                elements = partition_text(text_content)
            except Exception as e:
                return [], [], [], f"Error reading text file: {e}"
        elif file_type in ["png", "jpg", "jpeg", "tiff"]:
            try:
                from unstructured.partition.image import partition_image
                elements = partition_image(file_path, ocr_languages=ocr_languages, strategy=extra_params.get("strategy", "fast"))
            except ImportError:
                return [], [], [], "unstructured.partition.image not available"
        else:
            # Try using the general UnstructuredFileLoader for unknown types
            try:
                loader = UnstructuredFileLoader(file_path)
                docs = loader.load()
                # Convert to elements format
                elements = [{"text": doc.page_content} for doc in docs]
                return docs, [], [], None
            except Exception as general_e:
                return [], [], [], f"Unsupported file type for unstructured: {file_type}. Error: {general_e}"
        
        # Process extracted elements
        texts_for_chunks = []
        for i, element in enumerate(elements):
            # Process tables
            if hasattr(element, 'to_dict') and element.to_dict().get('type') == 'Table':
                # Handle table elements
                table_text = element.text if hasattr(element, 'text') else str(element)
                
                # Extract table data if available
                table_data = None
                if hasattr(element, 'metadata') and hasattr(element.metadata, 'text_as_html'):
                    try:
                        import pandas as pd
                        dfs = pd.read_html(element.metadata.text_as_html)
                        if dfs:
                            table_data = dfs[0]
                    except:
                        pass
                
                extracted_tables.append({
                    "table_id": f"table_unstructured_{i}",
                    "dataframe": table_data,
                    "raw_text": table_text,
                    "page": getattr(element, 'metadata', {}).get('page_number', 1)
                })
                texts_for_chunks.append(f"[TABLE {i}]\n{table_text}\n[/TABLE]")
            
            # Process images
            elif hasattr(element, 'to_dict') and element.to_dict().get('type') == 'Image':
                fig_id = f"figure_unstructured_{i}"
                # Try to get image data
                image_bytes = None
                if hasattr(element, 'metadata') and hasattr(element.metadata, 'image_bytes'):
                    image_bytes = element.metadata.image_bytes
                
                extracted_figures.append({
                    "figure_id": fig_id,
                    "page": getattr(element, 'metadata', {}).get('page_number', 1),
                    "image_bytes": image_bytes,
                    "caption": getattr(element, 'text', f"Figure {i}")
                })
                texts_for_chunks.append(f"[FIGURE {fig_id}]\n{getattr(element, 'text', '')}\n[/FIGURE]")
            
            # Handle standard text content
            elif hasattr(element, 'text'):
                texts_for_chunks.append(element.text)
            elif hasattr(element, 'to_dict'):
                # Extract text from dictionary representation
                el_dict = element.to_dict()
                if 'text' in el_dict:
                    texts_for_chunks.append(el_dict['text'])
            else:
                # Last resort: convert element to string
                texts_for_chunks.append(str(element))
        
        # Create document chunks
        doc_chunks = []
        if texts_for_chunks:
            combined_text = "\n\n".join(texts_for_chunks)
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
            doc_chunks = text_splitter.create_documents([combined_text], [{"source": file_path}])
            
        logger.info(f"Unstructured extraction complete: {len(doc_chunks)} chunks, {len(extracted_tables)} tables, {len(extracted_figures)} figures")
        return doc_chunks, extracted_tables, extracted_figures, None
        
    except Exception as e:
        logger.error(f"Error processing with unstructured: {e}", exc_info=True)
        return [], [], [], f"Error processing with unstructured: {str(e)}"


# --- Standard Document Processing (Langchain Loaders) ---
def process_document_standard(file_path, file_type):
    """Processes DOCX, PDF, TXT using standard Langchain loaders and local table/figure extraction."""
    try:
        if file_type == "pdf":
            loader = PyMuPDFLoader(file_path)
            # Use our specific table/figure detection for PDFs
            tables = detect_tables_from_pdf(file_path) # This will use fitz and our custom table_extraction
            figures = detect_figures_from_pdf(file_path) # This will use fitz and our custom chart_extraction
        elif file_type == "docx":
            loader = Docx2txtLoader(file_path)
            tables = detect_tables_from_docx(file_path)
            figures = detect_figures_from_docx(file_path)
        elif file_type == "txt":
            loader = TextLoader(file_path, autodetect_encoding=True)
            tables, figures = [], []
        else:
            return [], "Unsupported file type", [], []

        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200) # Larger chunks
        chunks = text_splitter.split_documents(documents)
        
        logger.info(f"Standard processing for {file_path}: {len(chunks)} chunks, {len(tables)} tables, {len(figures)} figures.")
        return chunks, None, tables, figures # Error is None if successful
    except Exception as e:
        logger.error(f"Error in standard document processing for {file_path}: {e}", exc_info=True)
        return [], f"Error in standard processing: {str(e)}", [], []


# --- Combined Document Processing Entry Point ---
def process_document_master(file_path, file_type, user_tier, use_unstructured_if_available=True):
    """Master function to process document, with enhanced handling for unstructured formats and errors."""
    logger.info(f"Processing document: {file_path} of type {file_type}")
    
    # Check if file exists and is readable
    if not os.path.exists(file_path):
        error_msg = f"File not found: {file_path}"
        logger.error(error_msg)
        return [], [], [], error_msg
        
    try:
        with open(file_path, "rb") as test_f:
            test_f.read(1)  # Try to read a byte to test file access
    except Exception as e:
        error_msg = f"Error accessing file {file_path}: {str(e)}"
        logger.error(error_msg)
        return [], [], [], error_msg
    
    # Get proper file type with enhanced detection
    detected_type, encoding, confidence = detect_file_type(file_path)
    if detected_type == 'unknown':
        logger.warning(f"Could not detect file type for {file_path}, falling back to extension: {file_type}")
    else:
        if detected_type != file_type:
            logger.info(f"Detected file type {detected_type} differs from extension {file_type}, using detected type")
            file_type = detected_type
    
    # Determine OCR languages based on user tier
    if isinstance(user_tier, str):
        tier_features = UserTier.get_tier_features().get(user_tier, {})
    else:
        tier_features = user_tier  # Assume it's already a dict of features
        
    ocr_languages = tier_features.get("multi_language_ocr", ["eng"])
    
    # First try unstructured.io if available and enabled
    if use_unstructured_if_available and UNSTRUCTURED_AVAILABLE:
        logger.info(f"Attempting to process {file_path} with unstructured.io")
        doc_chunks, tables, figures, error = process_with_unstructured(file_path, file_type, ocr_languages)
        if error is None and doc_chunks:
            return doc_chunks, tables, figures, None
        else:
            logger.warning(f"Unstructured processing failed: {error}. Falling back to standard processing.")
    
    # Fall back to standard processing
    logger.info(f"Processing {file_path} with standard methods")
    tables, figures = [], []
    
    # Check if this is an image that may contain handwritten content
    if file_type.lower() in ["png", "jpg", "jpeg", "tiff"]:
        logger.info(f"Processing possible handwritten content in image: {file_path}")
        
        # Convert language list to string format for pytesseract
        ocr_lang_str = "+".join(ocr_languages)
        
        # Process the image with handwritten-optimized OCR
        extracted_text = process_handwritten_image(file_path, language=ocr_lang_str)
        
        if extracted_text and len(extracted_text.strip()) > 0:
            # Create a Langchain Document from the extracted text
            doc_chunks = [Document(page_content=extracted_text, metadata={"source": file_path, "chunk_idx": 0})]
            logger.info(f"Successfully extracted handwritten text from image: {file_path}")
            return doc_chunks, tables, figures, None
        else:
            logger.warning(f"Failed to extract text from image: {file_path}")
            # Create a single document with an error message
            doc_chunks = [Document(page_content=f"[Image file contains no recognizable text: {os.path.basename(file_path)}]", 
                                  metadata={"source": file_path})]
            return doc_chunks, tables, figures, None
    
    # Continue with standard processing for non-image files
    loader = None
    try:
        if file_type == "pdf":
            loader = PyMuPDFLoader(file_path)
            tables = detect_tables_from_pdf(file_path)
            figures = detect_figures_from_pdf(file_path)
        elif file_type == "docx":
            loader = Docx2txtLoader(file_path)
            tables = detect_tables_from_docx(file_path)
            figures = detect_figures_from_docx(file_path)
        elif file_type == "txt":
            # Handle text files with encoding detection
            if encoding:
                loader = TextLoader(file_path, encoding=encoding)
            else:
                # Try with utf-8 and fallback to latin-1 if needed
                try:
                    loader = TextLoader(file_path, encoding='utf-8')
                except UnicodeDecodeError:
                    loader = TextLoader(file_path, encoding='latin-1')
        else:
            # Try to use a generic loader as last resort
            try:
                loader = UnstructuredFileLoader(file_path)
            except:
                return [], [], [], f"Unsupported file type for standard processing: {file_type}"
    except Exception as loader_error:
        return [], [], [], f"Error creating document loader: {str(loader_error)}"

    if loader is None:
         return [], [], [], "Loader could not be initialized."

    try:
        documents = loader.load()
        if not documents:
            return [], [], [], f"No content could be extracted from {file_path}"
            
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        doc_chunks = text_splitter.split_documents(documents)
        
        logger.info(f"Standard processing of {file_path} resulted in {len(doc_chunks)} chunks, {len(tables)} tables, {len(figures)} figures.")
        return doc_chunks, tables, figures, None
    except Exception as e:
        logger.error(f"Error in standard document processing for {file_path}: {e}", exc_info=True)
        return [], [], [], f"Error in standard document processing: {str(e)}"


# --- Table and Figure/Chart Detection specific to PDF/DOCX (Non-Unstructured full parse) ---
def detect_tables_from_pdf(pdf_path):
    """Detect tables from PDF using fitz and fallback to custom page-based unstructured if needed."""
    tables_data = []
    try:
        doc = open_pdf_fitz(pdf_path)
        for page_num, page in enumerate(doc):
            # 1. PyMuPDF's find_tables()
            found_tables_fitz = page.find_tables()
            if found_tables_fitz.tables:
                for i, fitz_table in enumerate(found_tables_fitz.tables):
                    raw_data = fitz_table.extract()
                    df = None
                    if raw_data:
                        try: # Pandas DataFrame creation
                            df = pd.DataFrame(raw_data[1:], columns=raw_data[0]) if len(raw_data) > 1 else pd.DataFrame(raw_data)
                            df = clean_column_names(df.copy())
                        except Exception as e_df: logger.warning(f"Pandas DF creation failed for fitz table: {e_df}")
                    
                    tables_data.append({
                        "table_id": f"fitz_{page_num+1}_{i}", "dataframe": df, "raw_data": raw_data,
                        "page": page_num + 1, "rect": list(fitz_table.bbox), "extraction_method": "fitz"
                    })
            
            # 2. Fallback to user-provided page-based unstructured table extraction (from table_extraction.py)
            # This is if fitz doesn't find tables or if you want to augment
            try:
                unstructured_page_tables = extract_tables_from_pdf_page_unstructured(page, page_num) # This is the imported one
                for i, table_chunk_text in enumerate(unstructured_page_tables):
                    if "Table:" in table_chunk_text and not table_chunk_text.startswith("[No tables") and not table_chunk_text.startswith("[Error"):
                        table_text_content = table_chunk_text.split("Table:", 1)[1].strip()
                        # Further parse table_text_content into raw_data and df
                        raw_data = [re.split(r'\s*\|\s*', line.strip('|')) for line in table_text_content.split('\n') if line.strip()]
                        df = None
                        if raw_data:
                            try:
                                df = pd.DataFrame(raw_data[1:], columns=raw_data[0]) if len(raw_data) > 1 else pd.DataFrame(raw_data)
                                df = clean_column_names(df.copy())
                            except Exception as e_df_us: logger.warning(f"Pandas DF creation failed for unstructured page table: {e_df_us}")
                        tables_data.append({
                            "table_id": f"uspage_{page_num+1}_{i}", "dataframe": df, "raw_data": raw_data,
                            "page": page_num + 1, "extraction_method": "unstructured_page_fallback"
                        })
            except Exception as e_us_page:
                logger.error(f"Error in page-based unstructured table extraction: {e_us_page}")

    except Exception as e:
        logger.error(f"Error detecting tables from PDF {pdf_path}: {e}")
    return tables_data


def detect_tables_from_docx(docx_path):
    tables_data = []
    try:
        doc = docx.Document(docx_path)
        for i, table_obj in enumerate(doc.tables):
            raw_data = []
            for row in table_obj.rows:
                raw_data.append([cell.text for cell in row.cells])
            df = None
            if raw_data:
                try:
                    df = pd.DataFrame(raw_data[1:], columns=raw_data[0]) if len(raw_data) > 1 else pd.DataFrame(raw_data)
                    df = clean_column_names(df.copy())
                except Exception as e_df: logger.warning(f"Pandas DF creation failed for docx table: {e_df}")
            tables_data.append({
                "table_id": f"docx_{i}", "dataframe": df, "raw_data": raw_data,
                "page": None, # Page numbers are harder to get reliably from python-docx for tables
                "extraction_method": "python-docx"
            })
    except Exception as e:
        logger.error(f"Error detecting tables from DOCX {docx_path}: {e}")
    return tables_data


def detect_figures_from_pdf(pdf_path):
    """Detect figures/images and attempt to identify charts from PDF."""
    figures_data = []
    try:
        doc = open_pdf_fitz(pdf_path)
        for page_num, page in enumerate(doc):
            # 1. Extract standard images using PyMuPDF
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    if not image_bytes: continue

                    pil_img = Image.open(io.BytesIO(image_bytes))
                    width, height = pil_img.size
                    img_b64 = base64.b64encode(image_bytes).decode('utf-8')
                    preview_src = f"data:image/{base_image['ext'].lower()};base64,{img_b64}"
                    
                    figures_data.append({
                        "figure_id": f"img_{page_num+1}_{img_index}", "page": page_num + 1,
                        "image_bytes": image_bytes, "rect": list(page.get_image_bbox(img_info)),
                        "width": width, "height": height, "image_type": base_image["ext"].lower(),
                        "preview_src": preview_src, "type": "image", "extraction_method": "fitz"
                    })
                except Exception as e_img:
                    logger.warning(f"Could not process image xref {xref} on page {page_num+1}: {e_img}")

            # 2. Attempt chart text extraction using user-provided chart_extraction.py
            try:
                chart_text_chunks = identify_and_extract_charts(page, page_num) # Imported function
                for i, chart_text_info in enumerate(chart_text_chunks):
                    if "Chart:" in chart_text_info and not "[No charts detected]" in chart_text_info and not "[Error:" in chart_text_info:
                        chart_content = chart_text_info.split("Chart:", 1)[1].strip()
                        # For charts identified this way, we might not have separate image bytes,
                        # but we can associate the text with the page.
                        # Optionally, render the page as an image for chart preview if this text is significant.
                        # This adds complexity and I/O. For now, just log the text.
                        figures_data.append({
                            "figure_id": f"charttext_{page_num+1}_{i}", "page": page_num + 1,
                            "text_content": chart_content, "type": "chart_text_unstructured",
                             "image_bytes": None, # Could attempt full page render here if needed
                            "preview_src": None
                        })
            except Exception as e_chart:
                logger.error(f"Error in page-based chart text extraction: {e_chart}")

    except Exception as e:
        logger.error(f"Error detecting figures/charts from PDF {pdf_path}: {e}")
    return figures_data


def detect_figures_from_docx(docx_path):
    """Extract images from DOCX. python-docx has limited direct image byte extraction,
       often relying on relationships. This is a simplified placeholder."""
    figures_data = []
    try:
        doc = docx.Document(docx_path)
        img_idx = 0
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image_type = image_part.content_type.split('/')[-1] # e.g., png, jpeg
                    
                    pil_img = Image.open(io.BytesIO(image_bytes))
                    width, height = pil_img.size
                    img_b64 = base64.b64encode(image_bytes).decode('utf-8')
                    preview_src = f"data:image/{image_type};base64,{img_b64}"

                    figures_data.append({
                        "figure_id": f"docx_img_{img_idx}", "page": None, # Page difficult to ascertain
                        "image_bytes": image_bytes, "image_type": image_type,
                        "width": width, "height": height,
                        "preview_src": preview_src, "type": "image", "extraction_method": "python-docx-rels"
                    })
                    img_idx += 1
                except Exception as e_img_docx:
                    logger.warning(f"Could not extract image rel {rel_id} from DOCX: {e_img_docx}")
    except Exception as e:
        logger.error(f"Error detecting figures from DOCX {docx_path}: {e}")
    return figures_data

# Configure pytesseract with Tesseract path if needed
# Uncomment and modify the line below if Tesseract is not in PATH
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Function to detect file type more accurately, beyond just file extension
def detect_file_type(file_path):
    """
    Detect the actual file type beyond just file extension.
    Returns both detected file type and confidence level.
    """
    try:
        # First check by file extension
        ext = os.path.splitext(file_path)[1].lower().strip('.')
        
        # Try to determine more accurate file type
        if ext == 'txt':
            # Check encoding for text files
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)  # Read a chunk to detect encoding
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                confidence = result['confidence']
                
                # Try to read the file with detected encoding
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        f.read(100)  # Try to read a bit to verify encoding
                    return 'txt', encoding, confidence
                except UnicodeDecodeError:
                    # If we can't decode as text, it might be binary
                    return 'binary', None, 0.0
                    
        elif ext in ['docx', 'doc']:
            # Verify it's a valid document
            try:
                docx.Document(file_path)
                return 'docx', None, 1.0
            except Exception:
                return 'unknown', None, 0.0
                
        elif ext == 'pdf':
            # Verify it's a valid PDF
            try:
                with fitz.open(file_path) as doc:
                    page_count = len(doc)
                return 'pdf', None, 1.0
            except Exception:
                return 'unknown', None, 0.0
                
        elif ext in ['png', 'jpg', 'jpeg', 'tiff']:
            # Verify it's a valid image
            try:
                img = Image.open(file_path)
                img.verify()
                return ext, None, 1.0
            except Exception:
                return 'unknown', None, 0.0
        
        # Default to extension-based detection
        return ext, None, 0.5
        
    except Exception as e:
        logger.error(f"Error detecting file type: {e}")
        return 'unknown', None, 0.0

# Improved handwritten image processing function
def process_handwritten_image(image_path, language='eng'):
    """
    Process an image with handwritten text using specialized OCR settings.
    Enhanced with multi-language support and various preprocessing techniques.
    
    Args:
        image_path: Path to the image file
        language: OCR language code(s)
        
    Returns:
        Extracted text from the image
    """
    try:
        # Load image
        img = cv2.imread(image_path)
        if img is None:
            logger.error(f"Could not read image: {image_path}")
            return "Image could not be read."
        
        # Store original for comparison
        original = img.copy()
        
        # Get image properties
        height, width = img.shape[:2]
        
        # Create a list to store results from different preprocessing methods
        results = []
        
        # Try multiple preprocessing techniques and select the best result
        
        # 1. Basic grayscale conversion
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # 2. Adaptive thresholding - good for varied lighting
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                     cv2.THRESH_BINARY_INV, 11, 2)
        
        # 3. Otsu's thresholding - good for bimodal images
        _, otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        
        # 4. Noise removal
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        
        # 5. Sharpening using unsharp mask
        gaussian = cv2.GaussianBlur(gray, (0, 0), 3)
        sharpened = cv2.addWeighted(gray, 1.5, gaussian, -0.5, 0)
        
        # 6. Edge enhancement
        edges = cv2.Canny(gray, 100, 200)
        dilated_edges = cv2.dilate(edges, None, iterations=1)
        
        # Process all preprocessed versions
        preprocessing_methods = [
            ("original", original),
            ("grayscale", gray),
            ("adaptive_threshold", thresh),
            ("otsu", otsu),
            ("denoised", denoised),
            ("sharpened", sharpened)
        ]
        
        # Configure OCR settings for handwritten text
        # Try different psm modes
        psm_modes = [6, 11, 13]  # 6: single block, 11: single line, 13: raw line
        
        # If we have a non-English language, force LSTM
        oem_mode = 3  # LSTM only mode (most accurate for handwritten, especially non-English)
        
        best_text = ""
        best_confidence = 0
        
        for name, img_processed in preprocessing_methods:
            for psm in psm_modes:
                # Skip invalid combinations (some modes don't work with all image types)
                if name in ["edges", "dilated_edges"] and psm != 6:
                    continue
                    
                # Create custom config with current parameters
                custom_config = f'--psm {psm} --oem {oem_mode} -l {language}'
                
                # Add additional parameters for better handwritten text recognition
                if "+" in language or language != "eng":  # Multi-language or non-English
                    custom_config += ' --tessdata-dir "C:\\Program Files\\Tesseract-OCR\\tessdata"'
                
                # Add preserve_interword_spaces for better spacing
                custom_config += ' -c preserve_interword_spaces=1'
                
                try:
                    # Run OCR
                    text = pytesseract.image_to_string(img_processed, config=custom_config)
                    
                    # Skip empty results
                    if not text.strip():
                        continue
                        
                    # Estimate text quality (word count, avg word length)
                    words = [w for w in text.split() if len(w) > 1]
                    word_count = len(words)
                    avg_word_length = sum(len(w) for w in words) / max(1, word_count)
                    
                    # Simple confidence metric (can be improved)
                    confidence = word_count * min(avg_word_length, 10) / 20.0
                    
                    logger.debug(f"Method: {name}, PSM: {psm}, Words: {word_count}, Confidence: {confidence:.2f}")
                    
                    # Save result if it's better than previous
                    if confidence > best_confidence:
                        best_text = text
                        best_confidence = confidence
                        
                except Exception as e:
                    logger.warning(f"OCR failed for method {name}, psm {psm}: {e}")
        
        # If no good result, try one more approach with deskewing
        if best_confidence < 0.3:
            try:
                # Attempt to deskew the image
                coords = np.column_stack(np.where(gray > 0))
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                    
                # Rotate the image to deskew
                (h, w) = gray.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                
                # Try OCR on deskewed image
                custom_config = f'--psm 6 --oem 3 -l {language} -c preserve_interword_spaces=1'
                text = pytesseract.image_to_string(rotated, config=custom_config)
                
                # Check if result is better
                words = [w for w in text.split() if len(w) > 1]
                word_count = len(words)
                if word_count > 0:
                    best_text = text
            except Exception as e:
                logger.warning(f"Deskewing attempt failed: {e}")
        
        # Return best result, or empty string if nothing found
        return best_text
        
    except Exception as e:
        logger.error(f"Error processing handwritten image: {e}", exc_info=True)
        return f"Error processing image: {str(e)}"