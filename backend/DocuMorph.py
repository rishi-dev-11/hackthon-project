import streamlit as st
import io
import base64
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
import pandas as pd
import json
import os
import re
import time
import tempfile
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes

# Function to add page numbers
def add_page_number_field(paragraph):
    run = paragraph.add_run()
    field_start = OxmlElement('w:fldChar')
    field_start.set(ns.qn('w:fldCharType'), 'begin')
    
    instruction_text = OxmlElement('w:instrText')
    instruction_text.set(ns.qn('xml:space'), 'preserve')
    instruction_text.text = 'PAGE'
    
    field_end = OxmlElement('w:fldChar')
    field_end.set(ns.qn('w:fldCharType'), 'end')
    
    run._r.append(field_start)
    run._r.append(instruction_text)
    run._r.append(field_end)

# New imports for AI enhancement
import requests
from sentence_transformers import SentenceTransformer
import numpy as np
from typing import List, Dict, Any, Tuple
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('DocuMorph')

# Set page configuration
st.set_page_config(
    page_title="DocuMorph AI - Intelligent Document Transformation",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize AI models
@st.cache_resource
def load_embedding_model():
    try:
        logger.info("Loading all-MiniLM-L6 embedding model...")
        model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("Embedding model loaded successfully")
        return model
    except Exception as e:
        logger.error(f"Error loading embedding model: {e}")
        st.error(f"Failed to load embedding model: {e}")
        return None

# Groq API client for LLM processing
class GroqClient:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.environ.get('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')
        self.base_url = "https://api.groq.com/openai/v1"
        
    def generate_text(self, prompt, model="meta-llama/llama-4-scout-17b-16e-instruct", temperature=0.7, max_tokens=1024):
        """Generate text using Groq's LLM API"""
        if not self.api_key:
            logger.error("No Groq API key provided")
            return "Error: No API key provided. Please set your Groq API key in the advanced settings."
            
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "max_tokens": max_tokens
            }
            
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=payload
            )
            
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"]
            else:
                logger.error(f"Groq API error: {response.status_code} - {response.text}")
                return f"Error: {response.status_code} - {response.text}"
                
        except Exception as e:
            logger.error(f"Exception when calling Groq API: {e}")
            return f"Error: {str(e)}"
            
    def analyze_document_structure(self, text):
        """Analyze document structure using Groq LLM"""
        prompt = f"""
        Analyze the following document text and identify its structure. 
        Return a JSON object with the following:
        1. "title": Extract the document title
        2. "sections": List of identified sections with their headings
        3. "style": Guess the document type (academic, business, technical, etc.)
        4. "structure_quality": Rate how well-structured the document is (1-10)
        5. "suggestions": Provide 2-3 formatting suggestions for improvement
        
        Document text:
        {text[:5000]}  # Limit to first 5000 chars for API efficiency
        
        Return ONLY valid JSON with no additional text.
        """
        
        result = self.generate_text(prompt)
        try:
            return json.loads(result)
        except json.JSONDecodeError:
            logger.error(f"Failed to parse JSON from Groq response: {result[:100]}...")
            return {
                "title": "Unknown",
                "sections": [],
                "style": "Unknown",
                "structure_quality": 0,
                "suggestions": ["Could not analyze document structure"]
            }

# Document section detection using embeddings
class SectionDetector:
    def __init__(self, embedding_model):
        self.embedding_model = embedding_model
        self.section_types = [
            "Title", "Abstract", "Introduction", "Methods", "Methodology", 
            "Results", "Discussion", "Conclusion", "References", "Appendix",
            "Executive Summary", "Problem Statement", "Background", "Analysis",
            "Recommendations", "Implementation Plan", "Budget", "Timeline"
        ]
        self.section_embeddings = self._compute_section_embeddings()
        
    def _compute_section_embeddings(self):
        """Pre-compute embeddings for section types"""
        if self.embedding_model is None:
            return {}
            
        return {
            section_type: self.embedding_model.encode(section_type)
            for section_type in self.section_types
        }
        
    def detect_sections(self, text: str) -> List[Dict[str, Any]]:
        """Detect document sections using semantic embeddings"""
        if self.embedding_model is None:
            return []
            
        # Split text into potential paragraphs/sections
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        sections = []
        current_section = None
        
        for i, paragraph in enumerate(paragraphs):
            # Only analyze short paragraphs that might be headings
            if len(paragraph) < 100:
                # Get embedding for the paragraph
                paragraph_embedding = self.embedding_model.encode(paragraph)
                
                # Find the most similar section type
                similarities = {
                    section_type: np.dot(paragraph_embedding, self.section_embeddings[section_type])
                    for section_type in self.section_types
                }
                
                best_match = max(similarities.items(), key=lambda x: x[1])
                
                # If similarity is high enough, consider it a section heading
                if best_match[1] > 0.5:  # Threshold for similarity
                    current_section = {
                        "name": best_match[0], 
                        "heading": paragraph,
                        "content": [],
                        "confidence": float(best_match[1])
                    }
                    sections.append(current_section)
            else:
                # This is content for the current section
                if current_section:
                    current_section["content"].append(paragraph)
        
        return sections

# Initialize session state
if 'templates' not in st.session_state:
    st.session_state.templates = {
        "Academic Paper": {
            "title_font": "Arial",
            "title_size": 16,
            "title_bold": True,
            "heading_font": "Arial",
            "heading_size": 14,
            "heading_bold": True,
            "subheading_font": "Arial",
            "subheading_size": 12,
            "subheading_bold": True,
            "body_font": "Times New Roman",
            "body_size": 12,
            "body_bold": False,
            "line_spacing": 1.5,
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "headers": True,
            "footers": True,
            "section_numbering": True,
            "caption_style": "below",
            "has_toc": True,
            "include_page_numbers": True
        },
        "Business Report": {
            "title_font": "Calibri",
            "title_size": 18,
            "title_bold": True,
            "heading_font": "Calibri",
            "heading_size": 16,
            "heading_bold": True,
            "subheading_font": "Calibri",
            "subheading_size": 14,
            "subheading_bold": True,
            "body_font": "Calibri",
            "body_size": 11,
            "body_bold": False,
            "line_spacing": 1.15,
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "headers": True,
            "footers": True,
            "section_numbering": False,
            "caption_style": "above",
            "has_toc": True,
            "include_page_numbers": True
        }
    }

for state_var in ['doc_content', 'logo_file', 'current_template', 'processed_doc', 'ai_analysis', 'embedding_model', 'groq_client', 'groq_api_key']:
    if state_var not in st.session_state:
        st.session_state[state_var] = None

# Custom CSS to improve the look and feel
st.markdown("""
<style>
    .main-header {
        font-size: 38px;
        font-weight: bold;
        color: #1E88E5;
        margin-bottom: 20px;
        text-align: center;
    }
    .sub-header {
        font-size: 24px;
        font-weight: bold;
        color: #0D47A1;
        margin-top: 30px;
        margin-bottom: 15px;
    }
    .section {
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .info-box {
        background-color: #E3F2FD;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 20px;
    }
    .ai-analysis-box {
        background-color: #F3E5F5;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 20px;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #F8F9FA;
        border-radius: 4px 4px 0px 0px;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #E3F2FD;
        border-bottom: 4px solid #1E88E5;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown('<p class="main-header">DocuMorph AI</p>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; font-size: 20px;">Intelligent Document Transformation Engine</p>', unsafe_allow_html=True)

# Create sidebar
st.sidebar.image("https://via.placeholder.com/150x80?text=DocuMorph+AI", use_column_width=True)
st.sidebar.markdown("### Document Transformation")

# Load the embedding model
if st.session_state.embedding_model is None:
    st.session_state.embedding_model = load_embedding_model()

# Initialize the Groq client
if st.session_state.groq_client is None:
    st.session_state.groq_client = GroqClient()

# Add API key configuration to sidebar
with st.sidebar.expander("AI Configuration", expanded=False):
    api_key = st.text_input("Groq API Key", type="password", 
                           help="Enter your Groq API key for enhanced AI analysis")
    if api_key:
        st.session_state.groq_api_key = api_key
        st.session_state.groq_client = GroqClient(api_key=api_key)
        st.success("API key configured!")
    
    # Model selection
    llm_model = st.selectbox(
        "LLM Model", 
        ["meta-llama/llama-4-scout-17b-16e-instruct", "mixtral-8x7b-32768", "llama2-70b-4096", "gemma-7b-it"],
        index=0,
        help="Select the LLM model to use for document analysis"
    )

# Main app tabs
tabs = st.tabs(["Upload & Format", "Template Management", "Advanced Features", "AI Analysis", "About"])

# Tab 1: Upload & Format
with tabs[0]:
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="sub-header">Upload Document</p>', unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader("Upload your unformatted document", 
                                        type=["docx", "txt", "pdf", "png", "jpg", "jpeg"],
                                        help="Upload a document in DOCX, TXT, PDF, or image format.")
        
        if uploaded_file is not None:
            file_details = {"Filename": uploaded_file.name, "FileType": uploaded_file.type, "FileSize": f"{uploaded_file.size / 1024:.2f} KB"}
            st.write("File Details:", file_details)
            
            # Process the uploaded file
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # DOCX file
                doc = Document(uploaded_file)
                st.session_state.doc_content = {
                    "type": "docx",
                    "content": doc,
                    "text": "\n".join([para.text for para in doc.paragraphs])
                }
                st.info(f"DOCX document with {len(doc.paragraphs)} paragraphs loaded successfully.")
                
            elif uploaded_file.type == "text/plain":
                # TXT file
                text_content = uploaded_file.getvalue().decode("utf-8")
                st.session_state.doc_content = {
                    "type": "txt",
                    "content": text_content,
                    "text": text_content
                }
                st.info(f"Text document with {len(text_content.splitlines())} lines loaded successfully.")
                
            elif uploaded_file.type in ["application/pdf", "image/png", "image/jpeg", "image/jpg"]:
                # PDF or Image file (OCR will be used)
                st.info("Processing document with OCR. This may take a moment...")
                
                if uploaded_file.type == "application/pdf":
                    try:
                        # Convert PDF to images
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                            tmp_file.write(uploaded_file.getvalue())
                            tmp_file_path = tmp_file.name
                        
                        images = convert_from_bytes(uploaded_file.getvalue())
                        extracted_text = ""
                        
                        for i, image in enumerate(images):
                            image_text = pytesseract.image_to_string(image)
                            extracted_text += f"[Page {i+1}]\n{image_text}\n\n"
                        
                        st.session_state.doc_content = {
                            "type": "ocr",
                            "content": extracted_text,
                            "text": extracted_text
                        }
                        st.success(f"PDF with {len(images)} pages processed with OCR.")
                        
                    except Exception as e:
                        st.error(f"Error processing PDF: {e}")
                        
                else:  # Image file
                    try:
                        image = Image.open(uploaded_file)
                        extracted_text = pytesseract.image_to_string(image)
                        
                        st.session_state.doc_content = {
                            "type": "ocr",
                            "content": extracted_text,
                            "text": extracted_text
                        }
                        st.success("Image processed with OCR.")
                        
                    except Exception as e:
                        st.error(f"Error processing image: {e}")
            
            # Once document is loaded, run AI analysis
            if st.session_state.doc_content and st.session_state.groq_api_key:
                with st.spinner("Analyzing document with AI..."):
                    text_content = st.session_state.doc_content["text"]
                    st.session_state.ai_analysis = st.session_state.groq_client.analyze_document_structure(text_content)
                    st.success("AI analysis complete!")
        
        # Logo upload (optional)
        st.markdown("#### Optional: Upload Logo")
        logo_file = st.file_uploader("Upload your logo (optional)", type=["png", "jpg", "jpeg"], 
                                    help="This logo will be added to the formatted document.")
        
        if logo_file is not None:
            st.session_state.logo_file = logo_file
            st.image(logo_file, width=100)
    
    with col2:
        st.markdown('<p class="sub-header">Select Template</p>', unsafe_allow_html=True)
        
        if st.session_state.templates:
            template_names = list(st.session_state.templates.keys())
            
            # AI template suggestion
            if st.session_state.ai_analysis and 'style' in st.session_state.ai_analysis:
                doc_style = st.session_state.ai_analysis['style']
                if doc_style.lower() == 'academic':
                    suggested_template = "Academic Paper"
                elif doc_style.lower() in ['business', 'professional']:
                    suggested_template = "Business Report"
                else:
                    suggested_template = template_names[0]
                
                st.info(f"💡 AI suggests using the '{suggested_template}' template based on document style.")
                selected_template = st.selectbox("Choose a template", template_names, 
                                               index=template_names.index(suggested_template) if suggested_template in template_names else 0)
            else:
                selected_template = st.selectbox("Choose a template", template_names)
            
            if selected_template:
                st.session_state.current_template = st.session_state.templates[selected_template]
                
                with st.expander("View Template Details"):
                    template_df = pd.DataFrame(
                        {
                            "Property": list(st.session_state.current_template.keys()),
                            "Value": [str(val) for val in st.session_state.current_template.values()]
                        }
                    )
                    st.table(template_df)
        else:
            st.warning("No templates available. Please create a template in the Template Management tab.")
        
        st.markdown('<p class="sub-header">Preview & Transform</p>', unsafe_allow_html=True)
        
        if st.session_state.doc_content and st.session_state.current_template:
            if st.button("Transform Document", type="primary"):
                with st.spinner("Transforming document..."):
                    # Create a new document with the selected template formatting
                    doc = Document()
                    
                    # Apply page setup from template
                    section = doc.sections[0]
                    margins = st.session_state.current_template["margins"]
                    section.top_margin = Inches(margins["top"])
                    section.bottom_margin = Inches(margins["bottom"])
                    section.left_margin = Inches(margins["left"])
                    section.right_margin = Inches(margins["right"])
                    
                    # Setup styles
                    # Title style
                    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
                    title_font = title_style.font
                    title_font.name = st.session_state.current_template["title_font"]
                    title_font.size = Pt(st.session_state.current_template["title_size"])
                    title_font.bold = st.session_state.current_template["title_bold"]
                    
                    # Heading style
                    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
                    heading_font = heading_style.font
                    heading_font.name = st.session_state.current_template["heading_font"]
                    heading_font.size = Pt(st.session_state.current_template["heading_size"])
                    heading_font.bold = st.session_state.current_template["heading_bold"]
                    
                    # Subheading style
                    subheading_style = doc.styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
                    subheading_font = subheading_style.font
                    subheading_font.name = st.session_state.current_template["subheading_font"]
                    subheading_font.size = Pt(st.session_state.current_template["subheading_size"])
                    subheading_font.bold = st.session_state.current_template["subheading_bold"]
                    
                    # Body style
                    body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
                    body_font = body_style.font
                    body_font.name = st.session_state.current_template["body_font"]
                    body_font.size = Pt(st.session_state.current_template["body_size"])
                    body_font.bold = st.session_state.current_template["body_bold"]

                    # Enhanced document processing with AI-detected sections
                    if st.session_state.ai_analysis and 'sections' in st.session_state.ai_analysis:
                        # Try to extract title from AI analysis
                        doc_title = st.session_state.ai_analysis.get('title', 'Untitled Document')
                        
                        # Add title page
                        title_para = doc.add_paragraph(doc_title, style='CustomTitle')
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()  # Add spacing
                        
                        # Create Table of Contents if specified in template
                        if st.session_state.current_template["has_toc"]:
                            toc_heading = doc.add_paragraph("Table of Contents", style='CustomHeading')
                            doc.add_paragraph("Right-click and select 'Update Field' to update TOC after editing.", style='CustomBody')
                            doc.add_paragraph()
                            doc.add_paragraph().add_run().add_break()
                            
                        # Process sections from AI analysis
                        section_counter = 1
                        ai_sections = st.session_state.ai_analysis['sections']
                        
                        for section in ai_sections:
                            if st.session_state.current_template["section_numbering"]:
                                heading_text = f"{section_counter}. {section}"
                                section_counter += 1
                            else:
                                heading_text = section
                                
                            doc.add_paragraph(heading_text, style='CustomHeading')
                            
                            # Add placeholder content for this section
                            doc.add_paragraph("This section content will be populated from your document.", style='CustomBody')
                        
                    else:
                        # Use traditional document processing approach
                        if st.session_state.doc_content["type"] == "docx":
                            # Extract basic structure from original document
                            original_doc = st.session_state.doc_content["content"]
                            
                            # Basic content classification
                            heading_pattern = re.compile(r'^#+\s+|^[A-Z\s]{5,}$|^\d+\.\s+[A-Z]')
                            section_counter = 1
                            
                            # Add title page if it's enabled
                            title_found = False
                            
                            for i, para in enumerate(original_doc.paragraphs):
                                text = para.text.strip()
                                
                                if not text:
                                    continue
                                    
                                # Try to identify the paragraph type
                                if i == 0 and not title_found:
                                    # First paragraph is usually the title
                                    title_para = doc.add_paragraph(text, style='CustomTitle')
                                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    title_found = True
                                    
                                elif heading_pattern.match(text) or (len(text) < 100 and text.isupper()):
                                    # Looks like a heading
                                    if st.session_state.current_template["section_numbering"]:
                                        heading_text = f"{section_counter}. {text.lstrip('#').strip()}"
                                        section_counter += 1
                                    else:
                                        heading_text = text.lstrip('#').strip()
                                        
                                    heading_para = doc.add_paragraph(heading_text, style='CustomHeading')
                                    
                                elif text.startswith('Figure') or text.startswith('Table'):
                                    # This is a caption
                                    caption_para = doc.add_paragraph(text)
                                    caption_para.style = 'Caption'
                                    
                                else:
                                    # Regular body text
                                    body_para = doc.add_paragraph(text, style='CustomBody')
                                    
                        elif st.session_state.doc_content["type"] in ["txt", "ocr"]:
                            # Process plain text document
                            text_content = st.session_state.doc_content["content"]
                            lines = text_content.splitlines()
                            
                            # Try to detect document structure
                            title_pattern = re.compile(r'^#+\s+|^[A-Z\s]{5,}$')
                            heading_pattern = re.compile(r'^#+\s+|^[A-Z][A-Za-z\s]{2,}:$|^\d+\.\s+[A-Z]')
                            
                            # Add title
                            title_found = False
                            section_counter = 1
                            
                            for i, line in enumerate(lines):
                                line = line.strip()
                                
                                if not line:
                                    continue
                                    
                                if i == 0 and not title_found:
                                    # First non-empty line is treated as title
                                    title_para = doc.add_paragraph(line, style='CustomTitle')
                                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    title_found = True
                                    doc.add_paragraph()  # Add spacing
                                    
                                elif heading_pattern.match(line) and len(line) < 100:
                                    # This looks like a heading
                                    if st.session_state.current_template["section_numbering"]:
                                        heading_text = f"{section_counter}. {line.lstrip('#').strip()}"
                                        section_counter += 1
                                    else:
                                        heading_text = line.lstrip('#').strip()
                                    
                                    # Add spacing before heading
                                    doc.add_paragraph()
                                    heading_para = doc.add_paragraph(heading_text, style='CustomHeading')
                                    
                                elif line.startswith('Figure') or line.startswith('Table'):
                                    # This is a caption
                                    caption_para = doc.add_paragraph(line)
                                    caption_para.style = 'Caption'
                                    
                                else:
                                    # Regular body text
                                    body_para = doc.add_paragraph(line, style='CustomBody')
                    
                    # Add logo if uploaded
                    if st.session_state.logo_file:
                        # Save the logo to a temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_logo:
                            tmp_logo.write(st.session_state.logo_file.getvalue())
                            logo_path = tmp_logo.name
                        
                        # Add logo to header
                        header = doc.sections[0].header
                        header_para = header.paragraphs[0]
                        header_run = header_para.add_run()
                        header_run.add_picture(logo_path, width=Inches(1.0))
                        
                        # Clean up temporary file
                        os.unlink(logo_path)
                    
                    # Add page numbers if enabled
                    if st.session_state.current_template["include_page_numbers"]:
                        footer = doc.sections[0].footer
                        footer_para = footer.paragraphs[0]
                        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_page_number_field(footer_para)
                    
                    # Save the document to a BytesIO object
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    
                    # Save to session state for download
                    st.session_state.processed_doc = doc_bytes
                    
                    st.success("Document transformation complete!")
                    
                    # Create download button
                    st.download_button(
                        label="Download Formatted Document",
                        data=doc_bytes,
                        file_name=f"DocuMorph_{uploaded_file.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.info("Please upload a document and select a template to transform your document.")

# Tab 2: Template Management
with tabs[1]:
    st.markdown('<p class="sub-header">Template Management</p>', unsafe_allow_html=True)
    
    template_action = st.radio("Choose action", ["Create New Template", "Edit Existing Template"])
    
    if template_action == "Create New Template":
        st.markdown("### Create New Template")
        
        template_name = st.text_input("Template Name", value="My Custom Template")
        
        with st.expander("General Settings", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                has_toc = st.checkbox("Include Table of Contents", value=True)
                include_page_numbers = st.checkbox("Include Page Numbers", value=True)
                section_numbering = st.checkbox("Enable Section Numbering", value=True)
            
            with col2:
                headers = st.checkbox("Include Headers", value=True)
                footers = st.checkbox("Include Footers", value=True)
                caption_style = st.selectbox("Caption Placement", ["below", "above"], index=0)
                
        with st.expander("Margins & Spacing"):
            col1, col2 = st.columns(2)
            
            with col1:
                margin_top = st.number_input("Top Margin (inches)", value=1.0, min_value=0.1, max_value=3.0, step=0.1)
                margin_left = st.number_input("Left Margin (inches)", value=1.0, min_value=0.1, max_value=3.0, step=0.1)
                
            with col2:
                margin_bottom = st.number_input("Bottom Margin (inches)", value=1.0, min_value=0.1, max_value=3.0, step=0.1)
                margin_right = st.number_input("Right Margin (inches)", value=1.0, min_value=0.1, max_value=3.0, step=0.1)
                
            line_spacing = st.slider("Line Spacing", min_value=1.0, max_value=2.0, value=1.15, step=0.05)
        
        with st.expander("Font Settings"):
            st.markdown("#### Title Font")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                title_font = st.selectbox("Title Font", ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], index=0)
            with col2:
                title_size = st.number_input("Title Size (pt)", value=16, min_value=8, max_value=72)
            with col3:
                title_bold = st.checkbox("Bold Title", value=True)
            
            st.markdown("#### Heading Font")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                heading_font = st.selectbox("Heading Font", ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], index=0)
            with col2:
                heading_size = st.number_input("Heading Size (pt)", value=14, min_value=8, max_value=72)
            with col3:
                heading_bold = st.checkbox("Bold Headings", value=True)
            
            st.markdown("#### Subheading Font")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                subheading_font = st.selectbox("Subheading Font", ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], index=0)
            with col2:
                subheading_size = st.number_input("Subheading Size (pt)", value=12, min_value=8, max_value=72)
            with col3:
                subheading_bold = st.checkbox("Bold Subheadings", value=True)
            
            st.markdown("#### Body Font")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                body_font = st.selectbox("Body Font", ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], index=1)
            with col2:
                body_size = st.number_input("Body Size (pt)", value=12, min_value=8, max_value=72)
            with col3:
                body_bold = st.checkbox("Bold Body", value=False)
        
        if st.button("Save Template", type="primary"):
            # Create a new template dictionary
            new_template = {
                "title_font": title_font,
                "title_size": title_size,
                "title_bold": title_bold,
                "heading_font": heading_font,
                "heading_size": heading_size,
                "heading_bold": heading_bold,
                "subheading_font": subheading_font,
                "subheading_size": subheading_size,
                "subheading_bold": subheading_bold,
                "body_font": body_font,
                "body_size": body_size,
                "body_bold": body_bold,
                "line_spacing": line_spacing,
                "margins": {"top": margin_top, "bottom": margin_bottom, "left": margin_left, "right": margin_right},
                "headers": headers,
                "footers": footers,
                "section_numbering": section_numbering,
                "caption_style": caption_style,
                "has_toc": has_toc,
                "include_page_numbers": include_page_numbers
            }
            
            # Add the template to session state
            st.session_state.templates[template_name] = new_template
            st.success(f"Template '{template_name}' created successfully!")
    
    else:  # Edit Existing Template
        if st.session_state.templates:
            template_to_edit = st.selectbox("Select template to edit", list(st.session_state.templates.keys()))
            selected_template = st.session_state.templates[template_to_edit]
            
            with st.expander("General Settings", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    has_toc = st.checkbox("Include Table of Contents", value=selected_template["has_toc"])
                    include_page_numbers = st.checkbox("Include Page Numbers", value=selected_template["include_page_numbers"])
                    section_numbering = st.checkbox("Enable Section Numbering", value=selected_template["section_numbering"])
                
                with col2:
                    headers = st.checkbox("Include Headers", value=selected_template["headers"])
                    footers = st.checkbox("Include Footers", value=selected_template["footers"])
                    caption_style = st.selectbox("Caption Placement", ["below", "above"], 
                                                 index=0 if selected_template["caption_style"] == "below" else 1)
            
            with st.expander("Margins & Spacing"):
                col1, col2 = st.columns(2)
                
                with col1:
                    margin_top = st.number_input("Top Margin (inches)", 
                                                  value=selected_template["margins"]["top"], 
                                                  min_value=0.1, max_value=3.0, step=0.1)
                    margin_left = st.number_input("Left Margin (inches)", 
                                                   value=selected_template["margins"]["left"], 
                                                   min_value=0.1, max_value=3.0, step=0.1)
                    
                with col2:
                    margin_bottom = st.number_input("Bottom Margin (inches)", 
                                                     value=selected_template["margins"]["bottom"], 
                                                     min_value=0.1, max_value=3.0, step=0.1)
                    margin_right = st.number_input("Right Margin (inches)", 
                                                    value=selected_template["margins"]["right"], 
                                                    min_value=0.1, max_value=3.0, step=0.1)
                    
                line_spacing = st.slider("Line Spacing", min_value=1.0, max_value=2.0, 
                                           value=selected_template["line_spacing"], step=0.05)
            
            with st.expander("Font Settings"):
                st.markdown("#### Title Font")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    title_font = st.selectbox("Title Font", 
                                                ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], 
                                                index=["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"].index(selected_template["title_font"]))
                with col2:
                    title_size = st.number_input("Title Size (pt)", 
                                                  value=selected_template["title_size"], 
                                                  min_value=8, max_value=72)
                with col3:
                    title_bold = st.checkbox("Bold Title", value=selected_template["title_bold"])
            
                st.markdown("#### Heading Font")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    heading_font = st.selectbox("Heading Font", 
                                                 ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], 
                                                 index=["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"].index(selected_template["heading_font"]))
                with col2:
                    heading_size = st.number_input("Heading Size (pt)", 
                                                    value=selected_template["heading_size"], 
                                                    min_value=8, max_value=72)
                with col3:
                    heading_bold = st.checkbox("Bold Headings", value=selected_template["heading_bold"])
            
                st.markdown("#### Subheading Font")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    subheading_font = st.selectbox("Subheading Font", 
                                                    ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], 
                                                    index=["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"].index(selected_template["subheading_font"]))
                with col2:
                    subheading_size = st.number_input("Subheading Size (pt)", 
                                                       value=selected_template["subheading_size"], 
                                                       min_value=8, max_value=72)
                with col3:
                    subheading_bold = st.checkbox("Bold Subheadings", value=selected_template["subheading_bold"])
            
                st.markdown("#### Body Font")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    body_font = st.selectbox("Body Font", 
                                               ["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"], 
                                               index=["Arial", "Times New Roman", "Calibri", "Georgia", "Verdana"].index(selected_template["body_font"]))
                with col2:
                    body_size = st.number_input("Body Size (pt)", 
                                                 value=selected_template["body_size"], 
                                                 min_value=8, max_value=72)
                with col3:
                    body_bold = st.checkbox("Bold Body", value=selected_template["body_bold"])
            
            if st.button("Update Template", type="primary"):
                # Update the template dictionary
                updated_template = {
                    "title_font": title_font,
                    "title_size": title_size,
                    "title_bold": title_bold,
                    "heading_font": heading_font,
                    "heading_size": heading_size,
                    "heading_bold": heading_bold,
                    "subheading_font": subheading_font,
                    "subheading_size": subheading_size,
                    "subheading_bold": subheading_bold,
                    "body_font": body_font,
                    "body_size": body_size,
                    "body_bold": body_bold,
                    "line_spacing": line_spacing,
                    "margins": {"top": margin_top, "bottom": margin_bottom, "left": margin_left, "right": margin_right},
                    "headers": headers,
                    "footers": footers,
                    "section_numbering": section_numbering,
                    "caption_style": caption_style,
                    "has_toc": has_toc,
                    "include_page_numbers": include_page_numbers
                }
                
                # Update the template in session state
                st.session_state.templates[template_to_edit] = updated_template
                st.success(f"Template '{template_to_edit}' updated successfully!")
                
                # Refresh the current template if it was the one edited
                if st.session_state.current_template and template_to_edit in st.session_state.templates:
                    st.session_state.current_template = st.session_state.templates[template_to_edit]
            
            if st.button("Delete Template"):
                # Delete the selected template
                del st.session_state.templates[template_to_edit]
                st.success(f"Template '{template_to_edit}' deleted successfully!")
                st.experimental_rerun()
        else:
            st.warning("No templates available to edit.")

# Tab 3: Advanced Features
with tabs[2]:
    st.markdown('<p class="sub-header">Advanced Features</p>', unsafe_allow_html=True)
    
    st.markdown("### Multi-Language Support")
    language = st.selectbox("Select document language", 
                           ["English", "Spanish", "French", "German", "Hindi"], 
                           help="Select the language for formatting rules and OCR processing.")
    
    # Language-specific font recommendations
    language_fonts = {
        "English": ["Arial", "Times New Roman", "Calibri"],
        "Spanish": ["Arial", "Times New Roman", "Georgia"],
        "French": ["Arial", "Garamond", "Verdana"],
        "German": ["Arial", "Helvetica", "Verdana"],
        "Hindi": ["Mangal", "Devanagari", "Arial Unicode MS"]
    }
    st.info(f"Recommended fonts for {language}: {', '.join(language_fonts.get(language, ['Arial']))}")
    
    st.markdown("### Smart Section Detection")
    if st.session_state.doc_content and st.session_state.embedding_model:
        section_detector = SectionDetector(st.session_state.embedding_model)
        
        if st.button("Detect Document Sections"):
            with st.spinner("Analyzing document structure with embeddings..."):
                text_content = st.session_state.doc_content["text"]
                sections = section_detector.detect_sections(text_content)
                
                # Display detected sections
                if sections:
                    st.markdown("#### Detected Sections")
                    for section in sections:
                        with st.expander(f"{section['name']} (Confidence: {section['confidence']:.2f})"):
                            st.write(f"**Heading**: {section['heading']}")
                            st.write("**Content Preview**:")
                            st.write("\n".join(section['content'][:3]) + "..." if len(section['content']) > 3 else "\n".join(section['content']))
                else:
                    st.warning("No recognizable sections detected.")
    else:
        st.info("Upload a document and ensure the embedding model is loaded to enable smart section detection.")
    
    st.markdown("### Version Control")
    if st.session_state.processed_doc:
        version_name = st.text_input("Save document version as", value="Version 1")
        if st.button("Save Version"):
            if 'version_history' not in st.session_state:
                st.session_state.version_history = {}
            
            # Save the document version
            version_key = f"{version_name}_{int(time.time())}"
            st.session_state.version_history[version_key] = st.session_state.processed_doc.getvalue()
            st.success(f"Version '{version_name}' saved!")
        
        # Display version history
        if 'version_history' in st.session_state and st.session_state.version_history:
            st.markdown("#### Version History")
            for version_key in st.session_state.version_history:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(version_key)
                with col2:
                    st.download_button(
                        label="Download",
                        data=st.session_state.version_history[version_key],
                        file_name=f"DocuMorph_{version_key}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    else:
        st.info("Transform a document to enable version control.")
    
    st.markdown("### Export Options")
    if st.session_state.processed_doc:
        export_format = st.selectbox("Export as", ["DOCX", "PDF", "Google Docs"])
        if export_format == "Google Docs":
            st.warning("Google Docs integration is not implemented in this version. Please download as DOCX and upload manually.")
        elif export_format == "PDF":
            st.warning("PDF export is not implemented in this version. Please convert the DOCX file using an external tool.")
        else:
            st.download_button(
                label="Download Formatted Document",
                data=st.session_state.processed_doc,
                file_name="DocuMorph_formatted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Transform a document to enable export options.")

# Tab 4: AI Analysis
with tabs[3]:
    st.markdown('<p class="sub-header">AI Document Analysis</p>', unsafe_allow_html=True)
    
    if st.session_state.ai_analysis and st.session_state.doc_content:
        st.markdown('<div class="ai-analysis-box">', unsafe_allow_html=True)
        st.markdown("### AI Analysis Results")
        
        # Display title
        st.markdown(f"**Document Title**: {st.session_state.ai_analysis.get('title', 'Unknown')}")
        
        # Display document style
        st.markdown(f"**Document Style**: {st.session_state.ai_analysis.get('style', 'Unknown')}")
        
        # Display structure quality
        st.markdown(f"**Structure Quality**: {st.session_state.ai_analysis.get('structure_quality', 0)}/10")
        
        # Display detected sections
        sections = st.session_state.ai_analysis.get('sections', [])
        if sections:
            st.markdown("**Detected Sections**:")
            for section in sections:
                st.markdown(f"- {section}")
        else:
            st.markdown("**Detected Sections**: None")
        
        # Display formatting suggestions
        suggestions = st.session_state.ai_analysis.get('suggestions', [])
        if suggestions:
            st.markdown("**Formatting Suggestions**:")
            for suggestion in suggestions:
                st.markdown(f"- {suggestion}")
        else:
            st.markdown("**Formatting Suggestions**: None")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Option to re-run analysis
        if st.button("Re-run AI Analysis"):
            with st.spinner("Re-analyzing document..."):
                text_content = st.session_state.doc_content["text"]
                st.session_state.ai_analysis = st.session_state.groq_client.analyze_document_structure(text_content)
                st.success("AI analysis updated!")
                st.experimental_rerun()
    else:
        st.info("Upload a document and ensure the Groq API key is configured to enable AI analysis.")

# Tab 5: About
with tabs[4]:
    st.markdown('<p class="sub-header">About DocuMorph AI</p>', unsafe_allow_html=True)
    st.markdown("""
    **DocuMorph AI** is an intelligent document transformation engine designed to streamline document formatting for professionals, researchers, and students. Powered by AI, including the Groq API for advanced analysis and Hugging Face's all-MiniLM-L6-v2 for semantic section detection, it automates the tedious process of styling documents to meet institutional or organizational standards, ensuring consistency and saving time.
    
    ### Key Features
    - **Automated Formatting**: Transform unformatted DOCX, TXT, PDF, or image-based documents into professionally styled outputs.
    - **Custom Templates**: Create, save, and apply reusable formatting templates.
    - **OCR Support**: Extract and format text from scanned documents or images using Tesseract.
    - **AI-Powered Analysis**: Leverage Groq's LLM for document structure analysis and formatting suggestions.
    - **Smart Section Detection**: Identify document sections like Abstract, Introduction, etc., using semantic embeddings.
    - **Version Control**: Save and manage multiple versions of formatted documents.
    
    ### Why It Matters
    DocuMorph AI eliminates repetitive formatting tasks, enhances productivity, and ensures brand/template consistency. It empowers users to focus on content creation rather than formatting, making it an invaluable tool for Tier 2/3 institutions and global content creators.
    
    Developed as part of **HackVyuha 2025** by Team Aerowire.
    """)

# Footer
st.markdown("""
---
<p style='text-align: center; color: #666;'>DocuMorph AI | Powered by Streamlit, Groq, and Hugging Face | HackVyuha 2025</p>
""", unsafe_allow_html=True)