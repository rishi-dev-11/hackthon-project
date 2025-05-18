import streamlit as st
import os
import uuid
import logging
from pymongo import MongoClient
import tempfile
import json
import base64
from datetime import datetime
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
import io
import pytesseract
import spacy
nlp = spacy.load("en_core_web_sm")
from langdetect import detect
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import magic
from transformers import M2M100ForConditionalGeneration, M2M100Tokenizer

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

st.set_page_config(page_title="DocuMorph AI", layout="wide")

# Check for Unstructured module
try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.staging.base import convert_to_dict
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logger.warning("Unstructured module not available.")

# Initialize NLP model
try:
    nlp = spacy.load("en_core_web_sm")
    NLP_AVAILABLE = True
except ImportError:
    NLP_AVAILABLE = False
    logger.warning("Spacy model not available.")

# Initialize MongoDB client
@st.cache_resource
def init_mongodb():
    try:
        client = MongoClient("mongodb://localhost:27017/")
        db = client["documorph_db"]
        templates_collection = db.get_collection("templates")
        documents_collection = db.get_collection("documents")
        client.server_info()
        logger.info("MongoDB connection established.")
        return templates_collection, documents_collection
    except Exception as e:
        logger.error(f"MongoDB connection error: {e}", exc_info=True)
        st.error(f"MongoDB connection error: {str(e)}")
        return None, None

# Load translation model
@st.cache_resource
def load_translation_model():
    try:
        model = M2M100ForConditionalGeneration.from_pretrained("facebook/m2m100_418M")
        tokenizer = M2M100Tokenizer.from_pretrained("facebook/m2m100_418M")
        logger.info("Translation model loaded successfully.")
        return model, tokenizer
    except Exception as e:
        logger.error(f"Translation model loading error: {e}")
        st.error(f"Translation model loading error: {str(e)}")
        return None, None

# Translation function
def translate(text, src_lang, tgt_lang, model, tokenizer):
    if not text or not model or not tokenizer:
        return text
    try:
        tokenizer.src_lang = src_lang
        encoded = tokenizer(text, return_tensors="pt")
        generated_tokens = model.generate(**encoded, forced_bos_token_id=tokenizer.get_lang_id(tgt_lang))
        return tokenizer.decode(generated_tokens[0], skip_special_tokens=True)
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return text

# OCR function
def extract_image_text(image_bytes, lang='eng'):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        img_array = np.array(image.convert('RGB'))
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        temp_path = f"temp_image_{uuid.uuid4()}.png"
        cv2.imwrite(temp_path, thresh)
        
        text = ""
        if UNSTRUCTURED_AVAILABLE:
            try:
                elements = partition_pdf(temp_path, strategy="hi_res")
                text = "\n".join([el.text for el in elements if hasattr(el, 'text') and el.text])
            except Exception as e:
                logger.warning(f"Unstructured failed: {e}")
        if not text and pytesseract:
            text = pytesseract.image_to_string(Image.open(temp_path), lang=lang).strip()
        os.remove(temp_path)
        return text if text else "No text detected"
    except Exception as e:
        logger.error(f"OCR error: {e}")
        return f"Error: {str(e)}"

# NLP-based section detection
def detect_sections(doc_text, lang):
    if lang != 'en' or not NLP_AVAILABLE:
        return []
    doc = nlp(doc_text)
    sections = []
    for sent in doc.sents:
        text = sent.text.strip()
        if len(text.split()) <= 10 and text.lower() in ["abstract", "introduction", "references", "conclusion"]:
            sections.append({"text": text, "type": "Heading"})
    return sections

# Google Docs export
def export_to_google_docs(file_path, file_name):
    try:
        creds = None
        if os.path.exists('credentials.json'):
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', ['https://www.googleapis.com/auth/drive.file'])
            creds = flow.run_local_server(port=0)
        if not creds:
            return "Google Docs export failed: Missing credentials."
        
        service = build('drive', 'v3', credentials=creds)
        file_metadata = {'name': file_name, 'mimeType': 'application/vnd.google-apps.document'}
        media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        return f"Exported to Google Docs: {file.get('id')}"
    except Exception as e:
        logger.error(f"Google Docs export error: {e}")
        return f"Error: {str(e)}"

def apply_template_to_document(content_doc, template_config, lang, logo_path=None):
    output_doc = Document()
    
    # Apply document-wide settings
    section = output_doc.sections[0]
    section.page_height = Inches(float(template_config.get('page_height', 11)))
    section.page_width = Inches(float(template_config.get('page_width', 8.5)))
    section.left_margin = Inches(float(template_config.get('margin_left', 1)))
    section.right_margin = Inches(float(template_config.get('margin_right', 1)))
    section.top_margin = Inches(float(template_config.get('margin_top', 1)))
    section.bottom_margin = Inches(float(template_config.get('margin_bottom', 1)))
    
    # Configure styles
    styles = output_doc.styles
    body_style = styles['Normal']
    body_font = body_style.font
    body_font.name = template_config.get('body_font', 'Calibri')
    body_font.size = Pt(int(template_config.get('body_size', 11)))
    
    # Bullet styles
    if template_config.get('bullet_style', 'default') != 'default':
        paragraph_format = body_style.paragraph_format
        if template_config['bullet_style'] == 'circle':
            body_style._element.xpath('.//w:buChar')[0].set(qn('w:char'), 'F0B7')  # Circle bullet
        elif template_config['bullet_style'] == 'square':
            body_style._element.xpath('.//w:buChar')[0].set(qn('w:char'), 'F0A7')  # Square bullet
    
    # Add logo to header if provided
    if logo_path and template_config.get('use_logo', False):
        for section in output_doc.sections:
            header = section.header
            header_para = header.paragraphs[0]
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(1))
    
    # Title page
    if template_config.get('use_title_page', False):
        title = output_doc.add_heading(template_config.get('document_title', 'Document Title'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if 'document_author' in template_config:
            author_para = output_doc.add_paragraph(template_config['document_author'], style='Normal')
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if template_config.get('include_date', False):
            date_formats = {
                'en': "%B %d, %Y",
                'fr': "%d %B %Y",
                'de': "%d. %B %Y",
                'es': "%d de %B de %Y"
            }
            date_str = datetime.now().strftime(date_formats.get(lang, "%Y-%m-%d"))
            date_para = output_doc.add_paragraph(date_str, style='Normal')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        output_doc.add_page_break()
    
    # Table of contents
    if template_config.get('include_toc', False):
        toc_title = template_config.get('toc_title', 'Table of Contents')
        output_doc.add_heading(toc_title, level=1)
        toc_paragraph = output_doc.add_paragraph()
        run = toc_paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._element.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        output_doc.add_page_break()
    
    # Process content with NLP section detection (only for English)
    full_text = "\n".join([p.text for p in content_doc.paragraphs if p.text.strip()])
    sections = detect_sections(full_text, lang)
    section_index = 0
    
    figure_count = 1
    table_count = 1
    for paragraph in content_doc.paragraphs:
        if not paragraph.text.strip():
            continue
        if lang == 'en' and section_index < len(sections) and paragraph.text.strip().lower() == sections[section_index]["text"].lower():
            output_doc.add_heading(paragraph.text, level=1)
            section_index += 1
        elif paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            output_doc.add_heading(paragraph.text, level=level)
        else:
            new_para = output_doc.add_paragraph(paragraph.text, style='Normal')
            new_para.paragraph_format.line_spacing = float(template_config.get('line_spacing', 1.15))
            new_para.paragraph_format.space_after = Pt(int(template_config.get('para_spacing', 10)))
    
    # Process tables with captions
    for table in content_doc.tables:
        new_table = output_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = template_config.get('table_style', 'Table Grid')
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text
        caption_prefix = template_config.get('table_caption_prefix', 'Table')
        caption_para = output_doc.add_paragraph(f"{caption_prefix} {table_count}: Table")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if template_config.get('caption_position', 'Below') == 'Above':
            caption_para._element.getparent().move(caption_para._element, new_table._element.getprevious())
        table_count += 1
    
    # Add headers/footers
    if template_config.get('use_header', False) or template_config.get('use_footer', False):
        for section in output_doc.sections:
            if template_config.get('use_header', False):
                header = section.header
                header_para = header.paragraphs[0]
                header_para.text = template_config.get('header_text', '')
            if template_config.get('use_footer', False):
                footer = section.footer
                footer_para = footer.paragraphs[0]
                footer_para.text = template_config.get('footer_text', '')
                if template_config.get('include_page_numbers', False):
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = footer_para.add_run()
                    run.add_text(" Page ")
                    fld_char = OxmlElement('w:fldChar')
                    fld_char.set(qn('w:fldCharType'), 'begin')
                    run._element.append(fld_char)
                    instr_text = OxmlElement('w:instrText')
                    instr_text.set(qn('xml:space'), 'preserve')
                    instr_text.text = "PAGE"
                    run._element.append(instr_text)
                    fld_char = OxmlElement('w:fldChar')
                    fld_char.set(qn('w:fldCharType'), 'end')
                    run._element.append(fld_char)
    
    return output_doc

def process_document(uploaded_file, template_config, logo_file=None):
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(uploaded_file.getvalue())
    logo_path = None
    
    if logo_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_logo:
            tmp_logo.write(logo_file.getvalue())
            logo_path = tmp_logo.name
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            tmp_file.write(uploaded_file.getvalue())
            input_doc = Document(tmp_file.name)
            text = " ".join([p.text for p in input_doc.paragraphs if p.text.strip()])
        elif file_type == 'text/plain':
            text = uploaded_file.getvalue().decode('utf-8')
            input_doc = Document()
            input_doc.add_paragraph(text)
            input_doc.save(tmp_file.name)
        elif file_type in ['application/pdf', 'image/png', 'image/jpeg']:
            text = extract_image_text(uploaded_file.getvalue(), lang='eng')
            input_doc = Document()
            input_doc.add_paragraph(text)
            input_doc.save(tmp_file.name)
        else:
            return None
        
        try:
            lang = detect(text)
        except:
            lang = 'en'
        
        if lang != 'en':
            model, tokenizer = load_translation_model()
            text_fields = ['header_text', 'footer_text', 'document_title', 'document_author', 'figure_caption_prefix', 'table_caption_prefix']
            for field in text_fields:
                if field in template_config:
                    template_config[field] = translate(template_config[field], 'en', lang, model, tokenizer)
            template_config['toc_title'] = translate("Table of Contents", 'en', lang, model, tokenizer)
        else:
            template_config['toc_title'] = "Table of Contents"
            template_config['figure_caption_prefix'] = template_config.get('figure_caption_prefix', 'Figure')
            template_config['table_caption_prefix'] = template_config.get('table_caption_prefix', 'Table')
        
        output_doc = apply_template_to_document(input_doc, template_config, lang, logo_path)
        output_path = os.path.join(tempfile.gettempdir(), f"formatted_{uploaded_file.name}.docx")
        output_doc.save(output_path)
        
        os.unlink(tmp_file.name)
        if logo_path:
            os.unlink(logo_path)
        return output_path

def create_template_form():
    template = {}
    st.subheader("Template Configuration")
    template_name = st.text_input("Template Name", "My Template")
    template["name"] = template_name
    template["shared"] = st.checkbox("Share Template with Team", value=False)
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Page Setup", "Typography", "Content Structure", "Tables & Figures", "Headers & Footers"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            template["page_width"] = st.number_input("Page Width (inches)", min_value=5.0, max_value=14.0, value=8.5, step=0.1)
            template["page_height"] = st.number_input("Page Height (inches)", min_value=5.0, max_value=20.0, value=11.0, step=0.1)
        with col2:
            template["margin_left"] = st.number_input("Left Margin (inches)", min_value=0.2, max_value=3.0, value=1.0, step=0.1)
            template["margin_right"] = st.number_input("Right Margin (inches)", min_value=0.2, max_value=3.0, value=1.0, step=0.1)
            template["margin_top"] = st.number_input("Top Margin (inches)", min_value=0.2, max_value=3.0, value=1.0, step=0.1)
            template["margin_bottom"] = st.number_input("Bottom Margin (inches)", min_value=0.2, max_value=3.0, value=1.0, step=0.1)
    
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            fonts = ["Calibri", "Arial", "Times New Roman", "Helvetica", "Georgia"]
            template["body_font"] = st.selectbox("Body Text Font", fonts, index=0)
            template["body_size"] = st.number_input("Body Text Size (pt)", min_value=8, max_value=16, value=11)
            template["line_spacing"] = st.number_input("Line Spacing", min_value=1.0, max_value=3.0, value=1.15, step=0.05)
            template["para_spacing"] = st.number_input("Paragraph Spacing (pt)", min_value=0, max_value=24, value=10)
            template["bullet_style"] = st.selectbox("Bullet Style", ["default", "circle", "square"], index=0)
    
    with tab3:
        template["use_title_page"] = st.checkbox("Include Title Page", value=True)
        if template["use_title_page"]:
            template["document_title"] = st.text_input("Default Document Title", "Document Title")
            template["document_author"] = st.text_input("Default Author", "Author Name")
            template["include_date"] = st.checkbox("Include Date on Title Page", value=True)
        template["include_toc"] = st.checkbox("Include Table of Contents", value=True)
    
    with tab4:
        template["table_style"] = st.selectbox("Table Style", ["Table Grid", "Table Normal", "Light Shading"], index=0)
        template["caption_position"] = st.radio("Caption Position", ["Below", "Above"], horizontal=True)
        template["figure_caption_prefix"] = st.text_input("Figure Caption Prefix", "Figure")
        template["table_caption_prefix"] = st.text_input("Table Caption Prefix", "Table")
    
    with tab5:
        template["use_header"] = st.checkbox("Use Header", value=False)
        if template["use_header"]:
            template["header_text"] = st.text_input("Header Text", "Document Header")
        template["use_footer"] = st.checkbox("Use Footer", value=True)
        if template["use_footer"]:
            template["footer_text"] = st.text_input("Footer Text", "")
            template["include_page_numbers"] = st.checkbox("Include Page Numbers", value=True)
        template["use_logo"] = st.checkbox("Use Logo in Header", value=False)
    
    return template

def main():
    st.title("DocuMorph AI")
    st.subheader("The Intelligent Document Transformation Engine")
    
    if 'user_id' not in st.session_state:
        st.session_state.user_id = str(uuid.uuid4())
    if 'templates' not in st.session_state:
        st.session_state.templates = []
    if 'current_template' not in st.session_state:
        st.session_state.current_template = None
    if 'formatted_doc_path' not in st.session_state:
        st.session_state.formatted_doc_path = None
    
    templates_collection, documents_collection = init_mongodb()
    
    with st.sidebar:
        st.header("Template Management")
        template_action = st.radio("Template Actions", ["Select Template", "Create New Template", "Edit Template"], horizontal=True)
        
        if template_action == "Select Template":
            if templates_collection:
                user_templates = list(templates_collection.find({"$or": [{"user_id": st.session_state.user_id}, {"shared": True}]}))
                template_names = [t["name"] for t in user_templates]
                if template_names:
                    selected_template_name = st.selectbox("Select a template", template_names)
                    selected_template = next((t for t in user_templates if t["name"] == selected_template_name), None)
                    if selected_template:
                        st.session_state.current_template = selected_template
                        st.success(f"Template '{selected_template_name}' loaded!")
                else:
                    st.info("No templates found. Please create a new template.")
        
        elif template_action == "Create New Template":
            with st.form("template_form"):
                template = create_template_form()
                if st.form_submit_button("Save Template"):
                    if templates_collection:
                        template["user_id"] = st.session_state.user_id
                        template["created_at"] = datetime.now()
                        templates_collection.insert_one(template)
                        st.session_state.current_template = template
                        st.success(f"Template '{template['name']}' saved successfully!")
        
        elif template_action == "Edit Template":
            if templates_collection:
                user_templates = list(templates_collection.find({"user_id": st.session_state.user_id}))
                template_names = [t["name"] for t in user_templates]
                if template_names:
                    selected_template_name = st.selectbox("Select template to edit", template_names)
                    selected_template = next((t for t in user_templates if t["name"] == selected_template_name), None)
                    if selected_template:
                        with st.form("edit_template_form"):
                            template = create_template_form()
                            if st.form_submit_button("Update Template"):
                                template["user_id"] = st.session_state.user_id
                                template["updated_at"] = datetime.now()
                                templates_collection.update_one({"_id": selected_template["_id"]}, {"$set": template})
                                st.session_state.current_template = template
                                st.success(f"Template '{template['name']}' updated successfully!")
    
    st.header("Document Transformation")
    uploaded_file = st.file_uploader("Upload unformatted document", type=["docx", "txt", "pdf", "png", "jpg"])
    logo_file = st.file_uploader("Upload logo (optional)", type=["png", "jpg"])
    
    if uploaded_file is not None:
        st.success(f"File '{uploaded_file.name}' uploaded successfully!")
        if st.session_state.current_template:
            st.info(f"Using template: {st.session_state.current_template['name']}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Transform Document"):
                    with st.spinner("Applying template and transforming document..."):
                        output_path = process_document(uploaded_file, st.session_state.current_template, logo_file)
                        if output_path:
                            st.session_state.formatted_doc_path = output_path
                            if documents_collection:
                                documents_collection.insert_one({
                                    "user_id": st.session_state.user_id,
                                    "file_name": uploaded_file.name,
                                    "output_path": output_path,
                                    "version": 1,
                                    "created_at": datetime.now()
                                })
                            st.success("Document transformed successfully!")
            with col2:
                if st.button("Export to Google Docs"):
                    result = export_to_google_docs(st.session_state.formatted_doc_path, uploaded_file.name)
                    st.write(result)
            
            if st.session_state.formatted_doc_path:
                with open(st.session_state.formatted_doc_path, "rb") as file:
                    st.download_button(
                        label="Download Formatted Document",
                        data=file,
                        file_name=f"formatted_{uploaded_file.name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning("Please select or create a template first.")

if __name__ == "__main__":
    main()