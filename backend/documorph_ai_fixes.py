import os
import io
import tempfile
import base64
import numpy as np
from PIL import Image
import pandas as pd
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import fitz  # PyMuPDF
from pymongo import MongoClient
import uuid

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Fix for PyMuPDF compatibility
def open_pdf(pdf_path):
    """Safely open a PDF with PyMuPDF using the correct method."""
    try:
        return fitz.open(pdf_path)
    except AttributeError:
        # Fallback for older versions
        try:
            return fitz.Document(pdf_path)
        except AttributeError:
            # Another fallback
            from pymupdf import Document
            return Document(pdf_path)

# Enhanced table extraction
def extract_tables_from_pdf(pdf_path):
    """Extract tables from a PDF document with better handling."""
    tables = []
    try:
        doc = open_pdf(pdf_path)
        for page_num, page in enumerate(doc):
            # Extract tables using PyMuPDF's built-in table detection
            tab_dict = page.find_tables()
            if tab_dict and tab_dict.tables:
                for i, table in enumerate(tab_dict.tables):
                    cells = table.extract()
                    rows = []
                    
                    # Process cells properly, handling various data types
                    for row in cells:
                        row_data = []
                        for cell in row:
                            # Handle different cell types
                            if hasattr(cell, 'text'):
                                row_data.append(cell.text)
                            elif isinstance(cell, str):
                                row_data.append(cell)
                            elif cell is None:
                                row_data.append("")
                            else:
                                row_data.append(str(cell))
                        rows.append(row_data)
                    
                    # Create pandas DataFrame with proper error handling
                    df = None
                    raw_data = rows
                    if rows:
                        try:
                            # Ensure we have at least one row for column headers
                            if len(rows) > 1:
                                df = pd.DataFrame(rows[1:], columns=rows[0])
                            else:
                                # Create single-row table with column indices
                                df = pd.DataFrame([rows[0]])
                        except Exception as e:
                            logger.warning(f"Could not create DataFrame for table: {e}")
                        
                    tables.append({
                        "page": page_num + 1,
                        "table_id": f"table_{page_num + 1}_{i + 1}",
                        "dataframe": df,
                        "raw_data": raw_data,
                        "rect": [table.rect.x0, table.rect.y0, table.rect.x1, table.rect.y1]
                    })
                    
                    # Store actual data for display
                    if df is not None:
                        tables[-1]["html"] = df.to_html(index=False)
                    else:
                        # Create HTML table manually if DataFrame creation failed
                        html = "<table>"
                        for row in raw_data:
                            html += "<tr>"
                            for cell in row:
                                html += f"<td>{cell}</td>"
                            html += "</tr>"
                        html += "</table>"
                        tables[-1]["html"] = html
                
        return tables
    except Exception as e:
        logger.error(f"Error extracting tables from PDF: {e}", exc_info=True)
        return []

# Enhanced figure extraction
def extract_figures_from_pdf(pdf_path):
    """Extract figures/images from a PDF document with better handling."""
    figures = []
    try:
        doc = open_pdf(pdf_path)
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    # Extract image with proper error handling
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Verify image data is valid
                    if not image_bytes:
                        logger.warning(f"Empty image data for image {img_index} on page {page_num+1}")
                        continue
                    
                    # Get image location on page
                    bbox = None
                    for img_obj in page.get_images(full=True):
                        if img_obj[0] == xref:
                            bbox = page.get_image_bbox(img_obj)
                            break
                    
                    # Create figure object with all necessary data
                    image_type = base_image["ext"].lower()
                    
                    # Validate image by attempting to open it
                    try:
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        width, height = pil_image.size
                        
                        # Create base64 representation for preview
                        img_b64 = base64.b64encode(image_bytes).decode('utf-8')
                        preview_src = f"data:image/{image_type};base64,{img_b64}"
                        
                        figures.append({
                            "page": page_num + 1,
                            "figure_id": f"figure_{page_num + 1}_{img_index + 1}",
                            "image_bytes": image_bytes,
                            "rect": bbox,
                            "width": width,
                            "height": height,
                            "image_type": image_type,
                            "preview_src": preview_src
                        })
                    except Exception as e:
                        logger.warning(f"Failed to validate image {img_index} on page {page_num+1}: {e}")
                        continue
                        
                except Exception as e:
                    logger.warning(f"Error extracting image {img_index} on page {page_num+1}: {e}")
                    continue
            
        return figures
    except Exception as e:
        logger.error(f"Error extracting figures from PDF: {e}", exc_info=True)
        return []

# Improved table integration into document
def add_table_to_docx(doc, table_data, caption):
    """Add a table to a DOCX document with proper formatting and caption."""
    try:
        # First add the caption
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.style = 'Caption'
        
        # Then add the table
        if "raw_data" in table_data and table_data["raw_data"]:
            rows = table_data["raw_data"]
            
            num_rows = len(rows)
            num_cols = max(len(row) for row in rows)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # Fill the table with data
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:  # Ensure we don't exceed columns
                        cell = table.cell(i, j)
                        cell.text = str(cell_text)
            
            # Format the header row if present
            if num_rows > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    cell_para = cell.paragraphs[0]
                    # Use a safe way to format text
                    if cell_para.runs:
                        for run in cell_para.runs:
                            run.bold = True
                    else:
                        run = cell_para.add_run(cell.text)
                        run.bold = True
                        cell.text = ""
            
            return True
        elif "dataframe" in table_data and table_data["dataframe"] is not None:
            # Alternative: use DataFrame if raw_data is not available
            df = table_data["dataframe"]
            table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Add headers
            for j, column in enumerate(df.columns):
                table.cell(0, j).text = str(column)
                
            # Add data
            for i, row in enumerate(df.itertuples(index=False)):
                for j, value in enumerate(row):
                    table.cell(i+1, j).text = str(value)
                    
            return True
        else:
            logger.warning("No table data available to add to document")
            p = doc.add_paragraph()
            p.add_run(f"[Table placeholder: {caption}]").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return False
            
    except Exception as e:
        logger.error(f"Error adding table to DOCX: {e}", exc_info=True)
        p = doc.add_paragraph()
        p.add_run(f"[Table error: {str(e)}]").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return False

# Improved figure integration into document
def add_figure_to_docx(doc, figure, caption):
    """Add a figure to a DOCX document with proper caption and formatting."""
    try:
        # Verify that image_bytes exists and is not None
        if "image_bytes" not in figure or figure["image_bytes"] is None:
            logger.warning(f"No image bytes for figure {figure.get('figure_id', 'unknown')}")
            p = doc.add_paragraph()
            p.add_run(f"[Figure placeholder: {caption}]").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the caption even if the image couldn't be loaded
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            return True
            
        # Add the image to the document
        p = doc.add_paragraph()
        r = p.add_run()
        
        # Create a temporary file to store the image
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{figure['image_type']}") as tmp:
            tmp_path = tmp.name
            tmp.write(figure["image_bytes"])
            
        try:
            # Add the image to the document with appropriate width
            width_inches = min(6, figure.get("width", 500) / 96)  # Convert pixels to inches, max 6 inches
            r.add_picture(tmp_path, width=Inches(width_inches))
            
            # Center the image
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the caption
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
            # Clean up temp file
            os.unlink(tmp_path)
            return True
        except Exception as e:
            logger.error(f"Error adding image to document: {e}")
            
            # Add a placeholder instead
            r.add_text(f"[Image could not be displayed: {e}]")
            r.italic = True
            
            # Add the caption anyway
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
            # Clean up temp file if it exists
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
            return False
    except Exception as e:
        logger.error(f"Error adding figure to DOCX: {e}", exc_info=True)
        
        # Add a placeholder instead
        p = doc.add_paragraph()
        p.add_run(f"[Figure error: {str(e)}]").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the caption anyway
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.style = 'Caption'
        
        return False

# Default templates for MongoDB
def create_default_templates():
    """Create default templates for different roles."""
    templates = [
        {
            "name": "Student Essay",
            "role": "Student",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Times New Roman",
            "body_font_size": 12,
            "heading_font": "Times New Roman",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 2.0,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "Page [Page] of [Pages]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        },
        {
            "name": "Business Report",
            "role": "Business Professional",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Company Name]",
            "footer_text": "Confidential | [Date] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Blue Professional"
        },
        {
            "name": "Research Paper",
            "role": "Researcher",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Times New Roman",
            "body_font_size": 12,
            "heading_font": "Times New Roman",
            "heading1_font_size": 14,
            "heading2_font_size": 12,
            "heading3_font_size": 12,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "[Page]",
            "include_tables_figures": True,
            "color_scheme": "Academic"
        },
        {
            "name": "Blog Post",
            "role": "Content Creator",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Georgia",
            "heading1_font_size": 18,
            "heading2_font_size": 16,
            "heading3_font_size": 14,
            "line_spacing": 1.5,
            "paragraph_spacing": 12,
            "header_text": "",
            "footer_text": "© [Year] [Author Name]",
            "include_tables_figures": True,
            "color_scheme": "Modern Gray"
        },
        {
            "name": "Multilingual Document",
            "role": "Multilingual User",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Arial",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "[Page]/[Pages]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        },
        {
            "name": "Book Chapter",
            "role": "Author",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Georgia",
            "body_font_size": 12,
            "heading_font": "Georgia",
            "heading1_font_size": 18,
            "heading2_font_size": 16,
            "heading3_font_size": 14,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "[Book Title] - Chapter [Chapter]",
            "footer_text": "[Author Name] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Earthy Tones"
        },
        {
            "name": "Team Document",
            "role": "Collaborator",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Team Name] - [Project Name]",
            "footer_text": "Last updated: [Date] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Blue Professional"
        },
        {
            "name": "Project Report",
            "role": "Project Manager",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Arial",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Project Name] - [Status]",
            "footer_text": "Confidential | Page [Page] of [Pages]",
            "include_tables_figures": True,
            "color_scheme": "Modern Gray"
        },
        {
            "name": "Custom Document",
            "role": "Others",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        }
    ]
    
    return templates

def initialize_db_templates(user_id):
    """Initialize MongoDB with default templates."""
    try:
        client = MongoClient("mongodb://localhost:27017/")
        db = client["documorph_db"]
        templates_collection = db["templates"]
        
        # Check if templates already exist for this user
        existing_templates = list(templates_collection.find({"user_id": user_id}))
        if existing_templates:
            logger.info(f"Templates already exist for user {user_id}")
            return True
            
        # Add default templates for this user
        templates = create_default_templates()
        for template in templates:
            template["user_id"] = user_id
            template["template_id"] = str(uuid.uuid4())
            template["is_custom"] = False
            templates_collection.insert_one(template)
            
        logger.info(f"Initialized default templates for user {user_id}")
        return True
    except Exception as e:
        logger.error(f"Error initializing DB templates: {e}", exc_info=True)
        return False

# Google Drive integration with proper redirect URI
def setup_google_drive_auth(client_id, client_secret):
    """Create Google Drive authentication flow with proper redirect URI."""
    try:
        # Import Google API libraries
        from google.oauth2.credentials import Credentials
        from google_auth_oauthlib.flow import InstalledAppFlow
        
        # Define proper scopes and redirect URI
        SCOPES = ['https://www.googleapis.com/auth/drive.file']
        REDIRECT_URI = "http://localhost:8000/oauth2callback"
        
        # Create OAuth flow
        flow = InstalledAppFlow.from_client_config(
            {
                "installed": {
                    "client_id": client_id,
                    "project_id": "documorph-ai",
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                    "client_secret": client_secret,
                    "redirect_uris": [REDIRECT_URI, "http://localhost", "urn:ietf:wg:oauth:2.0:oob"]
                }
            },
            SCOPES,
            redirect_uri=REDIRECT_URI
        )
        return flow
    except ImportError:
        logger.error("Google API client libraries not installed")
        return None
    except Exception as e:
        logger.error(f"Error setting up Google Drive auth: {e}", exc_info=True)
        return None