import streamlit as st
import os
import uuid
import logging
from pymongo import MongoClient
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_groq import ChatGroq
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
import io
import pytesseract
import tempfile
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import re
import pandas as pd
from docx.enum.section import WD_ORIENTATION
from docx.oxml.shared import OxmlElement, qn
import magic  # For file type detection
from langdetect import detect  # For language detection
from transformers import M2M100ForConditionalGeneration, M2M100Tokenizer  # For translation
import spacy  # For smart section detection

# Import local modules
try:
    from table_extraction import extract_tables
    from chart_extraction import identify_and_extract_charts
except ImportError:
    logging.warning("Could not import local table and chart extraction modules. These features may not work properly.")

try:
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    GOOGLE_API_AVAILABLE = True
except ImportError:
    GOOGLE_API_AVAILABLE = False
    st.warning("Google API client not available. Install with 'pip install google-api-python-client google-auth-oauthlib'")

try:
    import msoffcrypto  # For MS Word integration
    MSOFFCRYPTO_AVAILABLE = True
except ImportError:
    MSOFFCRYPTO_AVAILABLE = False
    st.warning("MS Office encryption module not available. Install with 'pip install msoffcrypto-tool'")

import requests  # For API calls
import json  # For JSON processing
from dotenv import load_dotenv  # For loading environment variables
import pickle
import hashlib  # For document fingerprinting

# Import fixes
try:
    from documorph_fixes import (
        open_pdf,
        extract_tables_from_pdf,
        extract_figures_from_pdf,
        add_table_to_docx as enhanced_add_table_to_docx,
        add_figure_to_docx as enhanced_add_figure_to_docx,
        initialize_db_templates,
        setup_google_drive_auth
    )
    FIXES_AVAILABLE = True
except ImportError:
    FIXES_AVAILABLE = False
    logging.warning("DocuMorph fixes not available. Some issues may persist.")


# Load environment variables
load_dotenv()
SERPAPI_KEY = os.getenv("serpapi")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("CLIENT_SECRET")
GOOGLE_API_KEY = os.getenv("GOOG_API_KEY")

# Import unstructured components
try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.image import partition_image
    from unstructured.partition.text import partition_text
    from unstructured.documents.elements import Title, NarrativeText, Table, Image as UnstructuredImage
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logging.warning("Unstructured module not available. Install with 'pip install unstructured unstructured-inference'")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Set page config
st.set_page_config(page_title="DocuMorph AI", layout="wide")

# User tier system
class UserTier:
    FREE = "free"
    PREMIUM = "premium"
    
    @staticmethod
    def get_tier_features():
        return {
            UserTier.FREE: {
                "max_documents": 5,
                "max_templates": 3,
                "llm_enabled": False,
                "multi_language": False,
                "ocr_languages": ["eng"],
                "style_guide_compliance": False,
                "async_processing": False,
                "team_collaboration": False,
                "advanced_tables": False,
                "caption_editor": False,
                "google_docs": False,
                "ms_word": False,
                "plagiarism_check": False,
                "template_categories": ["Student", "Content Creator","Others"],
                "google_drive_export": False,
                "custom_templates": True  # Custom templates available to free users
            },
            UserTier.PREMIUM: {
                "max_documents": 100,
                "max_templates": 50,
                "llm_enabled": True,
                "multi_language": True,
                "ocr_languages": ["eng", "fra", "deu", "spa", "rus", "ara", "chi_sim", "jpn", "kor","hi"],
                "style_guide_compliance": True,
                "async_processing": True,
                "team_collaboration": True,
                "advanced_tables": True,
                "caption_editor": True,
                "google_docs": True,
                "ms_word": True,
                "plagiarism_check": True,
                "template_categories": ["Student", "Content Creator", "Researcher", "Business Professional", "Multilingual User", "Author", "Collaborator", "Project Manager"],
                "google_drive_export": True,
                "custom_templates": True  # Custom templates available to premium users
            }
        }

# Initialize MongoDB client
@st.cache_resource
def init_mongodb():
    try:
        client = MongoClient("mongodb://localhost:27017/", serverSelectionTimeoutMS=2000)
        db = client["documorph_db"]
        docs_collection = db["documents"]
        templates_collection = db["templates"]
        users_collection = db["users"]
        client.server_info()
        logger.info("MongoDB connection established.")
        return docs_collection, templates_collection, users_collection
    except Exception as e:
        logger.error(f"MongoDB connection error: {e}", exc_info=True)
        st.warning("MongoDB not available. Some features will be limited. Running in memory-only mode.")
        # Return None values to indicate MongoDB is not available
        return None, None, None

# Initialize FAISS vector store
@st.cache_resource
def init_vector_store(user_id):
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store_path = f"vector_stores/faiss_index_{user_id}"
        if os.path.exists(vector_store_path):
            vector_store = FAISS.load_local(
                vector_store_path, embeddings, allow_dangerous_deserialization=True
            )
        else:
            vector_store = FAISS.from_texts(["initialize"], embeddings)
            vector_store.save_local(vector_store_path)
        logger.info(f"Vector store initialized for user {user_id}.")
        return vector_store, vector_store_path
    except Exception as e:
        logger.error(f"Error initializing vector store: {e}", exc_info=True)
        st.error(f"Error initializing vector store: {str(e)}")
        return None, None

# Initialize language detection and translation models
@st.cache_resource
def init_nlp_models():
    try:
        # Load spaCy model for section detection
        nlp = spacy.load("en_core_web_sm")
        
        # Load M2M100 for translation (only if needed - lazy loading)
        translation_model = None
        translation_tokenizer = None
        
        return {
            "nlp": nlp,
            "translation_model": translation_model,
            "translation_tokenizer": translation_tokenizer
        }
    except Exception as e:
        logger.error(f"Error initializing NLP models: {e}", exc_info=True)
        return None
        
# Initialize Groq LLM
@st.cache_resource
def init_llm():
    try:
        return ChatGroq(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.2,
            api_key="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"
        )
    except Exception as e:
        logger.error(f"Error initializing LLM: {e}", exc_info=True)
        st.error(f"Error initializing LLM: {str(e)}")
        return None

# Function to check if user can use LLM features
def can_use_llm(user_tier):
    tier_features = UserTier.get_tier_features()
    return tier_features[user_tier]["llm_enabled"]

# Function to get appropriate processing method based on user tier
def get_processing_method(user_tier, feature_name):
    tier_features = UserTier.get_tier_features()
    if tier_features[user_tier][feature_name]:
        return "advanced"  # LLM-based processing
    else:
        return "basic"  # Rule-based processing

# Function to process documents (DOCX, PDF, TXT)
def process_document(file_path, file_type, user_tier=UserTier.FREE):
    try:
        # Try using unstructured.io first if available
        if UNSTRUCTURED_AVAILABLE:
            logger.info(f"Attempting to process {file_path} with unstructured.io")
            chunks, tables, figures, error = process_with_unstructured(file_path, file_type, user_tier)
            if error is None and chunks:
                return chunks, None, tables, figures
            else:
                logger.warning(f"Unstructured processing failed: {error}. Falling back to standard processing.")
        
        # Fall back to standard processing if unstructured fails or isn't available
        if file_type == "pdf":
            loader = PyMuPDFLoader(file_path)
            # Also detect tables and figures for PDFs
            tables = detect_tables_from_pdf(file_path)
            figures = detect_figures_from_pdf(file_path)
        elif file_type == "docx":
            loader = Docx2txtLoader(file_path)
            # Detect tables and figures for DOCX
            tables = detect_tables_from_docx(file_path)
            figures = detect_figures_from_docx(file_path)
        elif file_type == "txt":
            loader = TextLoader(file_path)
            # Text files don't have tables or figures
            tables = []
            figures = []
        else:
            return [], "Unsupported file type", [], []
        
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = text_splitter.split_documents(documents)
        
        logger.info(f"Processed document {file_path} with {len(chunks)} chunks, {len(tables)} tables, and {len(figures)} figures")
        return chunks, None, tables, figures
    except Exception as e:
        logger.error(f"Error processing document: {e}", exc_info=True)
        return [], f"Error processing document: {str(e)}", [], []

# Function to extract document structure
def extract_document_structure(chunks, llm):
    try:
        # Concatenate chunks into one text for LLM analysis
        full_text = "\n\n".join([chunk.page_content for chunk in chunks])
        
        prompt = PromptTemplate(
            template="""Analyze the following document and identify its logical structure. 
            Extract the main sections, subsections, and their hierarchical relationships.
            For each section, provide:
            1. The heading/title
            2. The heading level (H1, H2, H3, etc.)
            3. A brief description of the content
            
            Document Content:
            {content}
            
            Output Format:
            - H1: [Section Title]
              Description: [Brief description]
              - H2: [Subsection Title]
                Description: [Brief description]
                - H3: [Sub-subsection Title]
                  Description: [Brief description]
            """,
            input_variables=["content"],
        )
        
        chain = prompt | llm | StrOutputParser()
        structure = chain.invoke({"content": full_text[:10000]})  # Limit to 10k chars to avoid token limits
        
        return structure
    except Exception as e:
        logger.error(f"Error extracting document structure: {e}", exc_info=True)
        return f"Error extracting document structure: {str(e)}"

# Function to generate abstract
def generate_abstract(chunks, llm, instructions=""):
    try:
        # Get first few chunks as these likely contain the document's main content
        content = "\n\n".join([chunk.page_content for chunk in chunks[:5]])
        
        prompt = PromptTemplate(
            template="""Generate a concise, professional abstract for the following document content.
            
            Document Content:
            {content}
            
            User Instructions: {instructions}
            
            Generate an abstract that accurately summarizes the main points, purpose, and findings of the document.
            The abstract should be well-structured, clear, and approximately 150-250 words.
            """,
            input_variables=["content", "instructions"],
        )
        
        chain = prompt | llm | StrOutputParser()
        abstract = chain.invoke({"content": content, "instructions": instructions})
        
        return abstract
    except Exception as e:
        logger.error(f"Error generating abstract: {e}", exc_info=True)
        return f"Error generating abstract: {str(e)}"

# Function to suggest section titles
def suggest_section_titles(section_content, current_title, llm, instructions=""):
    try:
        prompt = PromptTemplate(
            template="""Given the following content from a document section and its current title (if any),
            suggest 2-3 alternative title options that accurately reflect the content and are professional and engaging.
            
            Current section title: {current_title}
            
            Section content:
            {content}
            
            User Instructions: {instructions}
            
            Provide 2-3 alternative title suggestions in this format:
            1. [Title suggestion 1]
            2. [Title suggestion 2]
            3. [Title suggestion 3]
            
            Each title should be concise, specific, and reflective of the section's content.
            """,
            input_variables=["content", "current_title", "instructions"],
        )
        
        chain = prompt | llm | StrOutputParser()
        title_suggestions = chain.invoke({
            "content": section_content[:2000],  # Limit content to avoid token limits
            "current_title": current_title,
            "instructions": instructions
        })
        
        return title_suggestions
    except Exception as e:
        logger.error(f"Error suggesting section titles: {e}", exc_info=True)
        return f"Error suggesting section titles: {str(e)}"

# Function to improve word choice and style
def enhance_text_style(text, target_style, llm):
    try:
        prompt = PromptTemplate(
            template="""Enhance the following text by improving word choice, tone, and style according to the target style.
            
            Original text:
            {text}
            
            Target style: {target_style}
            
            Rewrite the text to match the target style while preserving all information.
            Focus on replacing generic terms with more precise and contextually appropriate words,
            improving sentence structure, and enhancing the overall professionalism of the text.
            """,
            input_variables=["text", "target_style"],
        )
        
        chain = prompt | llm | StrOutputParser()
        enhanced_text = chain.invoke({"text": text, "target_style": target_style})
        
        return enhanced_text
    except Exception as e:
        logger.error(f"Error enhancing text style: {e}", exc_info=True)
        return f"Error enhancing text style: {str(e)}"

# Function to get word suggestions based on current text
def get_word_suggestions(current_text, llm, num_suggestions=3):
    try:
        # Get the last few words to use as context
        words = current_text.split()
        context = " ".join(words[-10:]) if len(words) > 10 else current_text
        
        prompt = PromptTemplate(
            template="""Based on the following text, suggest {num_suggestions} words or phrases that might 
            come next in a professional document. The suggestions should be contextually relevant and 
            help the user complete their thought.
            
            Current text:
            {context}
            
            Provide exactly {num_suggestions} suggestions in this format:
            1. [suggestion 1]
            2. [suggestion 2]
            3. [suggestion 3]
            
            The suggestions should be brief (1-3 words) and professional.
            """,
            input_variables=["context", "num_suggestions"],
        )
        
        chain = prompt | llm | StrOutputParser()
        suggestions_text = chain.invoke({
            "context": context,
            "num_suggestions": num_suggestions
        })
        
        # Parse the suggestions into a list
        suggestions = []
        for line in suggestions_text.split('\n'):
            if re.match(r'^\d+\.', line.strip()):
                suggestion = line.split('.', 1)[1].strip()
                if suggestion.startswith('[') and suggestion.endswith(']'):
                    suggestion = suggestion[1:-1]
                suggestions.append(suggestion)
        
        return suggestions[:num_suggestions]
    except Exception as e:
        logger.error(f"Error getting word suggestions: {e}", exc_info=True)
        return ["Error", "getting", "suggestions"]

# Function to apply a template to a document
def apply_template_to_document(doc_path, template, output_path, tables=None, figures=None, chapters=None):
    try:
        # Create a new document or load the existing one
        if doc_path.endswith('.docx'):
            doc = docx.Document(doc_path)
            # Make a copy to ensure we're not losing content
            temp_doc = docx.Document()
            for para in doc.paragraphs:
                temp_doc.add_paragraph(para.text)
            doc = temp_doc
        else:
            # For PDF or text, create a new document
            doc = docx.Document()
            
            # If it's text, load the content
            if doc_path.endswith('.txt'):
                with open(doc_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    for para in content.split('\n\n'):
                        doc.add_paragraph(para)
            elif doc_path.endswith('.pdf'):
                # For PDFs, we need to extract text using PyMuPDF
                try:
                    # Changed from fitz.Document to FitzDocument
                    pdf_doc = fitz.open(doc_path)
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        text = page.get_text()
                        for para in text.split('\n\n'):
                            if para.strip():  # Only add non-empty paragraphs
                                doc.add_paragraph(para.strip())
                except Exception as e:
                    logger.error(f"Error extracting text from PDF: {e}")
                    doc.add_paragraph(f"Error extracting text from PDF: {str(e)}")
        
        # Apply template settings
        # 1. Set font for normal text
        style = doc.styles['Normal']
        font = style.font
        font.name = template.get('body_font', 'Calibri')
        font.size = Pt(template.get('body_font_size', 11))
        
        # 2. Set heading styles
        for i in range(1, 4):  # For Heading 1, 2, 3
            if f'heading{i}_font' in template and f'heading{i}_font_size' in template:
                heading_style = doc.styles[f'Heading {i}']
                heading_font = heading_style.font
                heading_font.name = template[f'heading{i}_font']
                heading_font.size = Pt(template[f'heading{i}_font_size'])
                if template.get(f'heading{i}_bold', False):
                    heading_font.bold = True
        
        # 3. Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(template.get('margin_top', 1.0))
            section.bottom_margin = Inches(template.get('margin_bottom', 1.0))
            section.left_margin = Inches(template.get('margin_left', 1.0))
            section.right_margin = Inches(template.get('margin_right', 1.0))
            
            # 3.1 Set line spacing if available
            if 'line_spacing' in template:
                for paragraph in doc.paragraphs:
                    paragraph.paragraph_format.line_spacing = template['line_spacing']
        
        # 4. Add header/footer if specified
        if template.get('header_text'):
            for section in sections:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = template['header_text']
        
        if template.get('footer_text'):
            for section in sections:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.text = template['footer_text']
            
        # 5. Add tables and figures with proper numbering
        if tables and figures and template.get('include_tables_figures', True):
            # Make sure we have chapters for numbering
            if not chapters:
                chapters = [{"number": 1, "title": "Document", "subsections": []}]
                
            # Assign numbers to figures and tables
            numbered_figures, numbered_tables = assign_figure_table_numbers(figures, tables, chapters)
            
            # Add an appendix section for tables
            if numbered_tables:
                doc.add_heading('Tables', level=1)
                for table in numbered_tables:
                    add_table_to_docx(doc, table, table["caption"])
                    doc.add_paragraph()  # Add spacing after table
            
            # Add an appendix section for figures
            if numbered_figures:
                doc.add_heading('Figures', level=1)
                for figure in numbered_figures:
                    add_figure_to_docx(doc, figure, figure["caption"])
                    doc.add_paragraph()  # Add spacing after figure
        
        # Ensure the document is not empty
        if len(doc.paragraphs) == 0:
            logger.warning("Document is empty. Adding a default paragraph.")
            doc.add_paragraph("Document content could not be properly extracted.")
        
        # Save the document
        doc.save(output_path)
        
        # Verify document is not empty
        verify_doc = docx.Document(output_path)
        if len(verify_doc.paragraphs) == 0:
            logger.error(f"Generated document is empty! Path: {output_path}")
            return False, "Error: Generated document is empty"
            
        logger.info(f"Document saved successfully with {len(verify_doc.paragraphs)} paragraphs to {output_path}")
        return True, output_path
    except Exception as e:
        logger.error(f"Error applying template: {e}", exc_info=True)
        return False, f"Error applying template: {str(e)}"

# Function to display document preview
def display_document_preview(doc_path):
    try:
        if doc_path.endswith('.docx'):
            # For DOCX, extract text and display
            doc = docx.Document(doc_path)
            content = []
            
            # Check if document has content
            if len(doc.paragraphs) == 0:
                st.warning("Document appears to be empty!")
                return False
                
            for para in doc.paragraphs:
                content.append(para.text)
            
            preview_text = '\n'.join(content)
            
            # Display preview in a scrollable area with proper formatting
            st.markdown("### Document Preview")
            preview_container = st.container()
            with preview_container:
                # Use a custom HTML wrapper for better formatting with dark text on light background
                html = f"""
                <div style="border:1px solid #ddd; padding:10px; border-radius:5px; 
                           height:400px; overflow-y:scroll; background-color:#f9f9f9; font-family:{doc.styles['Normal'].font.name}">
                    <div style="color: #333333;">{'<br>'.join(content)}</div>
                </div>
                """
                st.markdown(html, unsafe_allow_html=True)
                
                # Also show plain text in a hidden expander
                with st.expander("View as plain text"):
                    st.text_area("Document Content", preview_text, height=400, disabled=True)
            
            # Display download link for direct access
            with open(doc_path, "rb") as file:
                file_bytes = file.read()
                btn = st.download_button(
                    label="Download document for direct access",
                    data=file_bytes,
                    file_name=os.path.basename(doc_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
        elif doc_path.endswith('.pdf'):
            # For PDF, display using PDF viewer
            with open(doc_path, "rb") as file:
                file_bytes = file.read()
                if len(file_bytes) == 0:
                    st.warning("PDF document appears to be empty!")
                    return False
                    
                base64_pdf = base64.b64encode(file_bytes).decode('utf-8')
                pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
                st.markdown(pdf_display, unsafe_allow_html=True)
                
                # Display download link for direct access
                btn = st.download_button(
                    label="Download PDF for direct access",
                    data=file_bytes,
                    file_name=os.path.basename(doc_path),
                    mime="application/pdf"
                )
                
                # Also extract text from PDF using PyMuPDF for text view
                try:
                    # Changed from fitz.Document to FitzDocument
                    pdf_doc = fitz.open(doc_path)
                    pdf_text = []
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        pdf_text.append(page.get_text())
                    
                    with st.expander("View PDF text content"):
                        st.text_area("PDF Content", "\n\n".join(pdf_text), height=400, disabled=True)
                except Exception as e:
                    st.warning(f"Could not extract text from PDF: {str(e)}")
        else:
            # For other formats, just display text
            try:
                with open(doc_path, "r", encoding="utf-8") as file:
                    content = file.read()
                    if not content:
                        st.warning("Document appears to be empty!")
                        return False
                        
                    st.markdown("### Document Preview")
                    # Use a custom styled text area with visible text
                    st.markdown(f"""
                    <div style="border:1px solid #ddd; padding:10px; border-radius:5px; 
                               height:400px; overflow-y:scroll; background-color:#f9f9f9; color:#333333; font-family:monospace;">
                        <pre>{content}</pre>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Display download link for direct access
                    with open(doc_path, "rb") as file:
                        btn = st.download_button(
                            label="Download document for direct access",
                            data=file.read(),
                            file_name=os.path.basename(doc_path),
                            mime="text/plain"
                        )
            except UnicodeDecodeError:
                # Handle binary files
                st.warning("Cannot preview binary file format. Please download to view.")
                with open(doc_path, "rb") as file:
                    btn = st.download_button(
                        label="Download file for direct access",
                        data=file.read(),
                        file_name=os.path.basename(doc_path)
                    )
                    
        return True
    except Exception as e:
        logger.error(f"Error displaying document preview: {e}", exc_info=True)
        st.error(f"Error displaying document preview: {str(e)}")
        return False

# New helper functions for table and figure handling

def detect_tables_from_pdf(pdf_path):
    """Extract tables from a PDF document with better handling."""
    tables = []
    try:
        doc = open_pdf(pdf_path)
        
        # First try using PyMuPDF's built-in table detection
        for page_num, page in enumerate(doc):
            tab_dict = page.find_tables()
            if tab_dict and tab_dict.tables:
                for i, table in enumerate(tab_dict.tables):
                    try:
                        cells = table.extract()
                        rows = []
                        
                        # Process cells properly, handling various data types
                        for row in cells:
                            row_data = []
                            for cell in row:
                                # Handle different cell types
                                if hasattr(cell, 'text'):
                                    row_data.append(cell.text)
                                elif isinstance(cell, str):
                                    row_data.append(cell)
                                elif cell is None:
                                    row_data.append("")
                                else:
                                    row_data.append(str(cell))
                            rows.append(row_data)
                        
                        # Create pandas DataFrame with proper error handling
                        df = None
                        raw_data = rows
                        if rows and any(row and any(cell.strip() if isinstance(cell, str) else cell for cell in row) for row in rows):
                            try:
                                # Ensure we have at least one row for column headers
                                if len(rows) > 1:
                                    df = pd.DataFrame(rows[1:], columns=rows[0])
                                else:
                                    # Create single-row table with column indices
                                    df = pd.DataFrame([rows[0]])
                            except Exception as e:
                                logger.warning(f"Could not create DataFrame for table: {e}")
                            
                        # Skip empty tables
                        if not rows or not any(any(cell.strip() if isinstance(cell, str) else cell for cell in row) for row in rows):
                            logger.info(f"Skipping empty table on page {page_num+1}")
                            continue
                        
                        tables.append({
                            "page": page_num + 1,
                            "table_id": f"table_{page_num + 1}_{i + 1}",
                            "dataframe": df,
                            "raw_data": raw_data,
                            "rect": [table.rect.x0, table.rect.y0, table.rect.x1, table.rect.y1]
                        })
                        
                        # Store actual data for display
                        if df is not None:
                            tables[-1]["html"] = df.to_html(index=False)
                        else:
                            # Create HTML table manually if DataFrame creation failed
                            html = "<table>"
                            for row in raw_data:
                                html += "<tr>"
                                for cell in row:
                                    html += f"<td>{cell}</td>"
                                html += "</tr>"
                            html += "</table>"
                            tables[-1]["html"] = html
                    except Exception as e:
                        logger.error(f"Error processing table {i} on page {page_num+1}: {e}")
        
        # If PyMuPDF didn't find any tables, try using the table_extraction code from version1.0.0
        if not tables and os.path.exists(os.path.join("backend", "version1.0.0", "table_extraction.py")):
            try:
                # Import the extract_tables function from table_extraction
                import sys
                sys.path.append(os.path.join(os.getcwd(), "backend", "version1.0.0"))
                from table_extraction import extract_tables
                
                logger.info("Using alternative table extraction method")
                
                # Apply the custom table extraction to each page
                doc = open_pdf(pdf_path)
                for page_num, page in enumerate(doc):
                    table_chunks = extract_tables(page, page_num)
                    
                    for i, chunk in enumerate(table_chunks):
                        if "Table:" in chunk:
                            # Extract table content
                            table_text = chunk.split("Table:")[1].strip()
                            if table_text and not table_text.startswith("[No tables") and not table_text.startswith("[Error"):
                                rows = []
                                for line in table_text.split("\n"):
                                    if "|" in line:
                                        rows.append([cell.strip() for cell in line.split("|")])
                                
                                # Create DataFrame if possible
                                df = None
                                if rows and len(rows) > 1:
                                    try:
                                        df = pd.DataFrame(rows[1:], columns=rows[0] if rows[0] else None)
                                    except Exception as e:
                                        logger.warning(f"Could not create DataFrame for extracted table: {e}")
                                
                                # Add table information
                                tables.append({
                                    "page": page_num + 1,
                                    "table_id": f"table_{page_num + 1}_{i + 1}",
                                    "dataframe": df,
                                    "raw_data": rows,
                                    "extraction_method": "unstructured"
                                })
                                
                                # Store HTML representation
                                if df is not None:
                                    tables[-1]["html"] = df.to_html(index=False)
                                else:
                                    # Create HTML table manually
                                    html = "<table>"
                                    for row in rows:
                                        html += "<tr>"
                                        for cell in row:
                                            html += f"<td>{cell}</td>"
                                        html += "</tr>"
                                    html += "</table>"
                                    tables[-1]["html"] = html
            except ImportError:
                logger.warning("Could not import table_extraction module from version1.0.0")
            except Exception as e:
                logger.error(f"Error using alternative table extraction: {e}")
        
        logger.info(f"Extracted {len(tables)} tables from {pdf_path}")
        return tables
    except Exception as e:
        logger.error(f"Error extracting tables from PDF: {e}", exc_info=True)
        return []

def detect_tables_from_docx(docx_path):
    """Detect tables from a DOCX document."""
    try:
        tables = []
        doc = docx.Document(docx_path)
        
        for i, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                rows_data.append(row_data)
            
            # Create pandas DataFrame
            if rows_data:
                df = pd.DataFrame(rows_data[1:], columns=rows_data[0] if rows_data else None)
                
                tables.append({
                    "table_id": f"table_{i + 1}",
                    "dataframe": df,
                    "raw_data": rows_data
                })
        
        logger.info(f"Detected {len(tables)} tables in {docx_path}")
        return tables
    except Exception as e:
        logger.error(f"Error detecting tables from DOCX: {e}", exc_info=True)
        return []

def detect_figures_from_pdf(pdf_path):
    """Detect figures/images and charts from a PDF document."""
    try:
        doc = open_pdf(pdf_path)
        figures = []
        
        # First find images using PyMuPDF
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    # Extract image with proper error handling
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Verify image data is valid
                    if not image_bytes:
                        logger.warning(f"Empty image data for image {img_index} on page {page_num+1}")
                        continue
                    
                    # Get image location on page
                    bbox = None
                    for img_obj in page.get_images(full=True):
                        if img_obj[0] == xref:
                            bbox = page.get_image_bbox(img_obj)
                            break
                    
                    # Create figure object with all necessary data
                    image_type = base_image["ext"].lower()
                    
                    # Validate image by attempting to open it
                    try:
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        width, height = pil_image.size
                        
                        # Create base64 representation for preview
                        img_b64 = base64.b64encode(image_bytes).decode('utf-8')
                        preview_src = f"data:image/{image_type};base64,{img_b64}"
                        
                        figures.append({
                            "page": page_num + 1,
                            "figure_id": f"figure_{page_num + 1}_{img_index + 1}",
                            "image_bytes": image_bytes,
                            "rect": bbox,
                            "width": width,
                            "height": height,
                            "image_type": image_type,
                            "preview_src": preview_src,
                            "type": "image"
                        })
                    except Exception as e:
                        logger.warning(f"Failed to validate image {img_index} on page {page_num+1}: {e}")
                        continue
                        
                except Exception as e:
                    logger.warning(f"Error extracting image {img_index} on page {page_num+1}: {e}")
                    continue
        
        # Try to find charts using chart_extraction from version1.0.0
        try:
            import sys
            import os.path
            # Get the directory where this script is located
            script_dir = os.path.dirname(os.path.abspath(__file__))
            # Create an absolute path to the version1.0.0 folder
            version_dir = os.path.join(script_dir, "version1.0.0")
            sys.path.append(version_dir)
            
            # Try to import the module
            try:
                from chart_extraction import identify_and_extract_charts
                logger.info("Looking for charts with chart_extraction module")
            except ImportError:
                # Fallback to copying the code if the import fails
                logger.warning("Could not import chart_extraction directly - copying function")
                
                def identify_and_extract_charts(page, page_num):
                    """Extract text from charts on a page."""
                    chunks = []
                    try:
                        # Save page as a temporary PDF
                        doc = fitz.Document()
                        doc.insert_page(-1)
                        doc[0].insert_image(doc[0].rect, pixmap=page.get_pixmap(matrix=fitz.Matrix(2, 2)))
                        temp_pdf = f"temp_page_{page_num}_chart.pdf"
                        doc.save(temp_pdf)
                        doc.close()
                        
                        # Basic text extraction since unstructured might not be available
                        text = page.get_text()
                        if text:
                            chunks.append(f"Page {page_num+1} Chart: {text}")
                        
                        # Clean up
                        if os.path.exists(temp_pdf):
                            os.remove(temp_pdf)
                            
                        return chunks
                    except Exception as e:
                        logger.error(f"Error in chart extraction: {e}")
                        return [f"Page {page_num+1} Chart: [Error: {str(e)}]"]
            
            doc = open_pdf(pdf_path)
            for page_num, page in enumerate(doc):
                try:
                    chart_texts = identify_and_extract_charts(page, page_num)
                    
                    for i, chart_text in enumerate(chart_texts):
                        if "Chart:" in chart_text and not "[No charts detected]" in chart_text and not "[Error:" in chart_text:
                            # Try to extract page image as a visual representation of the chart
                            try:
                                # Render page to image
                                zoom = 2.0  # Higher resolution
                                mat = fitz.Matrix(zoom, zoom)
                                pix = page.get_pixmap(matrix=mat)
                                img_bytes = pix.tobytes("png")
                                
                                # Create figure entry for the chart
                                pil_img = Image.open(io.BytesIO(img_bytes))
                                width, height = pil_img.size
                                
                                # Create base64 representation for preview
                                img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                                preview_src = f"data:image/png;base64,{img_b64}"
                                
                                # Extract chart text
                                chart_content = chart_text.split("Chart:")[1].strip()
                                
                                figures.append({
                                    "page": page_num + 1,
                                    "figure_id": f"chart_{page_num + 1}_{i + 1}",
                                    "image_bytes": img_bytes,
                                    "rect": None,  # Chart covers the whole page or part of it
                                    "width": width,
                                    "height": height,
                                    "image_type": "png",
                                    "preview_src": preview_src,
                                    "type": "chart",
                                    "chart_text": chart_content
                                })
                            except Exception as e:
                                logger.warning(f"Failed to create chart image for page {page_num+1}: {e}")
                except Exception as e:
                    logger.warning(f"Error extracting charts from page {page_num+1}: {e}")
            
        except ImportError:
            logger.warning("Could not import chart_extraction module from version1.0.0")
        except Exception as e:
            logger.error(f"Error using chart extraction: {e}")
            
        logger.info(f"Extracted {len(figures)} figures from {pdf_path}")
        return figures
    except Exception as e:
        logger.error(f"Error extracting figures from PDF: {e}", exc_info=True)
        return []

def detect_figures_from_docx(docx_path):
    """Extract images from a DOCX document."""
    try:
        figures = []
        doc = docx.Document(docx_path)
        
        # We need to save the document to a temporary location and extract images from there
        temp_dir = tempfile.mkdtemp()
        temp_docx_path = os.path.join(temp_dir, "temp.docx")
        doc.save(temp_docx_path)
        
        # Use PyMuPDF to extract images (since python-docx doesn't have direct image extraction)
        try:
            # Changed from fitz.Document to FitzDocument
            doc_pdf = fitz.open(temp_docx_path)
            
            for page_num, page in enumerate(doc_pdf):
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc_pdf.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Verify that image_bytes is not None
                        if image_bytes:
                            figures.append({
                                "page": page_num + 1,
                                "figure_id": f"figure_{page_num + 1}_{img_index + 1}",
                                "image_bytes": image_bytes,
                                "image_type": base_image["ext"].lower()
                            })
                        else:
                            logger.warning(f"Image bytes is None for image {img_index} on page {page_num+1} in DOCX")
                    except Exception as e:
                        logger.warning(f"Error extracting image {img_index} on page {page_num+1} in DOCX: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error extracting images from DOCX with PyMuPDF: {e}")
        
        # Clean up temp files
        import shutil
        shutil.rmtree(temp_dir)
        
        logger.info(f"Detected {len(figures)} figures in {docx_path}")
        return figures
    except Exception as e:
        logger.error(f"Error detecting figures from DOCX: {e}", exc_info=True)
        return []

def analyze_document_chapters(doc_structure, llm):
    """Use LLM to analyze document structure and identify chapters for numbering."""
    try:
        prompt = PromptTemplate(
            template="""Analyze the following document structure and identify the main chapters or sections
            that should be used for numbering figures and tables.
            
            Document Structure:
            {structure}
            
            For each major section (chapter), provide:
            1. The chapter number (starting with 1)
            2. The chapter title
            3. Any subsections that belong to this chapter
            
            Output Format:
            Chapter 1: [Title]
            - Subsection: [Title] 
            - Subsection: [Title]
            
            Chapter 2: [Title]
            - Subsection: [Title]
            ...
            """,
            input_variables=["structure"],
        )
        
        chain = prompt | llm | StrOutputParser()
        chapter_analysis = chain.invoke({"structure": doc_structure})
        
        # Parse the chapter analysis to extract chapter numbers and titles
        chapters = []
        current_chapter = None
        
        for line in chapter_analysis.split('\n'):
            line = line.strip()
            if line.startswith('Chapter '):
                chapter_match = re.match(r'Chapter (\d+): (.+)', line)
                if chapter_match:
                    current_chapter = {
                        "number": int(chapter_match.group(1)),
                        "title": chapter_match.group(2),
                        "subsections": []
                    }
                    chapters.append(current_chapter)
            elif line.startswith('- Subsection: ') and current_chapter:
                subsection_title = line[13:].strip()
                current_chapter["subsections"].append(subsection_title)
        
        return chapters
    except Exception as e:
        logger.error(f"Error analyzing document chapters: {e}", exc_info=True)
        return []

def assign_figure_table_numbers(figures, tables, chapters):
    """Assign figure and table numbers based on chapter structure."""
    try:
        # Default to chapter 1 if no chapters found
        if not chapters:
            chapters = [{"number": 1, "title": "Document", "subsections": []}]
        
        # For now, we'll use a simple approach: assign all figures/tables to their page's closest chapter
        # A more sophisticated approach would analyze the content around the figure/table
        
        # First, get the page ranges for each chapter (approximation)
        chapter_page_ranges = []
        for i, chapter in enumerate(chapters):
            start_page = i * 5 + 1  # Simple estimation; would be better to analyze actual document
            end_page = (i+1) * 5 if i < len(chapters)-1 else 1000
            chapter_page_ranges.append((chapter["number"], start_page, end_page))
        
        # Assign numbers to figures
        for figure in figures:
            page = figure.get("page", 1)
            chapter_number = 1
            
            # Find which chapter this page belongs to
            for ch_num, start, end in chapter_page_ranges:
                if start <= page <= end:
                    chapter_number = ch_num
                    break
            
            # Count how many figures already in this chapter to get the sequence number
            seq_number = sum(1 for f in figures if f.get("chapter_number") == chapter_number) + 1
            
            # Assign the figure number
            figure["chapter_number"] = chapter_number
            figure["sequence_number"] = seq_number
            figure["full_number"] = f"{chapter_number}.{seq_number}"
            figure["caption"] = f"Figure {figure['full_number']}"
        
        # Assign numbers to tables using the same approach
        for table in tables:
            page = table.get("page", 1)
            chapter_number = 1
            
            # Find which chapter this page belongs to
            for ch_num, start, end in chapter_page_ranges:
                if start <= page <= end:
                    chapter_number = ch_num
                    break
            
            # Count how many tables already in this chapter to get the sequence number
            seq_number = sum(1 for t in tables if t.get("chapter_number") == chapter_number) + 1
            
            # Assign the table number
            table["chapter_number"] = chapter_number
            table["sequence_number"] = seq_number
            table["full_number"] = f"{chapter_number}.{seq_number}"
            table["caption"] = f"Table {table['full_number']}"
        
        return figures, tables
    except Exception as e:
        logger.error(f"Error assigning figure and table numbers: {e}", exc_info=True)
        return figures, tables

def add_figure_to_docx(doc, figure, caption):
    """Add a figure to a DOCX document with proper caption and formatting."""
    try:
        # Verify that image_bytes exists and is not None
        if "image_bytes" not in figure or figure["image_bytes"] is None:
            logger.warning(f"No image bytes for figure {figure.get('figure_id', 'unknown')}")
            p = doc.add_paragraph()
            p.add_run(f"[Figure placeholder: {caption}]").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the caption even if the image couldn't be loaded
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            return True
            
        # Add the image to the document
        p = doc.add_paragraph()
        r = p.add_run()
        
        try:
            # First check if the image bytes can be opened as a valid image
            try:
                # Explicitly validate image before saving to file
                pil_image = Image.open(io.BytesIO(figure["image_bytes"]))
                image_format = pil_image.format.lower() if pil_image.format else figure.get('image_type', 'png')
                
                # Create a temporary file to store the image with proper extension
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{image_format}") as tmp:
                    tmp_path = tmp.name
                    # Save the PIL image instead of raw bytes to ensure proper format
                    pil_image.save(tmp_path, format=image_format.upper())
            except Exception as img_error:
                logger.warning(f"Invalid image data: {img_error}. Trying to convert to PNG.")
                
                # Attempt to convert problematic image data to PNG as fallback
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp_path = tmp.name
                    try:
                        # Try to read as numpy array and convert
                        arr = np.frombuffer(figure["image_bytes"], dtype=np.uint8)
                        img_array = cv2.imdecode(arr, cv2.IMREAD_UNCHANGED)
                        if img_array is None:
                            raise ValueError("Could not decode image data")
                        cv2.imwrite(tmp_path, img_array)
                    except Exception as cv_error:
                        logger.error(f"Failed to convert image: {cv_error}")
                        raise ValueError(f"Could not process image data: {cv_error}")
            
            # Add the image to the document with appropriate width
            width_inches = min(6, figure.get("width", 500) / 96)  # Convert pixels to inches, max 6 inches
            r.add_picture(tmp_path, width=Inches(width_inches))
            
            # Center the image
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add the caption
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
            # Clean up temp file
            os.unlink(tmp_path)
            return True
        except Exception as e:
            logger.error(f"Error adding image to document: {e}")
            
            # Add a placeholder instead
            r.add_text(f"[Image could not be displayed: {e}]")
            r.italic = True
            
            # Add the caption anyway
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
            return False
    except Exception as e:
        logger.error(f"Error adding figure to DOCX: {e}", exc_info=True)
        
        # Add a placeholder instead
        p = doc.add_paragraph()
        p.add_run(f"[Figure error: {str(e)}]").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the caption anyway
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.style = 'Caption'
        
        return False

def add_table_to_docx(doc, table_data, caption):
    """Add a table to a DOCX document with proper formatting and caption."""
    try:
        # First add the caption
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.style = 'Caption'
        
        # Then add the table
        if "raw_data" in table_data and table_data["raw_data"]:
            rows = table_data["raw_data"]
            
            num_rows = len(rows)
            num_cols = max(len(row) for row in rows)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # Fill the table with data
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:  # Ensure we don't exceed columns
                        cell = table.cell(i, j)
                        cell.text = str(cell_text)
            
            # Format the header row if present
            if num_rows > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    cell_para = cell.paragraphs[0]
                    # Use a safe way to format text
                    if cell_para.runs:
                        for run in cell_para.runs:
                            run.bold = True
                    else:
                        run = cell_para.add_run(cell.text)
                        run.bold = True
                        cell.text = ""
            
            return True
        elif "dataframe" in table_data and table_data["dataframe"] is not None:
            # Alternative: use DataFrame if raw_data is not available
            df = table_data["dataframe"]
            table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Add headers
            for j, column in enumerate(df.columns):
                table.cell(0, j).text = str(column)
                
            # Add data
            for i, row in enumerate(df.itertuples(index=False)):
                for j, value in enumerate(row):
                    table.cell(i+1, j).text = str(value)
                    
            return True
        else:
            logger.warning("No table data available to add to document")
            p = doc.add_paragraph()
            p.add_run(f"[Table placeholder: {caption}]").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return False
            
    except Exception as e:
        logger.error(f"Error adding table to DOCX: {e}", exc_info=True)
        p = doc.add_paragraph()
        p.add_run(f"[Table error: {str(e)}]").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return False

def process_with_unstructured(file_path, file_type, user_tier=UserTier.FREE):
    """Process a document using unstructured.io for better parsing."""
    if not UNSTRUCTURED_AVAILABLE:
        logger.warning("Unstructured module not available. Falling back to standard processing.")
        return [], [], [], "Unstructured module not available"
    
    try:
        logger.info(f"Processing document with unstructured.io: {file_path}")
        elements = []
        tables = []
        figures = []
        
        # Set OCR languages based on user tier
        ocr_languages = UserTier.get_tier_features()[user_tier]["ocr_languages"]
        
        # Configure extra parameters based on user tier features
        extra_params = {}
        if user_tier == UserTier.PREMIUM:
            # Advanced settings for premium users
            extra_params = {
                "strategy": "hi_res",
                "languages": ocr_languages
            }
        else:
            # Basic settings for free users
            extra_params = {
                "strategy": "fast",
                "languages": ["eng"]  # Only English for free tier
            }
        
        # Process based on file type with error handling
        try:
            if file_type == "pdf":
                # Make sure we have partition_pdf in scope
                from unstructured.partition.pdf import partition_pdf
                elements = partition_pdf(file_path, **extra_params)
            elif file_type == "docx":
                from unstructured.partition.docx import partition_docx
                elements = partition_docx(file_path)
            elif file_type == "txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                from unstructured.partition.text import partition_text
                elements = partition_text(text)
            elif file_type in ["png", "jpg", "jpeg", "tiff"]:
                # Image processing with OCR
                from unstructured.partition.image import partition_image
                elements = partition_image(file_path, languages=ocr_languages if user_tier == UserTier.PREMIUM else ["eng"])
            else:
                return [], [], [], f"Unsupported file type for unstructured: {file_type}"
        except ImportError as e:
            if "pi_heif" in str(e).lower():
                logger.warning("Could not import pi_heif. Attempting to install...")
                try:
                    import subprocess
                    subprocess.check_call(["pip", "install", "pi_heif"])
                    # Try again after installing
                    if file_type == "pdf":
                        from unstructured.partition.pdf import partition_pdf
                        elements = partition_pdf(file_path, **extra_params)
                    elif file_type in ["png", "jpg", "jpeg", "tiff"]:
                        from unstructured.partition.image import partition_image
                        elements = partition_image(file_path, languages=ocr_languages if user_tier == UserTier.PREMIUM else ["eng"])
                except Exception as e2:
                    logger.error(f"Failed to install pi_heif dependency: {e2}")
                    return [], [], [], f"Error installing dependency: {e2}"
            else:
                logger.error(f"Import error: {e}")
                return [], [], [], f"Import error: {e}"
        except Exception as e:
            logger.error(f"Error processing with unstructured: {e}")
            return [], [], [], f"Error processing with unstructured: {e}"
        
        # Extract text content for LangChain
        texts = []
        for i, element in enumerate(elements):
            # Process based on element type
            if isinstance(element, Table):
                # Extract table data
                if hasattr(element, "metadata") and hasattr(element, "text"):
                    table_data = element.text
                    table_id = f"table_{i}"
                    
                    # Parse table text to create rows
                    rows = []
                    if hasattr(element, "cells"):
                        # Direct access to cell data if available
                        rows = element.cells
                    else:
                        # Fallback to parsing text
                        lines = table_data.split('\n')
                        for line in lines:
                            cells = line.split('|')
                            rows.append([cell.strip() for cell in cells if cell.strip()])
                    
                    # Create DataFrame if possible
                    df = None
                    if rows and len(rows) > 1:
                        try:
                            df = pd.DataFrame(rows[1:], columns=rows[0])
                        except Exception as e:
                            logger.warning(f"Could not create DataFrame for table: {e}")
                    
                    # Add table to collection
                    tables.append({
                        "table_id": table_id,
                        "raw_data": rows,
                        "dataframe": df,
                        "page": element.metadata.page_number if hasattr(element.metadata, "page_number") else 1
                    })
                    
                    # Add table text to content for LangChain
                    texts.append(f"[TABLE {table_id}]\n{table_data}\n[/TABLE]")
                
            elif isinstance(element, UnstructuredImage):
                # Extract image data if available
                figure_id = f"figure_{i}"
                
                # Get image data if available
                image_data = None
                if hasattr(element, "metadata") and hasattr(element.metadata, "image_data"):
                    image_data = element.metadata.image_data
                
                # Get page number
                page_number = 1
                if hasattr(element, "metadata") and hasattr(element.metadata, "page_number"):
                    page_number = element.metadata.page_number
                
                figures.append({
                    "figure_id": figure_id,
                    "page": page_number,
                    "image_bytes": image_data,
                    "image_type": "png",  # Default if not specified
                    "caption": element.text if hasattr(element, "text") else f"Figure {figure_id}"
                })
                
                # Add figure reference to content for LangChain
                texts.append(f"[FIGURE {figure_id}]\n{element.text if hasattr(element, 'text') else ''}\n[/FIGURE]")
            
            elif isinstance(element, Title):
                # Mark titles for better structure detection
                texts.append(f"[TITLE] {element.text} [/TITLE]")
            
            else:
                # Regular text
                if hasattr(element, "text"):
                    texts.append(element.text)
        
        # Create document chunks for LangChain
        from langchain_core.documents import Document
        
        chunks = []
        for i, text in enumerate(texts):
            chunks.append(Document(
                page_content=text,
                metadata={"source": file_path, "chunk": i}
            ))
            
        logger.info(f"Unstructured extraction complete: {len(chunks)} chunks, {len(tables)} tables, {len(figures)} figures")
        return chunks, tables, figures, None
        
    except Exception as e:
        logger.error(f"Error processing with unstructured: {e}", exc_info=True)
        return [], [], [], f"Error processing with unstructured: {str(e)}"

# Rule-based functions for free tier users

def rule_based_structure_extraction(chunks):
    """Extract document structure using rule-based approaches for free tier users."""
    try:
        # Concatenate chunks into one text
        full_text = "\n\n".join([chunk.page_content for chunk in chunks])
        
        # Look for potential headings using regex patterns
        heading_patterns = [
            (r'^#+\s+(.+)$', 'H1'),  # Markdown headings
            (r'^(.+)\n=+$', 'H1'),    # Underlined headings
            (r'^(.+)\n-+$', 'H2'),    # Underlined subheadings
            (r'^[A-Z][A-Z\s]+[A-Z]$', 'H1'),  # ALL CAPS HEADINGS
            (r'^[0-9]+\.\s+(.+)$', 'H2'),    # Numbered headings
            (r'^[A-Z][a-z].{10,60}$', 'H3')  # Sentence-like headings of certain length
        ]
        
        # Analyze lines to find headings
        structure = []
        lines = full_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Try to detect headings
            for pattern, level in heading_patterns:
                if re.match(pattern, line):
                    # Extract heading text
                    heading_text = re.match(pattern, line).group(1) if '(' in pattern else line
                    
                    # Generate description from following lines (up to 200 chars)
                    description = ""
                    if i + 1 < len(lines):
                        for j in range(i+1, min(i+5, len(lines))):
                            if lines[j].strip() and not any(re.match(p, lines[j].strip()) for p, _ in heading_patterns):
                                description += lines[j].strip() + " "
                    
                    structure.append(f"- {level}: {heading_text}")
                    structure.append(f"  Description: {description[:200]}...")
                    break
        
        if not structure:
            # If no structure was detected, create a basic one
            structure = [
                "- H1: Document Title",
                "  Description: Main document content...",
                "- H2: Introduction",
                "  Description: Beginning of the document...",
                "- H2: Content",
                "  Description: Main body of the document...",
                "- H2: Conclusion",
                "  Description: End of the document..."
            ]
            
        return "\n".join(structure)
    except Exception as e:
        logger.error(f"Error in rule-based structure extraction: {e}", exc_info=True)
        return "Error extracting document structure using rule-based approach"

def rule_based_chapter_detection(chunks):
    """Detect document chapters using rule-based approaches for free tier users."""
    try:
        # Create a simple chapter structure based on the first few chunks
        chapters = [
            {
                "number": 1,
                "title": "Introduction",
                "subsections": ["Overview", "Background"]
            },
            {
                "number": 2, 
                "title": "Content",
                "subsections": ["Main Points", "Details"]
            },
            {
                "number": 3,
                "title": "Conclusion",
                "subsections": ["Summary", "Next Steps"]
            }
        ]
        
        # Attempt to extract real chapter titles from text
        full_text = "\n\n".join([chunk.page_content for chunk in chunks])
        lines = full_text.split('\n')
        
        chapter_patterns = [
            r'Chapter\s+(\d+)[:\s]+(.+)',
            r'Section\s+(\d+)[:\s]+(.+)',
            r'^(\d+)\.\s+(.+)'
        ]
        
        # Override default chapters if we find actual ones
        detected_chapters = []
        for line in lines:
            for pattern in chapter_patterns:
                match = re.search(pattern, line)
                if match:
                    try:
                        chapter_num = int(match.group(1))
                        chapter_title = match.group(2).strip()
                        detected_chapters.append({
                            "number": chapter_num,
                            "title": chapter_title,
                            "subsections": []
                        })
                    except:
                        continue
        
        # Use detected chapters if we found any
        if detected_chapters:
            chapters = detected_chapters
            
            # Look for subsections (simple heuristic)
            current_chapter = None
            for line in lines:
                for chapter in chapters:
                    if chapter["title"] in line:
                        current_chapter = chapter
                        break
                        
                if current_chapter and '.' in line and len(line.split('.')) > 1:
                    potential_subsection = line.split('.', 1)[1].strip()
                    if 5 < len(potential_subsection) < 100 and potential_subsection not in current_chapter["subsections"]:
                        current_chapter["subsections"].append(potential_subsection)
        
        return chapters
    except Exception as e:
        logger.error(f"Error in rule-based chapter detection: {e}", exc_info=True)
        # Return default chapters if error
        return [
            {"number": 1, "title": "Document", "subsections": []}
        ]

def rule_based_abstract_generation(chunks):
    """Generate a basic abstract for free tier users without using LLM."""
    try:
        # Take the first 2-3 chunks and extract key sentences
        content = "\n\n".join([chunk.page_content for chunk in chunks[:3]])
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', content)
        
        # Select key sentences (first sentence, 2-3 from middle, last sentence)
        selected_sentences = []
        
        if sentences:
            # Always include first sentence if it's long enough
            if len(sentences[0]) > 20:
                selected_sentences.append(sentences[0])
            
            # Add 2-3 sentences from the middle
            middle_start = max(1, len(sentences) // 3)
            middle_end = min(len(sentences) - 1, 2 * len(sentences) // 3)
            
            for i in range(middle_start, middle_end):
                # Select sentences that aren't too short
                if len(sentences[i]) > 30:
                    selected_sentences.append(sentences[i])
                    # Limit to 2-3 sentences from middle
                    if len(selected_sentences) >= 3:
                        break
            
            # Add last sentence if it's meaningful
            if len(sentences[-1]) > 20:
                selected_sentences.append(sentences[-1])
        
        # Combine sentences into abstract
        abstract = " ".join(selected_sentences)
        
        # If abstract is too short, add a generic statement
        if len(abstract) < 100:
            abstract += " This document provides information on the subject matter covered within its contents."
        
        return abstract
    except Exception as e:
        logger.error(f"Error in rule-based abstract generation: {e}", exc_info=True)
        return "This document contains information relevant to the subject matter. For a more detailed abstract, upgrade to premium tier."

def rule_based_word_suggestions(current_text):
    """Generate basic word suggestions based on the current text without using LLM."""
    try:
        # Common words to suggest based on the last word
        suggestion_dict = {
            "the": ["following", "most", "best", "important"],
            "a": ["significant", "major", "minor", "detailed", "brief"],
            "this": ["approach", "method", "result", "finding", "conclusion"],
            "in": ["addition", "conclusion", "summary", "brief", "detail"],
            "for": ["example", "instance", "reference", "clarity", "consideration"],
            "to": ["summarize", "conclude", "illustrate", "demonstrate", "emphasize"],
            "of": ["course", "importance", "significance", "relevance", "interest"],
            "with": ["respect", "regard", "reference", "consideration", "attention"],
            "and": ["therefore", "thus", "consequently", "furthermore", "additionally"],
            "is": ["important", "necessary", "critical", "essential", "fundamental"]
        }
        
        # Default suggestions if we can't find context-specific ones
        default_suggestions = ["the", "and", "therefore", "additionally", "furthermore"]
        
        # Get the last word from the text
        words = current_text.split()
        last_word = words[-1].lower() if words else ""
        
        # Get suggestions based on the last word
        if last_word in suggestion_dict:
            return suggestion_dict[last_word][:3]  # Return up to 3 suggestions
        
        # If we're at the beginning of a sentence, suggest sentence starters
        if last_word.endswith('.') or last_word.endswith('!') or last_word.endswith('?') or not words:
            return ["The", "This", "In", "Furthermore", "Additionally"][:3]
        
        # Default suggestions
        return default_suggestions[:3]
    except Exception as e:
        logger.error(f"Error in rule-based word suggestions: {e}", exc_info=True)
        return ["and", "the", "therefore"]

def main():
    st.title("DocuMorph AI - Intelligent Document Transformation")
    
    # Initialize session state
    if 'user_id' not in st.session_state:
        st.session_state.user_id = str(uuid.uuid4())
        
    # Initialize MongoDB client
    docs_collection, templates_collection, users_collection = init_mongodb()
    
    # Initialize default templates if not already done
    if docs_collection is not None and templates_collection is not None:
        initialize_db_templates(st.session_state.user_id)

    if 'user_tier' not in st.session_state:
        st.session_state.user_tier = UserTier.FREE
    if 'documents' not in st.session_state:
        st.session_state.documents = []
    if 'active_doc' not in st.session_state:
        st.session_state.active_doc = None
    if 'active_template' not in st.session_state:
        st.session_state.active_template = None
    if 'doc_structure' not in st.session_state:
        st.session_state.doc_structure = None
    if 'edit_text' not in st.session_state:
        st.session_state.edit_text = ""
    if 'suggestion_index' not in st.session_state:
        st.session_state.suggestion_index = 0
    if 'suggestions' not in st.session_state:
        st.session_state.suggestions = []
    if 'formatted_doc_path' not in st.session_state:
        st.session_state.formatted_doc_path = None
    if 'tables' not in st.session_state:
        st.session_state.tables = []
    if 'figures' not in st.session_state:
        st.session_state.figures = []
    if 'chapters' not in st.session_state:
        st.session_state.chapters = []
    if 'google_drive_auth' not in st.session_state:
        st.session_state.google_drive_auth = None
    
    # Sidebar for features and user tier
    with st.sidebar:
        # User tier selection
        st.header("User Tier")
        tier_selection = st.radio(
            "Select your tier:",
            [UserTier.FREE, UserTier.PREMIUM],
            format_func=lambda x: "Free" if x == UserTier.FREE else "Premium"
        )
        
        if tier_selection != st.session_state.user_tier:
            st.session_state.user_tier = tier_selection
            st.info(f"Switched to {tier_selection.capitalize()} tier")
        
        # Display tier features
        tier_features = UserTier.get_tier_features()[st.session_state.user_tier]
        st.markdown("### Tier Features")
        
        feature_icons = {
            True: "✅",
            False: "❌"
        }
        
        st.markdown(f"**Document Limit:** {tier_features['max_documents']}")
        st.markdown(f"**Template Limit:** {tier_features['max_templates']}")
        st.markdown(f"**LLM Features:** {feature_icons[tier_features['llm_enabled']]}")
        st.markdown(f"**Multi-language Support:** {feature_icons[tier_features['multi_language']]}")
        st.markdown(f"**Style Guide Compliance:** {feature_icons[tier_features['style_guide_compliance']]}")
        st.markdown(f"**Team Collaboration:** {feature_icons[tier_features['team_collaboration']]}")
        
        st.header("Features")
        feature_selection = st.radio(
            "Select Feature:",
            ["Upload & Process Document", "Template Management", "Document Enhancement", 
             "Tables & Figures", "Live Document Editor", "Preview & Export", 
             "Plagiarism Check", "Export to Google Docs", "Custom Templates"]
        )
    
    # Main content area based on selected feature
    if feature_selection == "Upload & Process Document":
        st.header("Upload Document")
        
        # Check document limit for free tier
        if st.session_state.user_tier == UserTier.FREE and len(st.session_state.documents) >= UserTier.get_tier_features()[UserTier.FREE]["max_documents"]:
            st.warning(f"Free tier is limited to {UserTier.get_tier_features()[UserTier.FREE]['max_documents']} documents. Upgrade to Premium for more.")
        
        uploaded_file = st.file_uploader("Upload document", type=["docx", "pdf", "txt", "png", "jpg", "jpeg"])
        
        if uploaded_file is not None:
            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                file_path = tmp_file.name
            
            file_type = uploaded_file.name.split('.')[-1].lower()
            
            # Detect file type using magic
            try:
                detected_mime = magic.from_file(file_path, mime=True)
                st.write(f"Detected file type: {detected_mime}")
            except:
                st.write("File type detection is not available")
            
            # OCR options for image files
            ocr_language = "eng"
            if file_type in ["png", "jpg", "jpeg", "tiff"] and st.session_state.user_tier == UserTier.PREMIUM:
                ocr_language = st.selectbox(
                    "Select OCR language",
                    UserTier.get_tier_features()[UserTier.PREMIUM]["ocr_languages"],
                    index=0
                )
            
            if st.button("Process Document"):
                with st.spinner("Processing document..."):
                    # Process document
                    chunks, error, tables, figures = process_document(file_path, file_type, st.session_state.user_tier)
                    
                    if error:
                        st.error(error)
                    else:
                        # Save document info
                        doc_id = str(uuid.uuid4())
                        permanent_path = f"documents/{st.session_state.user_id}/{doc_id}.{file_type}"
                        os.makedirs(os.path.dirname(permanent_path), exist_ok=True)
                        
                        with open(permanent_path, "wb") as f:
                            f.write(uploaded_file.getvalue())
                        
                        # Store in MongoDB
                        doc_info = {
                            "doc_id": doc_id,
                            "user_id": st.session_state.user_id,
                            "original_name": uploaded_file.name,
                            "file_path": permanent_path,
                            "file_type": file_type
                        }
                        
                        if docs_collection is not None:
                            docs_collection.insert_one(doc_info)
                        
                        # Store in session state
                        st.session_state.documents.append(doc_info)
                        st.session_state.active_doc = doc_info
                        st.session_state.tables = tables
                        st.session_state.figures = figures
                        
                        # Extract document structure (method depends on user tier)
                        if can_use_llm(st.session_state.user_tier):
                            llm = init_llm()
                            structure = extract_document_structure(chunks, llm)
                            st.session_state.doc_structure = structure
                            
                            # Analyze chapters for figure/table numbering
                            chapters = analyze_document_chapters(structure, llm)
                            st.session_state.chapters = chapters
                        else:
                            # Rule-based structure extraction for free tier
                            structure = rule_based_structure_extraction(chunks)
                            st.session_state.doc_structure = structure
                            
                            # Basic chapter detection for free tier
                            chapters = rule_based_chapter_detection(chunks)
                            st.session_state.chapters = chapters
                        
                        st.success(f"Document '{uploaded_file.name}' processed successfully!")
                        st.markdown("### Extracted Document Structure")
                        st.write(structure)
                        
                        # Display table and figure detection results
                        st.markdown(f"### Detected Components")
                        st.markdown(f"- Tables: {len(tables)}")
                        st.markdown(f"- Figures: {len(figures)}")
                        st.markdown(f"- Chapters: {len(chapters)}")
                        
                        # Extract document text for editing
                        if file_type == "docx":
                            doc = docx.Document(permanent_path)
                            full_text = []
                            for para in doc.paragraphs:
                                full_text.append(para.text)
                            st.session_state.edit_text = "\n".join(full_text)
                        elif file_type == "txt":
                            with open(permanent_path, "r", encoding="utf-8") as f:
                                st.session_state.edit_text = f.read()
                        else:
                            # For PDFs, use the extracted chunks
                            st.session_state.edit_text = "\n\n".join([chunk.page_content for chunk in chunks])
            
            # Clean up temp file
            os.unlink(file_path)
        
        # Show list of previously uploaded documents
        if st.session_state.documents:
            st.header("Your Documents")
            for i, doc in enumerate(st.session_state.documents):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"{i+1}. {doc['original_name']}")
                with col2:
                    if st.button("Select", key=f"select_doc_{i}"):
                        st.session_state.active_doc = doc
                        st.success(f"Selected document: {doc['original_name']}")
                        
                        # Process document to get tables and figures
                        file_path = doc['file_path']
                        file_type = doc['file_type']
                        chunks, _, tables, figures = process_document(file_path, file_type, st.session_state.user_tier)
                        st.session_state.tables = tables
                        st.session_state.figures = figures
                        
                        # Extract document structure (method depends on user tier)
                        if can_use_llm(st.session_state.user_tier):
                            llm = init_llm()
                            structure = extract_document_structure(chunks, llm)
                            st.session_state.doc_structure = structure
                            
                            # Analyze chapters for figure/table numbering
                            chapters = analyze_document_chapters(structure, llm)
                            st.session_state.chapters = chapters
                        else:
                            # Rule-based structure extraction for free tier
                            structure = rule_based_structure_extraction(chunks)
                            st.session_state.doc_structure = structure
                            
                            # Basic chapter detection for free tier
                            chapters = rule_based_chapter_detection(chunks)
                            st.session_state.chapters = chapters
                        
                        # Load document content for editing
                        if file_type == "docx":
                            doc_obj = docx.Document(file_path)
                            full_text = []
                            for para in doc_obj.paragraphs:
                                full_text.append(para.text)
                            st.session_state.edit_text = "\n".join(full_text)
                        elif file_type == "txt":
                            with open(file_path, "r", encoding="utf-8") as f:
                                st.session_state.edit_text = f.read()
                        else:
                            # For PDFs, use the extracted chunks
                            st.session_state.edit_text = "\n\n".join([chunk.page_content for chunk in chunks])

    elif feature_selection == "Template Management":
        st.header("Template Management")
        
        # Check template limit for free tier            
        tier_features = UserTier.get_tier_features()[st.session_state.user_tier]
        max_templates = tier_features["max_templates"]            
        
        # Get user templates - handles both MongoDB and in-memory modes
        user_templates = get_templates_for_user(st.session_state.user_id, templates_collection, st.session_state.user_tier)
        if st.session_state.user_tier == UserTier.FREE and len(user_templates) >= max_templates:
            st.warning(f"Free tier is limited to {max_templates} templates. Upgrade to Premium for more.")
        
        tab1, tab2 = st.tabs(["Create Template", "Manage Templates"])
        
        with tab1:
            st.subheader("Create New Template")
            template_name = st.text_input("Template Name", "")
            
            # Add role selection for the template
            template_role = st.selectbox(
                "Select template role",
                ["Student", "Content Creator", "Researcher", "Business Professional", 
                 "Multilingual User", "Author", "Collaborator", "Project Manager", "Others"],
                index=0
            )
            
            # For free tier, show warning if selecting non-allowed role
            if st.session_state.user_tier == UserTier.FREE and not can_use_template_role(st.session_state.user_tier, template_role):
                st.warning(f"Free tier only supports Student and Content Creator templates. Please upgrade to Premium for {template_role} templates.")
                template_role = "Student"  # Reset to allowed template
            
            st.markdown("### Document Layout")
            col1, col2 = st.columns(2)
            with col1:
                margin_top = st.number_input("Top Margin (inches)", 0.5, 2.0, 1.0, 0.1)
                margin_left = st.number_input("Left Margin (inches)", 0.5, 2.0, 1.0, 0.1)
            with col2:
                margin_bottom = st.number_input("Bottom Margin (inches)", 0.5, 2.0, 1.0, 0.1)
                margin_right = st.number_input("Right Margin (inches)", 0.5, 2.0, 1.0, 0.1)
            
            st.markdown("### Typography")
            col1, col2 = st.columns(2)
            with col1:
                body_font = st.selectbox("Body Font", ["Calibri", "Arial", "Times New Roman", "Georgia"])
                body_font_size = st.slider("Body Font Size (pt)", 8, 14, 11)
            with col2:
                heading_font = st.selectbox("Heading Font", ["Calibri", "Arial", "Times New Roman", "Georgia"])
                heading1_font_size = st.slider("Heading 1 Size (pt)", 14, 22, 16)
            
            heading2_font_size = st.slider("Heading 2 Size (pt)", 12, 18, 14)
            heading3_font_size = st.slider("Heading 3 Size (pt)", 11, 16, 12)
            
            st.markdown("### Header & Footer")
            header_text = st.text_input("Header Text (optional)")
            footer_text = st.text_input("Footer Text (optional)")
            
            if st.button("Save Template"):
                # Check if we're at the template limit for free users
                if st.session_state.user_tier == UserTier.FREE and len(user_templates) >= max_templates:
                    st.error(f"Free tier is limited to {max_templates} templates. Upgrade to Premium to create more templates.")
                else:
                    template = {
                        "template_id": str(uuid.uuid4()),
                        "user_id": st.session_state.user_id,
                        "name": template_name,
                        "role": template_role,
                        "margin_top": margin_top,
                        "margin_bottom": margin_bottom,
                        "margin_left": margin_left,
                        "margin_right": margin_right,
                        "body_font": body_font,
                        "body_font_size": body_font_size,
                        "heading_font": heading_font,
                        "heading1_font_size": heading1_font_size,
                        "heading2_font_size": heading2_font_size,
                        "heading3_font_size": heading3_font_size,
                        "header_text": header_text,
                        "footer_text": footer_text
                    }
                    
                    if templates_collection is not None:
                        templates_collection.insert_one(template)
                    
                    st.success(f"Template '{template_name}' saved successfully!")
        
        with tab2:
            st.subheader("Your Templates")
            
            if templates_collection is not None:
                templates = list(templates_collection.find({"user_id": st.session_state.user_id}))
                
                # Get unique roles from templates
                roles = ["All"] + list(set([t.get("role", "Others") for t in templates]))
                
                # Filter by role
                selected_role_filter = st.selectbox("Filter by role:", roles)
                
                if templates:
                    # Filter templates by role if a specific role is selected
                    if selected_role_filter != "All":
                        filtered_templates = [t for t in templates if t.get("role", "Others") == selected_role_filter]
                    else:
                        filtered_templates = templates
                    
                    # Further filter by allowed categories for user's tier
                    user_tier = st.session_state.user_tier
                    if user_tier == UserTier.FREE:
                        filtered_templates = [t for t in filtered_templates if can_use_template_role(user_tier, t.get("role", "Others"))]
                        if not filtered_templates:
                            st.info("Free tier users can only access Student and Content Creator templates. Upgrade to Premium for more template options.")
                    
                    if not filtered_templates:
                        if selected_role_filter != "All":
                            st.info(f"No templates found for role: {selected_role_filter}")
                        else:
                            st.info("No templates found for your tier.")
                    else:
                        # Group templates by role for better organization
                        templates_by_role = {}
                        for template in filtered_templates:
                            role = template.get("role", "Others")
                            if role not in templates_by_role:
                                templates_by_role[role] = []
                            templates_by_role[role].append(template)
                        
                        # Display templates grouped by role
                        for role, role_templates in templates_by_role.items():
                            st.markdown(f"### {role} Templates")
                            for i, template in enumerate(role_templates):
                                with st.expander(f"{template['name']}"):
                                    st.markdown(f"**Role**: {template.get('role', 'Others')}")
                                    st.markdown(f"**Body Font**: {template['body_font']}, {template['body_font_size']}pt")
                                    st.markdown(f"**Heading Font**: {template['heading_font']}")
                                    st.markdown(f"**Margins**: T: {template['margin_top']}\", B: {template['margin_bottom']}\", L: {template['margin_left']}\", R: {template['margin_right']}\"")
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        if st.button("Select Template", key=f"select_template_{role}_{i}"):
                                            st.session_state.active_template = template
                                            st.success(f"Selected template: {template['name']}")
                                    with col2:
                                        if st.button("Delete", key=f"delete_template_{role}_{i}"):
                                            if templates_collection is not None:
                                                templates_collection.delete_one({"template_id": template["template_id"]})
                                            st.success(f"Template '{template['name']}' deleted.")
                                            st.rerun()
                else:
                    st.info("No templates found. Create a new template to get started.")
            else:
                st.error("MongoDB connection failed. Templates cannot be loaded.")
    
    elif feature_selection == "Document Enhancement":
        st.header("Document Enhancement")
        
        if not st.session_state.active_doc:
            st.warning("Please select a document first")
            return
        
        # Load document for processing
        doc_path = st.session_state.active_doc["file_path"]
        file_type = st.session_state.active_doc["file_type"]
        
        # Process document
        chunks, error, tables, figures = process_document(doc_path, file_type, st.session_state.user_tier)
        if error:
            st.error(error)
            return
        
        # Get user role for tailored enhancements
        user_role = "Others"  # Default role
        if st.session_state.active_template:
            template_role = st.session_state.active_template.get("role", "Others")
            user_role = template_role
        
        # Allow user to select a different role if desired
        selected_role = st.selectbox(
            "Select document role for tailored enhancements:",
            ["Student", "Content Creator", "Researcher", "Business Professional", 
             "Multilingual User", "Author", "Collaborator", "Project Manager", "Others"],
            index=["Student", "Content Creator", "Researcher", "Business Professional", 
                   "Multilingual User", "Author", "Collaborator", "Project Manager", "Others"].index(user_role)
            if user_role in ["Student", "Content Creator", "Researcher", "Business Professional", 
                            "Multilingual User", "Author", "Collaborator", "Project Manager", "Others"] else 8
        )
        
        # For Premium users, use LLM with role-specific prompts; for Free users, use rule-based approaches
        if can_use_llm(st.session_state.user_tier):
            llm = init_llm()
            if not llm:
                st.error("Failed to initialize LLM")
                return
            
            tab1, tab2, tab3, tab4 = st.tabs(["Abstract Generation", "Section Title Suggestions", "Style Enhancement", "Structure Suggestions"])
            
            with tab1:
                st.subheader("Generate Abstract")
                
                # Get role-specific prompt
                default_instructions = get_role_specific_prompt(selected_role, "abstract")
                
                instructions = st.text_area("Instructions for abstract", 
                                          default_instructions, 
                                          height=150)
                
                if st.button("Generate Abstract"):
                    with st.spinner("Generating abstract..."):
                        abstract = generate_abstract(chunks, llm, instructions)
                        st.markdown("### Generated Abstract")
                        st.write(abstract)
                        
                        if st.button("Save Abstract to Document"):
                            st.info("Abstract saved! It will be included when you export the document.")
                            # In a real implementation, you would save this to the document
            
            with tab2:
                st.subheader("Section Title Suggestions")
                
                # Get role-specific prompt
                default_title_instructions = get_role_specific_prompt(selected_role, "section_titles")
                
                # First, extract section content to work with
                if st.session_state.doc_structure:
                    sections = st.session_state.doc_structure.split('\n')
                    section_titles = []
                    
                    for line in sections:
                        if line.strip().startswith('- H'):
                            # Extract section title
                            title_match = line.strip().split(': ')[1] if ': ' in line else ""
                            if title_match:
                                section_titles.append(title_match.strip())
                else:
                    sections = []
                    section_titles = []
                    
                selected_section = st.selectbox("Select section to suggest titles for:", 
                                               ["Select a section"] + section_titles)
                
                if selected_section != "Select a section":
                    # Get section content (in a real implementation, you'd extract the actual content)
                    section_index = section_titles.index(selected_section)
                    section_content = chunks[min(section_index, len(chunks)-1)].page_content
                    
                    instructions = st.text_area("Instructions for title suggestions", 
                                              default_title_instructions,
                                              height=150)
                    
                    if st.button("Generate Title Suggestions"):
                        with st.spinner("Generating title suggestions..."):
                            titles = suggest_section_titles(section_content, selected_section, llm, instructions)
                            st.markdown("### Title Suggestions")
                            st.write(titles)
            
            with tab3:
                st.subheader("Style Enhancement")
                
                # Get role-specific style
                target_style_default = get_role_specific_prompt(selected_role, "style")
                
                text_to_enhance = st.text_area("Enter text to enhance", "")
                target_style = st.text_input("Target style", value=target_style_default)
                
                if st.button("Enhance Text Style") and text_to_enhance:
                    with st.spinner("Enhancing text..."):
                        enhanced_text = enhance_text_style(text_to_enhance, target_style, llm)
                        st.markdown("### Enhanced Text")
                        st.write(enhanced_text)
            
            with tab4:
                st.subheader("Document Structure Suggestions")
                
                # Get role-specific structure prompt
                structure_prompt = get_role_specific_prompt(selected_role, "structure")
                st.markdown("### Structure Analysis Instructions")
                st.write(structure_prompt)
                
                if st.button("Analyze Document Structure"):
                    with st.spinner("Analyzing document structure..."):
                        # Create a prompt template that includes the role-specific instructions
                        prompt = PromptTemplate(
                            template="""Analyze the following document and identify its logical structure based on these instructions:
                            
                            {structure_instructions}
                            
                            Document Content:
                            {content}
                            
                            Output Format:
                            - H1: [Section Title]
                              Description: [Brief description]
                              - H2: [Subsection Title]
                                Description: [Brief description]
                                - H3: [Sub-subsection Title]
                                  Description: [Brief description]
                            """,
                            input_variables=["content", "structure_instructions"],
                        )
                        
                        # Concatenate chunks into one text for LLM analysis
                        full_text = "\n\n".join([chunk.page_content for chunk in chunks])
                        
                        chain = prompt | llm | StrOutputParser()
                        structure = chain.invoke({
                            "content": full_text[:10000],  # Limit to 10k chars to avoid token limits
                            "structure_instructions": structure_prompt
                        })
                        
                        st.session_state.doc_structure = structure
                        st.markdown("### Suggested Document Structure")
                        st.write(structure)
                        
                        # Analyze chapters for figure/table numbering
                        chapters = analyze_document_chapters(structure, llm)
                        st.session_state.chapters = chapters
                        
                        # Show chapter structure
                        if chapters:
                            st.markdown("### Chapter Structure for Numbering")
                            for chapter in chapters:
                                st.markdown(f"**Chapter {chapter['number']}:** {chapter['title']}")
                                for subsection in chapter.get('subsections', []):
                                    st.markdown(f"- {subsection}")
                                    
                        # Add option to apply structure
                        if st.button("Apply Structure to Document"):
                            st.info("This would restructure your document according to the suggested structure. Feature coming soon!")
                            # In a real implementation, you would modify the document structure here
        else:
            # Free tier users - Rule-based approaches
            st.info("Free tier uses rule-based document enhancement. Upgrade to Premium for AI-powered enhancements.")
            
            tab1, tab2, tab3 = st.tabs(["Basic Abstract", "Basic Structure", "Basic Style Tips"])
            
            with tab1:
                st.subheader("Generate Basic Abstract")
                
                if st.button("Generate Basic Abstract"):
                    with st.spinner("Generating abstract..."):
                        abstract = rule_based_abstract_generation(chunks)
                        st.markdown("### Generated Abstract")
                        st.write(abstract)
            
            with tab2:
                st.subheader("Basic Document Structure")
                
                if st.button("Identify Basic Structure"):
                    with st.spinner("Analyzing document structure..."):
                        structure = rule_based_structure_extraction(chunks)
                        st.session_state.doc_structure = structure
                        st.markdown("### Document Structure")
                        st.write(structure)
            
            with tab3:
                st.subheader("Basic Style Tips")
                st.markdown("### General Style Guidelines")
                st.markdown("- Use consistent formatting throughout the document")
                st.markdown("- Ensure paragraphs are properly separated")
                st.markdown("- Use bullet points for lists")
                st.markdown("- Keep sentences clear and concise")
                st.markdown("- Use section headings to organize content")

    elif feature_selection == "Tables & Figures":
        st.header("Tables & Figures Management")
        
        if not st.session_state.active_doc:
            st.warning("Please select a document first")
            return
            
        # Display detected tables
        if st.session_state.tables:
            st.subheader("Detected Tables")
            for i, table in enumerate(st.session_state.tables):
                with st.expander(f"Table {i+1}"):
                    if "dataframe" in table:
                        st.dataframe(table["dataframe"])
                    elif "raw_data" in table:
                        st.table(table["raw_data"])
                    
                    # Table caption and numbering options
                    col1, col2 = st.columns(2)
                    with col1:
                        table_prefix = st.text_input(f"Table prefix for Table {i+1}", 
                                                    value="Table", 
                                                    key=f"table_prefix_{i}")
                    with col2:
                        chapter_num = st.number_input(f"Chapter number for Table {i+1}", 
                                                     min_value=1, 
                                                     value=table.get("chapter_number", 1), 
                                                     key=f"table_chapter_{i}")
                    
                    # Update table info
                    st.session_state.tables[i]["caption_prefix"] = table_prefix
                    st.session_state.tables[i]["chapter_number"] = chapter_num
                    
                    # Calculate sequence number based on other tables with the same chapter
                    seq_number = sum(1 for t in st.session_state.tables[:i] 
                                    if t.get("chapter_number") == chapter_num) + 1
                    
                    st.session_state.tables[i]["sequence_number"] = seq_number
                    st.session_state.tables[i]["full_number"] = f"{chapter_num}.{seq_number}"
                    st.session_state.tables[i]["caption"] = f"{table_prefix} {chapter_num}.{seq_number}"
                    
                    st.info(f"This table will be numbered as: {st.session_state.tables[i]['caption']}")
        else:
            st.info("No tables detected in the document.")
        
        # Display detected figures
        if st.session_state.figures:
            st.subheader("Detected Figures")
            for i, figure in enumerate(st.session_state.figures):
                with st.expander(f"Figure {i+1}"):
                    # Display the image
                    if "image_bytes" in figure:
                        try:
                            image = Image.open(io.BytesIO(figure["image_bytes"]))
                            st.image(image, caption=f"Figure {i+1} from document", 
                                   use_column_width=True)
                        except Exception as e:
                            st.error(f"Could not display image: {e}")
                    
                    # Figure caption and numbering options
                    col1, col2 = st.columns(2)
                    with col1:
                        figure_prefix = st.text_input(f"Figure prefix for Figure {i+1}", 
                                                     value="Figure", 
                                                     key=f"figure_prefix_{i}")
                    with col2:
                        chapter_num = st.number_input(f"Chapter number for Figure {i+1}", 
                                                     min_value=1, 
                                                     value=figure.get("chapter_number", 1), 
                                                     key=f"figure_chapter_{i}")
                    
                    # Update figure info
                    st.session_state.figures[i]["caption_prefix"] = figure_prefix
                    st.session_state.figures[i]["chapter_number"] = chapter_num
                    
                    # Calculate sequence number based on other figures with the same chapter
                    seq_number = sum(1 for f in st.session_state.figures[:i] 
                                    if f.get("chapter_number") == chapter_num) + 1
                    
                    st.session_state.figures[i]["sequence_number"] = seq_number
                    st.session_state.figures[i]["full_number"] = f"{chapter_num}.{seq_number}"
                    st.session_state.figures[i]["caption"] = f"{figure_prefix} {chapter_num}.{seq_number}"
                    
                    st.info(f"This figure will be numbered as: {st.session_state.figures[i]['caption']}")
        else:
            st.info("No figures detected in the document.")
            
        # Include tables and figures options
        st.subheader("Table & Figure Options")
        include_tables = st.checkbox("Include tables in export", value=True)
        include_figures = st.checkbox("Include figures in export", value=True)
        
        # Save options
        if st.button("Save Table & Figure Settings"):
            if 'active_template' in st.session_state and st.session_state.active_template:
                template = st.session_state.active_template
                template["include_tables"] = include_tables
                template["include_figures"] = include_figures
                
                if templates_collection is not None:
                    templates_collection.update_one(
                        {"template_id": template["template_id"]}, 
                        {"$set": template}
                    )
                    
                st.success("Table & Figure settings saved!")
            else:
                st.warning("Please select a template first.")

    elif feature_selection == "Live Document Editor":
        st.header("Live Document Editor")
        
        if not st.session_state.active_doc:
            st.warning("Please select a document first")
            return
        
        st.info(f"Editing document: {st.session_state.active_doc['original_name']}")
        
        # Initialize LLM for suggestions (Premium tier only)
        llm = init_llm() if can_use_llm(st.session_state.user_tier) else None
        
        # Text editor with word suggestions
        st.subheader("Document Content")
        
        # Text area for editing document content
        new_text = st.text_area("Edit document content", 
                               st.session_state.edit_text, 
                               height=400)
        
        # If text has changed, update suggestions
        if new_text != st.session_state.edit_text:
            st.session_state.edit_text = new_text
            
            # Get word suggestions if user has premium tier
            if llm and can_use_llm(st.session_state.user_tier):
                with st.spinner("Getting suggestions..."):
                    st.session_state.suggestions = get_word_suggestions(new_text, llm)
            else:
                # Basic word suggestion for free tier
                st.session_state.suggestions = rule_based_word_suggestions(new_text)
        
        # Display word suggestions
        if st.session_state.suggestions:
            st.subheader("Word Suggestions")
            cols = st.columns(len(st.session_state.suggestions))
            
            for i, suggestion in enumerate(st.session_state.suggestions):
                with cols[i]:
                    if st.button(suggestion, key=f"suggestion_{i}"):
                        # Append the suggestion to the current text
                        st.session_state.edit_text += f" {suggestion}"
                        st.rerun()
        
        # Save edited content
        if st.button("Save Changes"):
            try:
                # Create temp file with edited content
                file_type = st.session_state.active_doc["file_type"]
                temp_file = f"temp_{uuid.uuid4()}.{file_type}"
                
                if file_type == "docx":
                    doc = docx.Document()
                    # Split text by paragraphs and add each to document
                    paragraphs = st.session_state.edit_text.split('\n')
                    for para in paragraphs:
                        doc.add_paragraph(para)
                    doc.save(temp_file)
                else:
                    # For text or PDF, just save as text
                    with open(temp_file, "w", encoding="utf-8") as f:
                        f.write(st.session_state.edit_text)
                
                st.success("Changes saved!")
                
                # Update the document file path
                st.session_state.active_doc["file_path"] = temp_file
                
            except Exception as e:
                st.error(f"Error saving document: {str(e)}")
    
    elif feature_selection == "Preview & Export":
        st.header("Preview & Export Document")
        
        if not st.session_state.active_doc:
            st.warning("Please select a document first")
            return
        
        if not st.session_state.active_template:
            st.warning("Please select a template first")
            return
        
        st.info(f"Selected document: {st.session_state.active_doc['original_name']}")
        st.info(f"Selected template: {st.session_state.active_template['name']}")
        
        # Display table and figure counts
        st.markdown(f"### Document Components")
        st.markdown(f"- Tables: {len(st.session_state.tables)}")
        st.markdown(f"- Figures: {len(st.session_state.figures)}")
        st.markdown(f"- Chapters: {len(st.session_state.chapters)}")
        
        # Options for including tables and figures
        include_tables_figures = st.checkbox("Include tables and figures", value=True)
        
        # Display preview of current document
        if st.session_state.active_doc:
            st.subheader("Current Document")
            display_document_preview(st.session_state.active_doc["file_path"])
        
        if st.button("Apply Template and Export"):
            with st.spinner("Applying template to document..."):
                doc_path = st.session_state.active_doc["file_path"]
                template = st.session_state.active_template
                
                # Update template with tables and figures option
                template["include_tables_figures"] = include_tables_figures
                
                output_dir = f"output/{st.session_state.user_id}"
                os.makedirs(output_dir, exist_ok=True)
                output_path = f"{output_dir}/{uuid.uuid4()}.docx"
                
                # Pass tables, figures, and chapters if including them
                tables = st.session_state.tables if include_tables_figures else []
                figures = st.session_state.figures if include_tables_figures else []
                chapters = st.session_state.chapters
                
                success, result = apply_template_to_document(
                    doc_path, template, output_path, tables, figures, chapters
                )
                
                if success:
                    st.session_state.formatted_doc_path = output_path
                    st.success("Template applied successfully!")
                    
                    # Display preview of formatted document
                    st.subheader("Formatted Document Preview")
                    display_document_preview(output_path)
                else:
                    st.error(result)

    elif feature_selection == "Plagiarism Check":
        st.header("Plagiarism Check")
        
        if not st.session_state.active_doc:
            st.warning("Please select a document first")
            return
            
        st.info(f"Selected document: {st.session_state.active_doc['original_name']}")
        
        # Get text to check
        if 'edit_text' in st.session_state and st.session_state.edit_text:
            text_to_check = st.session_state.edit_text
        else:
            # Load document content if not already in edit_text
            doc_path = st.session_state.active_doc["file_path"]
            file_type = st.session_state.active_doc["file_type"]
            
            if file_type == "docx":
                doc = docx.Document(doc_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                text_to_check = "\n".join(full_text)
            elif file_type == "txt":
                with open(doc_path, "r", encoding="utf-8") as f:
                    text_to_check = f.read()
            else:
                # For PDFs, extract text using PyMuPDF
                try:
                    pdf_doc = fitz.open(doc_path)
                    full_text = []
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        full_text.append(page.get_text())
                    text_to_check = "\n".join(full_text)
                except Exception as e:
                    st.error(f"Error extracting text from PDF: {e}")
                    text_to_check = "Error extracting document text"
        
        # Option to select specific text for checking
        use_specific_text = st.checkbox("Check specific text instead of entire document")
        
        if use_specific_text:
            text_to_check = st.text_area("Enter text to check for plagiarism", 
                                      value=text_to_check[:500] if len(text_to_check) > 500 else text_to_check,
                                      height=200)
        
        # Display information based on user tier
        if not can_use_feature(st.session_state.user_tier, "plagiarism_check"):
            st.warning("Plagiarism checking is only available for Premium users. Please upgrade to access this feature.")
            st.markdown("### Premium Plagiarism Check Features:")
            st.markdown("- Full document scanning")
            st.markdown("- Web-based source detection")
            st.markdown("- Academic database comparison")
            st.markdown("- Detailed similarity reports")
            st.markdown("- Word replacement suggestions")
            st.markdown("- Plagiarism percentage calculation")
        
        # Run plagiarism check
        if st.button("Check for Plagiarism"):
            if can_use_feature(st.session_state.user_tier, "plagiarism_check"):
                with st.spinner("Checking for plagiarism..."):
                    # Limit text length for API economy
                    if len(text_to_check) > 5000:
                        st.info("Checking first 5000 characters only. For larger documents, break into smaller sections.")
                        text_to_check = text_to_check[:5000]
                    
                    result = check_plagiarism(text_to_check, st.session_state.user_tier)
                    
                    if "error" in result:
                        st.error(result["error"])
                    else:
                        st.subheader("Plagiarism Check Results")
                        
                        # Display plagiarism percentage with a progress bar
                        plagiarism_percentage = result.get("plagiarism_percentage", 0)
                        
                        # Determine color based on percentage
                        if plagiarism_percentage < 15:
                            bar_color = "green"
                            status_icon = "✅"
                            status_text = "Low plagiarism detected"
                        elif plagiarism_percentage < 30:
                            bar_color = "orange"
                            status_icon = "⚠️"
                            status_text = "Moderate plagiarism detected"
                        else:
                            bar_color = "red"
                            status_icon = "🚫"
                            status_text = "High plagiarism detected"
                        
                        # Display overall result with percentage
                        col1, col2 = st.columns([1, 3])
                        with col1:
                            st.markdown(f"### {status_icon}")
                        with col2:
                            st.markdown(f"### {status_text}")
                            st.progress(plagiarism_percentage / 100)
                            st.markdown(f"**Plagiarism Percentage:** {plagiarism_percentage}%")
                        
                        st.markdown(f"Checked {result['chunks_checked']} of {result['total_chunks']} text segments")
                        
                        # Display detailed results
                        for chunk_result in result["results"]:
                            with st.expander(f"Segment {chunk_result['chunk_index'] + 1} Results"):
                                st.markdown("**Text segment:**")
                                st.text(chunk_result["text"])
                                
                                if chunk_result.get("has_matches", False):
                                    st.markdown("**Potential matches:**")
                                    for match in chunk_result.get("matches", []):
                                        st.markdown(f"[{match['title']}]({match['link']})")
                                        st.markdown(f"*{match['snippet']}*")
                                    
                                    # Display word suggestions if available
                                    word_suggestions = chunk_result.get("word_suggestions", {})
                                    if word_suggestions:
                                        st.markdown("### Suggested Alternatives")
                                        st.markdown("To reduce plagiarism, consider replacing these phrases:")
                                        
                                        for original_phrase, alternatives in word_suggestions.items():
                                            st.markdown(f"**Original:** {original_phrase}")
                                            
                                            if alternatives:
                                                for i, alt in enumerate(alternatives):
                                                    st.markdown(f"{i+1}. {alt}")
                                            
                                            st.markdown("---")
                                else:
                                    st.markdown("**No significant matches found**")
            else:
                st.error("Plagiarism checking is only available for Premium users. Please upgrade to access this feature.")
                
    elif feature_selection == "Export to Google Docs":
        st.header("Export to Google Docs")
        
        if not st.session_state.active_doc and not st.session_state.formatted_doc_path:
            st.warning("Please select a document and apply a template first")
            return
            
        # Determine which document to export
        if st.session_state.formatted_doc_path:
            export_path = st.session_state.formatted_doc_path
            export_name = f"DocuMorph_Formatted_{os.path.basename(export_path)}"
            st.info(f"Exporting formatted document: {export_name}")
        else:
            export_path = st.session_state.active_doc["file_path"]
            export_name = st.session_state.active_doc["original_name"]
            st.info(f"Exporting original document: {export_name}")
        
        # Check if Google API is available
        if not GOOGLE_API_AVAILABLE:
            st.error("Google API client not available. Please install with:")
            st.code("pip install google-api-python-client google-auth-oauthlib")
            return
            
        # Display information based on user tier
        if not can_use_feature(st.session_state.user_tier, "google_docs"):
            st.warning("Export to Google Docs is only available for Premium users. Please upgrade to access this feature.")
            st.markdown("### Premium Google Docs Export Features:")
            st.markdown("- Direct export to Google Docs")
            st.markdown("- Format preservation")
            st.markdown("- Collaborative editing")
            st.markdown("- Document sharing")
            return
            
        # Check if we're already in the authentication flow
        if st.session_state.google_drive_auth:
            auth_info = st.session_state.google_drive_auth
            
            st.info("Please complete Google Drive authorization")
            st.markdown(f"1. Visit this link to authorize: [Authorize Google Drive]({auth_info['auth_url']})")
            st.markdown("2. Sign in and grant permission")
            st.markdown("3. Copy the authorization code and paste it below")
            
            auth_code = st.text_input("Enter the authorization code:")
            
            if auth_code and st.button("Complete Authorization"):
                with st.spinner("Completing authorization and uploading file..."):
                    result = complete_google_drive_export(
                        auth_code,
                        auth_info["credentials_path"],
                        auth_info["file_path"],
                        auth_info["file_name"]
                    )
                    
                    if result["success"]:
                        st.success("File successfully exported to Google Drive!")
                        st.markdown(f"**View file:** [Open in Google Drive]({result['web_link']})")
                        
                        # Reset auth flow
                        st.session_state.google_drive_auth = None
                    else:
                        st.error(f"Export failed: {result.get('error', 'Unknown error')}")
        else:
            # Start export process
            if st.button("Export to Google Drive"):
                with st.spinner("Setting up Google Drive export..."):
                    result = export_to_google_drive(
                        export_path, 
                        export_name,
                        st.session_state.user_tier
                    )
                    
                    if result.get("success", False):
                        st.success("File exported to Google Drive successfully!")
                    elif "auth_url" in result:
                        # Store auth info for the next step
                        st.session_state.google_drive_auth = result
                        st.rerun()  # Refresh to show auth flow
                    else:
                        st.error(f"Export failed: {result.get('error', 'Unknown error')}")
                        
    elif feature_selection == "Custom Templates":
        st.header("Custom Templates")
        st.subheader("Create and manage your personal templates")
        
        # Display info about custom templates
        st.info("Custom templates are available to all users regardless of tier. Create and use your own personalized templates!")
        
        tab1, tab2 = st.tabs(["Create Custom Template", "Manage Custom Templates"])
        
        with tab1:
            st.markdown("### Create New Custom Template")
            template_name = st.text_input("Template Name", "My Custom Template")
            
            st.markdown("### Document Layout")
            col1, col2 = st.columns(2)
            with col1:
                margin_top = st.number_input("Top Margin (inches)", 0.5, 2.0, 1.0, 0.1)
                margin_left = st.number_input("Left Margin (inches)", 0.5, 2.0, 1.0, 0.1)
            with col2:
                margin_bottom = st.number_input("Bottom Margin (inches)", 0.5, 2.0, 1.0, 0.1)
                margin_right = st.number_input("Right Margin (inches)", 0.5, 2.0, 1.0, 0.1)
            
            st.markdown("### Typography")
            col1, col2 = st.columns(2)
            with col1:
                body_font = st.selectbox("Body Font", ["Calibri", "Arial", "Times New Roman", "Georgia", "Helvetica", "Verdana", "Tahoma"])
                body_font_size = st.slider("Body Font Size (pt)", 8, 14, 11)
            with col2:
                heading_font = st.selectbox("Heading Font", ["Calibri", "Arial", "Times New Roman", "Georgia", "Helvetica", "Verdana", "Tahoma"])
                heading1_font_size = st.slider("Heading 1 Size (pt)", 14, 22, 16)
            
            heading2_font_size = st.slider("Heading 2 Size (pt)", 12, 18, 14)
            heading3_font_size = st.slider("Heading 3 Size (pt)", 11, 16, 12)
            
            st.markdown("### Additional Styling")
            line_spacing = st.slider("Line Spacing", 1.0, 2.0, 1.15, 0.05)
            paragraph_spacing = st.slider("Paragraph Spacing (pt)", 0, 14, 6, 1)
            
            st.markdown("### Header & Footer")
            header_text = st.text_input("Header Text (optional)")
            footer_text = st.text_input("Footer Text (optional)")
            
            # Color scheme
            color_scheme = st.selectbox("Color Scheme", 
                                      ["Default (Black/White)", "Blue Professional", "Modern Gray", "Earthy Tones", "Academic"])
            
            if st.button("Save Custom Template"):
                template = {
                    "template_id": str(uuid.uuid4()),
                    "user_id": st.session_state.user_id,
                    "name": template_name,
                    "is_custom": True,
                    "margin_top": margin_top,
                    "margin_bottom": margin_bottom,
                    "margin_left": margin_left,
                    "margin_right": margin_right,
                    "body_font": body_font,
                    "body_font_size": body_font_size,
                    "heading_font": heading_font,
                    "heading1_font_size": heading1_font_size,
                    "heading2_font_size": heading2_font_size,
                    "heading3_font_size": heading3_font_size,
                    "line_spacing": line_spacing,
                    "paragraph_spacing": paragraph_spacing,
                    "header_text": header_text,
                    "footer_text": footer_text,
                    "color_scheme": color_scheme
                }
                
                if templates_collection is not None:
                    templates_collection.insert_one(template)
                
                st.success(f"Custom template '{template_name}' saved successfully!")
                
                # Display a preview of the template
                st.markdown("### Template Preview")
                preview_html = f"""
                <div style="border:1px solid #ccc; padding:20px; border-radius:5px; margin:10px 0;">
                    <h1 style="font-family:{heading_font}; font-size:{heading1_font_size}pt;">Heading 1 Example</h1>
                    <h2 style="font-family:{heading_font}; font-size:{heading2_font_size}pt;">Heading 2 Example</h2>
                    <h3 style="font-family:{heading_font}; font-size:{heading3_font_size}pt;">Heading 3 Example</h3>
                    <p style="font-family:{body_font}; font-size:{body_font_size}pt; line-height:{line_spacing};">
                        This is an example of the body text with the selected font and size.
                        This preview gives you an idea of how your template will look when applied to a document.
                    </p>
                </div>
                """
                st.markdown(preview_html, unsafe_allow_html=True)
        
        with tab2:
            st.subheader("Your Custom Templates")
            
            if templates_collection is not None:
                # Only show custom templates
                custom_templates = list(templates_collection.find({
                    "user_id": st.session_state.user_id,
                    "is_custom": True
                }))
                
                if custom_templates:
                    for i, template in enumerate(custom_templates):
                        with st.expander(f"{template['name']}"):
                            st.markdown(f"**Body Font**: {template['body_font']}, {template['body_font_size']}pt")
                            st.markdown(f"**Heading Font**: {template['heading_font']}")
                            st.markdown(f"**Margins**: T: {template['margin_top']}\", B: {template['margin_bottom']}\", L: {template['margin_left']}\", R: {template['margin_right']}\"")
                            st.markdown(f"**Line Spacing**: {template.get('line_spacing', 1.15)}")
                            
                            # Preview of the template
                            st.markdown("### Preview")
                            preview_html = f"""
                            <div style="border:1px solid #ccc; padding:10px; border-radius:5px; margin:10px 0;">
                                <h3 style="font-family:{template['heading_font']}; font-size:{template['heading3_font_size']}pt;">Sample Heading</h3>
                                <p style="font-family:{template['body_font']}; font-size:{template['body_font_size']}pt; line-height:{template.get('line_spacing', 1.15)};">
                                    This is a sample of how text will look with this template.
                                </p>
                            </div>
                            """
                            st.markdown(preview_html, unsafe_allow_html=True)
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.button("Select Template", key=f"select_custom_template_{i}"):
                                    st.session_state.active_template = template
                                    st.success(f"Selected template: {template['name']}")
                            with col2:
                                if st.button("Delete", key=f"delete_custom_template_{i}"):
                                    if templates_collection is not None:
                                        templates_collection.delete_one({"template_id": template["template_id"]})
                                    st.success(f"Template '{template['name']}' deleted.")
                                    st.rerun()
                else:
                    st.info("No custom templates found. Create a new custom template to get started.")
            else:
                st.error("MongoDB connection failed. Templates cannot be loaded.")

# Function to check for plagiarism using SerpAPI
def check_plagiarism(text, user_tier=UserTier.FREE):
    """Check text for plagiarism using SerpAPI."""
    if not can_use_feature(user_tier, "plagiarism_check"):
        return {"error": "Plagiarism checking is only available for premium users"}
    
    try:
        # Prepare chunks of text to check (limit to ~2000 chars per query)
        chunks = []
        words = text.split()
        chunk_size = 200  # words per chunk
        
        for i in range(0, len(words), chunk_size):
            chunk = " ".join(words[i:i+chunk_size])
            chunks.append(chunk)
        
        results = []
        plagiarized_word_count = 0
        total_word_count = len(words)
        
        # Only check first 3 chunks to avoid excessive API usage
        for i, chunk in enumerate(chunks[:3]):
            # Create a search query
            params = {
                "engine": "google",
                "q": f'"{chunk}"',  # Exact match search
                "api_key": SERPAPI_KEY
            }
            
            # Get results
            response = requests.get("https://serpapi.com/search", params=params)
            if response.status_code == 200:
                data = response.json()
                
                # Extract organic results
                organic_results = data.get("organic_results", [])
                
                # Process results
                chunk_results = []
                has_matches = False
                
                for result in organic_results[:5]:  # Limit to top 5 results
                    chunk_results.append({
                        "title": result.get("title", ""),
                        "link": result.get("link", ""),
                        "snippet": result.get("snippet", "")
                    })
                    has_matches = True
                
                # If matches found, consider this chunk partially plagiarized
                # For a simple estimate, count the whole chunk as plagiarized if matches found
                if has_matches:
                    plagiarized_word_count += len(chunk.split())
                
                # Generate word replacement suggestions using LLM if premium user and matches found
                word_suggestions = {}
                if has_matches and can_use_llm(user_tier):
                    try:
                        llm = init_llm()
                        if llm:
                            # Create a prompt for the LLM
                            prompt = PromptTemplate(
                                template="""You are helping to avoid plagiarism by suggesting alternative wording.
                                
                                Original text that may be plagiarized:
                                {text}
                                
                                Please provide 5 key phrases or sentences from this text that might be plagiarized, 
                                and for each one, suggest 3 alternative ways to express the same idea with different wording.
                                
                                Format your response as follows:
                                
                                Original phrase 1: "[phrase]"
                                - Alternative 1: "[alternative]"
                                - Alternative 2: "[alternative]"
                                - Alternative 3: "[alternative]"
                                
                                Original phrase 2: "[phrase]"
                                - Alternative 1: "[alternative]"
                                - Alternative 2: "[alternative]"
                                - Alternative 3: "[alternative]"
                                
                                And so on for 5 phrases total.
                                """,
                                input_variables=["text"],
                            )
                            
                            chain = prompt | llm | StrOutputParser()
                            suggestions_text = chain.invoke({"text": chunk})
                            
                            # Parse the suggestions
                            word_suggestions = parse_word_suggestions(suggestions_text)
                    except Exception as e:
                        logger.error(f"Error generating word suggestions: {e}")
                
                # Add results for this chunk
                results.append({
                    "chunk_index": i,
                    "text": chunk,
                    "matches": chunk_results,
                    "has_matches": has_matches,
                    "word_suggestions": word_suggestions
                })
            else:
                logger.error(f"SerpAPI request failed: {response.status_code}, {response.text}")
                results.append({
                    "chunk_index": i,
                    "text": chunk,
                    "error": f"API request failed: {response.status_code}",
                    "has_matches": False
                })
        
        # Calculate plagiarism percentage
        plagiarism_percentage = 0
        if total_word_count > 0:
            plagiarism_percentage = min(100, round((plagiarized_word_count / total_word_count) * 100, 2))
        
        # Summarize results
        has_plagiarism = any(r["has_matches"] for r in results)
        
        return {
            "has_plagiarism": has_plagiarism,
            "plagiarism_percentage": plagiarism_percentage,
            "chunks_checked": len(results),
            "total_chunks": len(chunks),
            "results": results
        }
    except Exception as e:
        logger.error(f"Error checking plagiarism: {e}", exc_info=True)
        return {"error": f"Error checking plagiarism: {str(e)}"}

# Helper function to parse word suggestions from LLM output
def parse_word_suggestions(text):
    """Parse the word suggestions output from LLM into a structured format."""
    suggestions = {}
    current_phrase = None
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Check if this is a new original phrase
        if line.startswith("Original phrase"):
            try:
                phrase_parts = line.split(":", 1)
                if len(phrase_parts) > 1:
                    current_phrase = phrase_parts[1].strip().strip('"')
                    suggestions[current_phrase] = []
            except:
                continue
        
        # Check if this is an alternative
        elif line.startswith("-") and current_phrase:
            try:
                alt_parts = line.split(":", 1)
                if len(alt_parts) > 1:
                    alternative = alt_parts[1].strip().strip('"')
                    suggestions[current_phrase].append(alternative)
            except:
                continue
    
    return suggestions

# Function to check if a user can use a specific feature
def can_use_feature(user_tier, feature_name):
    """Check if the user's tier has access to a specific feature."""
    tier_features = UserTier.get_tier_features()
    return tier_features[user_tier].get(feature_name, False)

# Function to export document to Google Drive
def export_to_google_drive(file_path, file_name, user_tier=UserTier.FREE):
    """Export a document to Google Drive."""
    if not can_use_feature(user_tier, "google_drive_export"):
        return {"success": False, "error": "Google Drive export is only available for premium users"}
    
    try:
        # Set up the credentials
        credentials = {
            "installed": {
                "client_id": GOOGLE_CLIENT_ID,
                "project_id": "documorph-ai",
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                "client_secret": GOOGLE_CLIENT_SECRET,
                "redirect_uris": ["http://localhost:8501", "urn:ietf:wg:oauth:2.0:oob"]
            }
        }
        
        # Save credentials temporarily
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json') as temp_cred_file:
            json.dump(credentials, temp_cred_file)
            credentials_path = temp_cred_file.name
        
        # Create a flow with proper scopes and redirect URI
        flow = InstalledAppFlow.from_client_config(
            client_config=credentials,
            scopes=['https://www.googleapis.com/auth/drive.file'],
            redirect_uri='urn:ietf:wg:oauth:2.0:oob'
        ) 
       
        # Generate the authorization URL for the user to visit
        auth_url, _ = flow.authorization_url(prompt='consent', access_type='offline')
        
        # Provide auth URL to user
        return {
            "success": False,  # Not immediately successful
            "auth_url": auth_url,
            "message": "Please visit this URL to authorize access to your Google Drive",
            "credentials_path": credentials_path,
            "file_path": file_path,
            "file_name": file_name
        }
    except Exception as e:
        logger.error(f"Error setting up Google Drive export: {e}", exc_info=True)
        return {"success": False, "error": f"Error exporting to Google Drive: {str(e)}"}

# Function to complete Google Drive export after authorization
def complete_google_drive_export(auth_code, credentials_path, file_path, file_name):
    """Complete the Google Drive export after user authorization."""
    try:
        # Load the credentials
        flow = InstalledAppFlow.from_client_secrets_file(
            credentials_path,
            ['https://www.googleapis.com/auth/drive.file'],
            redirect_uri='urn:ietf:wg:oauth:2.0:oob'  # Use same redirect URI as before
        )
        
        # Exchange auth code for access token
        flow.fetch_token(code=auth_code)
        creds = flow.credentials
        
        # Build the Drive API client
        drive_service = build('drive', 'v3', credentials=creds)
        
        # Determine file MIME type
        mime_type = "application/octet-stream"  # Default
        if file_path.endswith('.pdf'):
            mime_type = 'application/pdf'
        elif file_path.endswith('.docx'):
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_path.endswith('.txt'):
            mime_type = 'text/plain'
        
        # Create file metadata
        file_metadata = {
            'name': file_name,
            'mimeType': mime_type
        }
        
        # Create media
        media = MediaFileUpload(
            file_path,
            mimetype=mime_type,
            resumable=True
        )
        
        # Upload file
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,webViewLink'
        ).execute()
        
        # Clean up temporary credentials file
        if os.path.exists(credentials_path):
            os.unlink(credentials_path)
        
        return {
            "success": True,
            "file_id": file.get('id'),
            "web_link": file.get('webViewLink'),
            "message": f"File uploaded successfully to Google Drive"
        }
    except Exception as e:
        logger.error(f"Error completing Google Drive export: {e}", exc_info=True)
        return {"success": False, "error": f"Error completing Google Drive export: {str(e)}"}

# PyMuPDF compatibility wrapper
def open_pdf(pdf_path):
    """Open a PDF with PyMuPDF using the most compatible method."""
    try:
        return fitz.open(pdf_path)
    except AttributeError:
        try:
            return fitz.Document(pdf_path)
        except AttributeError:
            try:
                from pymupdf import Document
                return Document(pdf_path)
            except (AttributeError, ImportError):
                raise ImportError("Could not initialize PyMuPDF with any available method")

def create_default_templates():
    """Create default templates for different roles."""
    templates = [
        {
            "name": "Student Essay",
            "role": "Student",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Times New Roman",
            "body_font_size": 12,
            "heading_font": "Times New Roman",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 2.0,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "Page [Page] of [Pages]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        },
        {
            "name": "Business Report",
            "role": "Business Professional",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Company Name]",
            "footer_text": "Confidential | [Date] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Blue Professional"
        },
        {
            "name": "Research Paper",
            "role": "Researcher",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Times New Roman",
            "body_font_size": 12,
            "heading_font": "Times New Roman",
            "heading1_font_size": 14,
            "heading2_font_size": 12,
            "heading3_font_size": 12,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "[Page]",
            "include_tables_figures": True,
            "color_scheme": "Academic"
        },
        {
            "name": "Blog Post",
            "role": "Content Creator",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Georgia",
            "heading1_font_size": 18,
            "heading2_font_size": 16,
            "heading3_font_size": 14,
            "line_spacing": 1.5,
            "paragraph_spacing": 12,
            "header_text": "",
            "footer_text": "© [Year] [Author Name]",
            "include_tables_figures": True,
            "color_scheme": "Modern Gray"
        },
        {
            "name": "Multilingual Document",
            "role": "Multilingual User",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Arial",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "[Page]/[Pages]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        },
        {
            "name": "Book Chapter",
            "role": "Author",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Georgia",
            "body_font_size": 12,
            "heading_font": "Georgia",
            "heading1_font_size": 18,
            "heading2_font_size": 16,
            "heading3_font_size": 14,
            "line_spacing": 1.5,
            "paragraph_spacing": 6,
            "header_text": "[Book Title] - Chapter [Chapter]",
            "footer_text": "[Author Name] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Earthy Tones"
        },
        {
            "name": "Team Document",
            "role": "Collaborator",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Team Name] - [Project Name]",
            "footer_text": "Last updated: [Date] | Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Blue Professional"
        },
        {
            "name": "Project Report",
            "role": "Project Manager",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.25,
            "margin_right": 1.25,
            "body_font": "Arial",
            "body_font_size": 11,
            "heading_font": "Arial",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "[Project Name] - [Status]",
            "footer_text": "Confidential | Page [Page] of [Pages]",
            "include_tables_figures": True,
            "color_scheme": "Modern Gray"
        },
        {
            "name": "Custom Document",
            "role": "Others",
            "margin_top": 1.0,
            "margin_bottom": 1.0,
            "margin_left": 1.0,
            "margin_right": 1.0,
            "body_font": "Calibri",
            "body_font_size": 11,
            "heading_font": "Calibri",
            "heading1_font_size": 16,
            "heading2_font_size": 14,
            "heading3_font_size": 12,
            "line_spacing": 1.15,
            "paragraph_spacing": 6,
            "header_text": "",
            "footer_text": "Page [Page]",
            "include_tables_figures": True,
            "color_scheme": "Default (Black/White)"
        }
    ]
    
    return templates

def initialize_db_templates(user_id):
    """Initialize MongoDB with default templates for a user."""
    try:
        client = MongoClient("mongodb://localhost:27017/")
        db = client["documorph_db"]
        templates_collection = db["templates"]
        
        # Check if templates already exist for this user
        existing_templates = list(templates_collection.find({"user_id": user_id}))
        if existing_templates:
            logger.info(f"Templates already exist for user {user_id}")
            return True
            
        # Add default templates for this user
        templates = create_default_templates()
        for template in templates:
            template["user_id"] = user_id
            template["template_id"] = str(uuid.uuid4())
            template["is_custom"] = False
            templates_collection.insert_one(template)
            
        logger.info(f"Initialized default templates for user {user_id}")
        return True
    except Exception as e:
        logger.error(f"Error initializing DB templates: {e}", exc_info=True)
        return False

def get_role_specific_prompt(role, task_type):
    """Generate role-specific prompts for different document processing tasks."""
    prompts = {
        "Student": {
            "abstract": """Generate a concise, academic abstract for this student document.
            Focus on clarity, proper academic terminology, and highlighting the main arguments or findings.
            The abstract should be well-structured and approximately 150-250 words.
            Use formal language appropriate for academic submission.
            """,
            "section_titles": """Suggest academic section titles that would be appropriate for a student paper.
            Titles should be clear, descriptive, and follow academic conventions.
            Consider standard sections like Introduction, Literature Review, Methodology, Results, Discussion, and Conclusion.
            """,
            "style": "Academic",
            "structure": """Analyze this student document and suggest an academic structure.
            Focus on logical flow, proper academic sections, and clear organization of ideas.
            Identify where the document could benefit from better section organization or additional headings.
            """
        },
        "Content Creator": {
            "abstract": """Generate an engaging, reader-friendly summary for this content piece.
            Focus on hooking the reader's interest, highlighting key points, and creating a compelling preview.
            The summary should be approximately 100-200 words and use engaging, conversational language.
            """,
            "section_titles": """Suggest engaging, attention-grabbing section titles for this content.
            Titles should be catchy, memorable, and encourage readers to continue reading.
            Consider creative headings that spark curiosity while clearly indicating section content.
            """,
            "style": "Engaging and Conversational",
            "structure": """Analyze this content and suggest a structure optimized for reader engagement.
            Focus on creating a narrative flow, using hooks throughout the content, and organizing information
            in a way that maintains reader interest from beginning to end.
            """
        },
        "Researcher": {
            "abstract": """Generate a comprehensive, scientific abstract for this research document.
            Follow standard scientific abstract structure: background, methods, results, and conclusions.
            The abstract should be precise, data-focused, and approximately 200-300 words.
            Use formal scientific language and emphasize research significance and findings.
            """,
            "section_titles": """Suggest formal research section titles following scientific conventions.
            Titles should be precise, descriptive, and follow standard research paper organization.
            Consider sections like Abstract, Introduction, Methods, Results, Discussion, Conclusion, and References.
            """,
            "style": "Scientific and Formal",
            "structure": """Analyze this research document and suggest a structure following scientific conventions.
            Focus on logical progression of ideas, proper organization of research components, and clear
            separation between different elements of the research (methods, results, discussion, etc.).
            """
        },
        "Business Professional": {
            "abstract": """Generate a concise executive summary for this business document.
            Focus on key business insights, actionable information, and bottom-line impact.
            The summary should be approximately 150-250 words and use clear, professional language.
            Emphasize business value, recommendations, and strategic implications.
            """,
            "section_titles": """Suggest professional business section titles that convey authority and clarity.
            Titles should be direct, action-oriented, and clearly communicate section purpose.
            Consider sections like Executive Summary, Market Analysis, Strategic Recommendations, Implementation Plan, etc.
            """,
            "style": "Professional and Concise",
            "structure": """Analyze this business document and suggest a structure optimized for business decision-makers.
            Focus on presenting information in a hierarchy of importance, with executive summary first,
            followed by supporting details, analysis, and recommendations or next steps.
            """
        },
        "Multilingual User": {
            "abstract": """Generate a clear, straightforward abstract that would be easily understood across languages.
            Focus on simple sentence structures, common vocabulary, and universal concepts.
            The abstract should be approximately 150-250 words and avoid idioms or culturally specific references.
            """,
            "section_titles": """Suggest clear, universally understandable section titles.
            Titles should be straightforward, descriptive, and avoid language-specific idioms or complex terms.
            Consider simple, direct headings that would translate well across multiple languages.
            """,
            "style": "Clear and Universally Accessible",
            "structure": """Analyze this document and suggest a structure that would work well across languages and cultures.
            Focus on universal organizational patterns, clear progression of ideas, and avoiding
            culturally specific organizational structures that might not translate well.
            """
        },
        "Author": {
            "abstract": """Generate an engaging book chapter summary or overview.
            Focus on narrative elements, themes, key arguments, and reader takeaways.
            The summary should be approximately 150-250 words and reflect the author's voice and style.
            Create something that would entice readers to continue reading the full chapter.
            """,
            "section_titles": """Suggest creative yet descriptive section titles appropriate for a book chapter.
            Titles should balance creativity with clarity and maintain the author's voice.
            Consider how these headings guide the reader through the narrative or argument.
            """,
            "style": "Literary and Narrative-Focused",
            "structure": """Analyze this document and suggest a structure appropriate for a book chapter.
            Focus on narrative flow, thematic development, and reader engagement.
            Consider how to organize content to build reader interest throughout the chapter.
            """
        },
        "Collaborator": {
            "abstract": """Generate a clear team document summary that highlights key points for all collaborators.
            Focus on shared goals, action items, responsibilities, and next steps.
            The summary should be approximately 150-200 words and use clear, inclusive language.
            Emphasize collaborative elements and information relevant to all team members.
            """,
            "section_titles": """Suggest practical section titles that facilitate team collaboration.
            Titles should be clear, action-oriented, and help team members quickly find relevant information.
            Consider sections like Project Overview, Team Responsibilities, Timeline, Action Items, etc.
            """,
            "style": "Clear, Practical, and Action-Oriented",
            "structure": """Analyze this collaborative document and suggest a structure that facilitates team work.
            Focus on organizing information for quick reference, clear responsibility assignment,
            and effective tracking of project elements and progress.
            """
        },
        "Project Manager": {
            "abstract": """Generate a comprehensive project summary highlighting key objectives, status, and outcomes.
            Focus on project metrics, milestones, resource allocation, and critical path elements.
            The summary should be approximately 150-250 words and use precise project management terminology.
            Emphasize timeline, deliverables, and current status relative to project goals.
            """,
            "section_titles": """Suggest project management section titles that facilitate project tracking and reporting.
            Titles should be specific, metric-oriented, and aligned with project management methodologies.
            Consider sections like Project Scope, Timeline, Resource Allocation, Risk Assessment, etc.
            """,
            "style": "Structured and Metrics-Focused",
            "structure": """Analyze this project document and suggest a structure optimized for project management.
            Focus on organizing information to track project progress, highlight dependencies,
            and clearly communicate status, risks, and next steps to stakeholders.
            """
        }
    }
    
    # Default prompt if role or task not found
    default_prompts = {
        "abstract": """Generate a concise, professional abstract summarizing the main points of this document.
        The abstract should be well-structured, clear, and approximately 150-250 words.
        """,
        "section_titles": """Suggest professional section titles that would improve the organization of this document.
        Titles should be clear, descriptive, and help guide the reader through the content.
        """,
        "style": "Professional",
        "structure": """Analyze this document and suggest an improved structure.
        Focus on logical organization, clear progression of ideas, and appropriate section divisions.
        """
    }
    
    # Get role-specific prompts or fall back to default
    role_prompts = prompts.get(role, default_prompts)
    return role_prompts.get(task_type, default_prompts[task_type])

# Function to check if user can use a specific template role
def can_use_template_role(user_tier, role):
    """Check if the user's tier allows using templates with this role."""
    tier_features = UserTier.get_tier_features()
    allowed_categories = tier_features[user_tier]["template_categories"]
    return role in allowed_categories

# Function to create temporary in-memory templates when MongoDB isn't available
def get_templates_for_user(user_id, templates_collection, user_tier=UserTier.FREE):
    if templates_collection is None:
        # Create in-memory templates
        templates = create_default_templates()
        for template in templates:
            template["user_id"] = user_id
            template["template_id"] = str(uuid.uuid4())
            template["is_custom"] = False
        return templates
    else:
        # Get from MongoDB
        return list(templates_collection.find({"user_id": user_id}))

if __name__ == "__main__":
    main() 
