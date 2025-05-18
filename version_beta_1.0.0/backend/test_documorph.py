import streamlit as st
import os
import sys

# This is a simple test to make sure the imports work
def test_imports():
    try:
        print("Testing import: docx")
        import docx
        print("Testing import: docx.shared")
        from docx.shared import Pt, RGBColor, Inches
        print("Testing import: docx.enum.text")
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("Testing import: langchain_community.document_loaders")
        from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
        print("Testing import: langchain_groq")
        from langchain_groq import ChatGroq
        print("All imports successful!")
        return True
    except ImportError as e:
        print(f"Import error: {e}")
        print(f"Python path: {sys.path}")
        return False

# Test creating a simple document with python-docx
def test_docx_creation():
    try:
        print("Testing document creation...")
        import docx
        from docx.shared import Pt
        
        # Create a new document
        doc = docx.Document()
        
        # Add a title
        doc.add_heading('DocuMorph AI Test Document', 0)
        
        # Add a paragraph
        p = doc.add_paragraph('This is a test document created by python-docx to verify functionality.')
        
        # Style the paragraph
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        
        # Save the document
        os.makedirs('test_output', exist_ok=True)
        doc.save('test_output/test_document.docx')
        print("Document created successfully at: test_output/test_document.docx")
        return True
    except Exception as e:
        print(f"Error creating document: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Testing DocuMorph AI dependencies...")
    print(f"Python version: {sys.version}")
    print(f"Current directory: {os.getcwd()}")
    
    imports_ok = test_imports()
    docx_ok = test_docx_creation()
    
    if imports_ok and docx_ok:
        print("All tests passed! DocuMorph AI dependencies are working correctly.")
    else:
        print("Some tests failed. Please check the error messages above.") 