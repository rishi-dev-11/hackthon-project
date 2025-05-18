import streamlit as st
import os
import uuid
import logging
from pymongo import MongoClient
import tempfile
import json
# import base64 # Not actively used, but kept from original for potential future use
from datetime import datetime
import docx
from docx import Document as PythonDocxDocument # Renamed to avoid conflict
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import fitz  # PyMuPDF

import cv2
import numpy as np
from PIL import Image
import io
import pytesseract
import spacy
from langdetect import detect
from google_auth_oauthlib.flow import InstalledAppFlow # type: ignore
from googleapiclient.discovery import build # type: ignore
from googleapiclient.http import MediaFileUpload # type: ignore
from google.auth.transport.requests import Request # type: ignore
from google.oauth2.credentials import Credentials # type: ignore
import magic
from transformers import M2M100ForConditionalGeneration, M2M100Tokenizer # type: ignore
import zipfile
from pymongo import MongoClient
from dotenv import load_dotenv
import os
load_dotenv()

# ReportLab Imports
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. PDF generation will be disabled. Install with 'pip install reportlab'")

# LangChain & LLM Imports
try:
    from langchain_groq import ChatGroq
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    logging.warning("LangChain or related libraries not available. LLM features will be disabled. Install with 'pip install langchain langchain-groq sentence-transformers faiss-cpu'")


# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

st.set_page_config(page_title="DocuMorph AI", layout="wide")

# Check for Unstructured module (already present)
try:
    from unstructured.partition.pdf import partition_pdf
    UNSTRUCTURED_AVAILABLE = True
    logger.info("Unstructured module loaded successfully.")
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logger.warning("Unstructured module not available (used for advanced OCR). Install with: pip install unstructured[all-docs]")


# Initialize NLP model (spaCy)
@st.cache_resource
def load_nlp_model():
    try:
        nlp_model = spacy.load("en_core_web_sm")
        logger.info("Spacy NLP model loaded successfully.")
        return nlp_model, True
    except IOError:
        logger.warning("Spacy model 'en_core_web_sm' not found. Please download it (python -m spacy download en_core_web_sm). NLP features will be limited.")
        return None, False
    except ImportError:
        logger.warning("Spacy not installed. NLP features will be disabled.")
        return None, False

nlp, NLP_AVAILABLE = load_nlp_model()

# Load environment variables at the start
load_dotenv(override=True)  # Force reload environment variables
mongo_uri = os.getenv("MONGO_URI", "mongodb://localhost:27017")  # Default to local MongoDB if not specified
if mongo_uri:
    logger.info("MongoDB URI loaded")
    if st.sidebar:  # Only show if sidebar exists
        st.sidebar.success("MongoDB URI loaded")
else:
    logger.error("MongoDB URI not loaded")
    if st.sidebar:
        st.sidebar.error("MongoDB URI not loaded")

# Initialize MongoDB client
@st.cache_resource
def init_mongodb():
    try:
        # Get MongoDB URI from environment variable or use local default
        mongo_uri = os.getenv("MONGO_URI", "mongodb://localhost:27017")
        logger.info("Attempting to connect to MongoDB...")

        # Create MongoDB client with connection options
        client = MongoClient(
            mongo_uri,
            serverSelectionTimeoutMS=5000,  # Reduced timeout for local connection
            connectTimeoutMS=5000,
            socketTimeoutMS=5000,
            appname="documorph"
        )
        
        # Test connection
        client.admin.command('ping')
        
        # Get database and collections
        db = client.get_database("documorph_db")
        templates_collection = db.get_collection("templates")
        documents_collection = db.get_collection("documents")
        
        logger.info("MongoDB connection established successfully")
        return templates_collection, documents_collection
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"MongoDB connection error: {error_msg}", exc_info=True)
        
        if "Connection refused" in error_msg:
            st.error("MongoDB connection refused. Please check if MongoDB is running locally.")
        else:
            st.error(f"MongoDB connection failed: {error_msg}")
        
        return None, None

# Load translation model
@st.cache_resource
def load_translation_model():
    try:
        model = M2M100ForConditionalGeneration.from_pretrained("facebook/m2m100_418M")
        tokenizer = M2M100Tokenizer.from_pretrained("facebook/m2m100_418M")
        logger.info("Translation model loaded successfully.")
        return model, tokenizer
    except Exception as e:
        logger.error(f"Translation model loading error: {e}")
        return None, None

# --- LLM and Embeddings Initialization ---
@st.cache_resource
def init_llm(groq_api_key):
    if not LANGCHAIN_AVAILABLE:
        return None
    try:
        # Consider making the model name configurable
        llm = ChatGroq(
            temperature=0.2,
            model="llama3-70b-8192", # Or other Llama3 models available on Groq like llama3-8b-8192
            api_key=groq_api_key
        )
        logger.info("ChatGroq LLM initialized.")
        return llm
    except Exception as e:
        logger.error(f"Error initializing Groq LLM: {e}", exc_info=True)
        # st.error(f"Error initializing LLM. Ensure API key is correct and Groq service is accessible.")
        return None

@st.cache_resource
def init_embeddings_model():
    if not LANGCHAIN_AVAILABLE:
        return None
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("HuggingFace embeddings model initialized.")
        return embeddings
    except Exception as e:
        logger.error(f"Error initializing embeddings model: {e}", exc_info=True)
        # st.error(f"Error initializing embeddings model.")
        return None

# Placeholder for vector store init if needed later for RAG-like features
@st.cache_resource
def init_vector_store(texts, embeddings, chat_id_or_doc_id):
    if not LANGCHAIN_AVAILABLE or not texts or not embeddings:
        return None
    try:
        # vector_store_path = f"vector_stores/faiss_index_{chat_id_or_doc_id}" # Example path
        # if os.path.exists(vector_store_path):
        #     vector_store = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
        # else:
        vector_store = FAISS.from_texts(texts, embeddings)
        # vector_store.save_local(vector_store_path) # Persist if needed
        logger.info(f"Vector store initialized for {chat_id_or_doc_id}.")
        return vector_store
    except Exception as e:
        logger.error(f"Error initializing vector store: {e}", exc_info=True)
        return None

# Translation function (already present)
def translate(text, src_lang, tgt_lang, model, tokenizer):
    # ... (same as before)
    if not text or not model or not tokenizer:
        return text
    try:
        tokenizer.src_lang = src_lang
        encoded = tokenizer(text, return_tensors="pt") # type: ignore
        generated_tokens = model.generate(**encoded, forced_bos_token_id=tokenizer.get_lang_id(tgt_lang))
        return tokenizer.decode(generated_tokens[0], skip_special_tokens=True)
    except Exception as e:
        logger.error(f"Translation error for text '{text[:50]}...': {e}")
        return text # Fallback to original text

# OCR function (already present, slightly enhanced error handling)
def extract_text_from_image_or_pdf_page(file_bytes, lang='eng', is_pdf=False):
    # ... (same as before)
    try:
        if is_pdf: # Handle PDF page as image for OCR
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            if len(pdf_doc) > 0:
                page = pdf_doc[0] # Process first page for simplicity
                pix = page.get_pixmap()
                image_bytes_for_ocr = pix.tobytes("png")
            else:
                return "Empty PDF"
            pdf_doc.close()
        else:
            image_bytes_for_ocr = file_bytes

        image = Image.open(io.BytesIO(image_bytes_for_ocr))
        img_array = np.array(image.convert('RGB'))
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        temp_filename = f"temp_ocr_image_{uuid.uuid4().hex}.png"
        cv2.imwrite(temp_filename, thresh)
        text = ""
        if UNSTRUCTURED_AVAILABLE:
            try:
                elements = partition_pdf(filename=temp_filename, strategy="hi_res")
                text_parts = [el.text for el in elements if hasattr(el, 'text') and el.text] # type: ignore
                if text_parts: text = "\n".join(text_parts)
                logger.info(f"OCR with Unstructured extracted: {text[:100]}...")
            except Exception as e_unstructured:
                logger.warning(f"Unstructured OCR failed for {temp_filename}: {e_unstructured}")
        if not text:
            try:
                text = pytesseract.image_to_string(Image.open(temp_filename), lang=lang).strip()
                logger.info(f"OCR with Pytesseract extracted: {text[:100]}...")
            except Exception as e_pytesseract:
                logger.warning(f"Pytesseract OCR failed for {temp_filename}: {e_pytesseract}")
        os.remove(temp_filename)
        return text if text else "No text detected in image"
    except Exception as e:
        logger.error(f"Error during OCR process: {e}", exc_info=True)
        if 'temp_filename' in locals() and os.path.exists(temp_filename): os.remove(temp_filename)
        return f"Error during OCR: {str(e)}"


# NLP-based section detection (modified to use template config)
def detect_sections(doc_text, lang, template_config):
    # ... (same as before)
    if lang != 'en' or not NLP_AVAILABLE or not nlp: return []
    section_keywords_str = template_config.get("section_keywords", "abstract,introduction,references,conclusion,methodology,results,discussion")
    user_defined_keywords = [kw.strip().lower() for kw in section_keywords_str.split(',')]
    doc = nlp(doc_text[:1000000]) # Process up to 1M characters for performance
    sections = []
    current_section_text = []
    current_section_title = None

    for para_text in doc_text.split('\n\n'): # A bit more robust than sentence splitting for headings
        clean_para_text = para_text.strip()
        if not clean_para_text: continue

        # Check if the paragraph itself is a keyword-based heading
        # This is a simple heuristic
        first_line = clean_para_text.split('\n')[0].strip()
        if len(first_line.split()) <= 7 and first_line.lower() in user_defined_keywords:
            if current_section_title and current_section_text: # Save previous section
                 sections.append({"title": current_section_title, "level": 1, "content": "\n\n".join(current_section_text)})
            current_section_title = first_line
            current_section_text = [clean_para_text.replace(first_line, "", 1).strip()] if clean_para_text != first_line else []
            logger.info(f"Detected section by keyword: {current_section_title}")
        elif current_section_title: # If we are inside a section, append text
            current_section_text.append(clean_para_text)
        # else: # Text before any recognized section heading (could be part of an implicit first section)
            # if not sections and not current_section_title: # Start an implicit first section
            #     current_section_title = "Introduction" # Default name
            #     current_section_text.append(clean_para_text)


    if current_section_title and current_section_text: # Save the last section
        sections.append({"title": current_section_title, "level": 1, "content": "\n\n".join(current_section_text)})
    
    # If no sections detected by keywords but text exists, treat whole doc as one section for title suggestion
    if not sections and doc_text.strip():
        sections.append({"title": "Full Document Content", "level": 1, "content": doc_text.strip()})

    return sections


# Google Docs export (already present)
def export_to_google_docs(file_path, file_name):
    # ... (same as before, ensure token.json handling for web apps or use service accounts if deployed)
    # For Streamlit Cloud, client-side OAuth flow (InstalledAppFlow) won't work directly.
    # You'd typically use st.secrets for service account credentials or a more complex OAuth backend.
    # This version assumes local execution or a setup where token.json can be managed.
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', ['https://www.googleapis.com/auth/drive.file'])
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception as e_refresh:
                logger.error(f"Failed to refresh Google token: {e_refresh}")
                # Fallback to re-running flow if refresh fails
                if os.path.exists('credentials.json'):
                    flow = InstalledAppFlow.from_client_secrets_file('credentials.json', ['https://www.googleapis.com/auth/drive.file'])
                    creds = flow.run_local_server(port=0) # This blocks in a deployed app
                else:
                    st.error("Google API credentials.json missing.")
                    return "Export failed: credentials.json missing."
        else:
            if os.path.exists('credentials.json'):
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', ['https://www.googleapis.com/auth/drive.file'])
                creds = flow.run_local_server(port=0)
            else:
                st.error("Google API credentials.json missing.")
                return "Export failed: credentials.json missing."
        if creds: # Save the credentials for the next run
            with open('token.json', 'w') as token:
                token.write(creds.to_json())
        else:
            return "Could not obtain Google credentials."
    try:
        service = build('drive', 'v3', credentials=creds)
        # ... (rest of the Google Docs export logic, unchanged)
        file_metadata = {'name': file_name, 'mimeType': 'application/vnd.google-apps.document'}
        media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        request = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink')
        response = None; progress_bar = st.progress(0); status_text = st.empty()
        status_text.text("Starting Google Docs export..."); # type: ignore
        while response is None:
            status, response = request.next_chunk() # type: ignore
            if status: progress = int(status.progress() * 100); progress_bar.progress(progress); status_text.text(f"Uploading to Google Docs: {progress}%") # type: ignore
        progress_bar.progress(100); status_text.success(f"Exported to Google Docs: {response.get('webViewLink')}") # type: ignore
        return f"Exported to Google Docs: {response.get('webViewLink')}"
    except Exception as e:
        logger.error(f"Google Docs export error: {e}", exc_info=True)
        st.error(f"Google Docs export error: {e}")
        return f"Error during Google Docs export: {str(e)}"


def apply_template_to_document(content_doc, template_config, lang, logo_path=None, llm_structure_suggestions=None):
    # ... (The core python-docx formatting engine - largely unchanged from before)
    # One potential modification: if llm_structure_suggestions are provided,
    # it might influence how headings are applied. For now, this function remains
    # focused on the template_config. LLM suggestions would typically be used
    # to *prepare* the content_doc or guide the user *before* this final formatting step.
    # However, if llm_structure_suggestions included identified text for headings,
    # we could try to match and apply styles here. This is complex.

    # For now, apply_template_to_document remains procedural based on template_config.
    # LLM suggestions would be used to help the *user* structure their input,
    # or the `input_docx_obj` could be pre-processed based on LLM suggestions
    # before being passed to this function.
    output_doc = PythonDocxDocument()
    # ... (rest of the function is the same as the previous version)
    # Apply document-wide settings from template
    section = output_doc.sections[0]
    section.page_height = Inches(float(template_config.get('page_height', 11)))
    section.page_width = Inches(float(template_config.get('page_width', 8.5)))
    section.left_margin = Inches(float(template_config.get('margin_left', 1)))
    section.right_margin = Inches(float(template_config.get('margin_right', 1)))
    section.top_margin = Inches(float(template_config.get('margin_top', 1)))
    section.bottom_margin = Inches(float(template_config.get('margin_bottom', 1)))
    styles = output_doc.styles
    body_style_name = 'DocuMorphBody'; # type: ignore
    try: body_style = styles.add_style(body_style_name, WD_STYLE_TYPE.PARAGRAPH) # type: ignore
    except: body_style = styles[body_style_name] # type: ignore
    body_font = body_style.font; body_font.name = template_config.get('body_font', 'Calibri'); body_font.size = Pt(int(template_config.get('body_size', 11))) # type: ignore
    for i in range(1, 4):
        heading_style_name = f'DocuMorphHeading{i}'; # type: ignore
        try: heading_style = styles.add_style(heading_style_name, WD_STYLE_TYPE.PARAGRAPH) # type: ignore
        except: heading_style = styles[heading_style_name] # type: ignore
        heading_style.base_style = styles['Normal']; h_font = heading_style.font # type: ignore
        h_font.name = template_config.get(f'h{i}_font', template_config.get('body_font', 'Calibri')); h_font.size = Pt(int(template_config.get(f'h{i}_size', 16 - 2*i))); h_font.bold = template_config.get(f'h{i}_bold', True) # type: ignore
        pf = heading_style.paragraph_format; pf.space_before = Pt(int(template_config.get(f'h{i}_space_before', 12))); pf.space_after = Pt(int(template_config.get(f'h{i}_space_after', 6))) # type: ignore
    if logo_path and os.path.exists(logo_path) and template_config.get('use_logo', False):
        try:
            for current_section in output_doc.sections:
                header = current_section.header; # type: ignore
                if not header.paragraphs: header_para = header.add_paragraph() # type: ignore
                else: header_para = header.paragraphs[0] # type: ignore
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT; run = header_para.add_run(); run.add_picture(logo_path, width=Inches(float(template_config.get('logo_width_inches', 1.0)))) # type: ignore
        except Exception as e_logo: logger.error(f"Error adding logo: {e_logo}"); st.warning(f"Could not add logo: {e_logo}")
    if template_config.get('use_title_page', False):
        title_text = template_config.get('document_title', 'Document Title'); title_para = output_doc.add_paragraph(); title_run = title_para.add_run(title_text) # type: ignore
        title_run.font.name = template_config.get('title_font', template_config.get('body_font', 'Calibri')); title_run.font.size = Pt(int(template_config.get('title_size', 24))); title_run.font.bold = True # type: ignore
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER; title_para.paragraph_format.space_after = Pt(18) # type: ignore
        if 'document_author' in template_config and template_config['document_author']: author_para = output_doc.add_paragraph(template_config['document_author']); author_para.style = body_style_name; author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
        if template_config.get('include_date_on_title_page', False):
            date_formats = {'en': "%B %d, %Y", 'fr': "%d %B %Y", 'de': "%d. %B %Y", 'es': "%d de %B de %Y"}; date_str = datetime.now().strftime(date_formats.get(lang, "%Y-%m-%d"))
            date_para = output_doc.add_paragraph(date_str); date_para.style = body_style_name; date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
        output_doc.add_page_break()
    if template_config.get('include_toc', False):
        toc_title = template_config.get('toc_title', 'Table of Contents'); output_doc.add_paragraph(toc_title, style='DocuMorphHeading1') # type: ignore
        toc_paragraph = output_doc.add_paragraph(); run = toc_paragraph.add_run('[Table of Contents Placeholder - Update in Word]'); run.font.italic = True # type: ignore
        fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin'); instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'TOC \\o "1-3" \\h \\z \\u'; fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end') # type: ignore
        p_element = toc_paragraph._p; p_element.append(fldChar); p_element.append(instrText); p_element.append(fldChar2) # type: ignore
        output_doc.add_page_break()
    full_text_for_nlp = "\n".join([p.text for p in content_doc.paragraphs if p.text.strip()])
    detected_doc_sections = detect_sections(full_text_for_nlp, lang, template_config); detected_section_idx = 0
    for element in content_doc.element.body: # type: ignore
        if element.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(element, content_doc); text = para.text.strip() # type: ignore
            if not text: continue
            is_section_heading = False
            if lang == 'en' and detected_section_idx < len(detected_doc_sections):
                if text.lower() == detected_doc_sections[detected_section_idx]["title"].lower(): # Use "title" from detect_sections
                    level_to_apply = detected_doc_sections[detected_section_idx].get("level", 1)
                    output_doc.add_paragraph(text, style=f'DocuMorphHeading{level_to_apply}') # type: ignore
                    detected_section_idx += 1; is_section_heading = True
            if not is_section_heading:
                is_heading_from_style = False
                if para.style and para.style.name.lower().startswith('heading'): # type: ignore
                    try:
                        level = int(para.style.name[-1]); # type: ignore
                        if 1 <= level <= 3: output_doc.add_paragraph(text, style=f'DocuMorphHeading{level}'); is_heading_from_style = True # type: ignore
                    except ValueError: pass
                if not is_heading_from_style:
                    new_para = output_doc.add_paragraph(text, style=body_style_name); pf = new_para.paragraph_format # type: ignore
                    pf.line_spacing = float(template_config.get('line_spacing', 1.15)); pf.space_before = Pt(int(template_config.get('para_space_before_pt', 0))); pf.space_after = Pt(int(template_config.get('para_space_after_pt', 8))) # type: ignore
        elif element.tag.endswith('tbl'):
            table = docx.table.Table(element, content_doc) # type: ignore
            if template_config.get('caption_position', 'Below') == 'Above': output_doc.add_paragraph(f"{template_config.get('table_caption_prefix', 'Table')} {len(output_doc.tables) + 1}: [Table Caption]", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
            new_table = output_doc.add_table(rows=0, cols=len(table.columns)); new_table.style = template_config.get('table_style', 'TableGrid') # type: ignore
            for row_idx, row in enumerate(table.rows):
                new_table.add_row(); new_row = new_table.rows[-1] # type: ignore
                for cell_idx, cell in enumerate(row.cells): new_row.cells[cell_idx].text = cell.text # type: ignore
            if template_config.get('caption_position', 'Below') == 'Below': output_doc.add_paragraph(f"{template_config.get('table_caption_prefix', 'Table')} {len(output_doc.tables)}: [Table Caption]", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    for current_section in output_doc.sections:
        if template_config.get('use_header', False):
            header = current_section.header; # type: ignore
            if not header.paragraphs: header.add_paragraph(); # type: ignore
            header_para = header.paragraphs[0]; header_para.text = template_config.get('header_text', '') # type: ignore
        if template_config.get('use_footer', False):
            footer = current_section.footer; # type: ignore
            if not footer.paragraphs: footer.add_paragraph(); # type: ignore
            footer_para = footer.paragraphs[0]; footer_para.text = template_config.get('footer_text', '') # type: ignore
            if template_config.get('include_page_numbers', False):
                if footer_para.text: footer_para.add_run("\t"); # type: ignore
                footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT # type: ignore
                run = footer_para.add_run(); fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin'); instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'; fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate'); fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end') # type: ignore
                run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3) # type: ignore
    return output_doc


def process_uploaded_document(uploaded_file, template_config, logo_file=None):
    # ... (same as before: file reading, MIME type, OCR, language detection, logo handling)
    full_text_content = ""
    input_docx_obj = None 
    temp_input_file_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_uploaded_file:
            tmp_uploaded_file.write(uploaded_file.getvalue()); temp_input_file_path = tmp_uploaded_file.name
        mime_type = magic.Magic(mime=True).from_file(temp_input_file_path)
        logger.info(f"Processing '{uploaded_file.name}' with MIME type: {mime_type}")
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            input_docx_obj = PythonDocxDocument(temp_input_file_path)
            full_text_content = "\n\n".join([p.text for p in input_docx_obj.paragraphs])
        elif mime_type == 'text/plain':
            with open(temp_input_file_path, 'r', encoding='utf-8') as f: full_text_content = f.read()
            input_docx_obj = PythonDocxDocument(); input_docx_obj.add_paragraph(full_text_content)
        elif mime_type == 'application/pdf':
            pdf_doc = fitz.open(temp_input_file_path); text_parts = [page.get_text("text") for page in pdf_doc]; full_text_content = "\n\n".join(text_parts); pdf_doc.close()
            if not full_text_content.strip() or len(full_text_content.strip()) < 100:
                logger.info("Minimal text from PDF, trying OCR..."); ocr_text = extract_text_from_image_or_pdf_page(uploaded_file.getvalue(), is_pdf=True)
                if ocr_text and "Error" not in ocr_text and "No text detected" not in ocr_text: full_text_content = ocr_text
            input_docx_obj = PythonDocxDocument() # Create a new docx for PDF content
            # Attempt to add PDF paragraphs more intelligently
            for para_text in full_text_content.split('\n\n'): # Split by double newline as heuristic for paragraphs
                if para_text.strip(): input_docx_obj.add_paragraph(para_text.strip())
            if not input_docx_obj.paragraphs: input_docx_obj.add_paragraph(" ") # Ensure at least one paragraph
        elif mime_type in ['image/png', 'image/jpeg', 'image/tiff']:
            full_text_content = extract_text_from_image_or_pdf_page(uploaded_file.getvalue())
            input_docx_obj = PythonDocxDocument(); input_docx_obj.add_paragraph(full_text_content)
        else:
            st.error(f"Unsupported file type: {mime_type} for {uploaded_file.name}")
            if temp_input_file_path and os.path.exists(temp_input_file_path): os.remove(temp_input_file_path)
            return None, None
        try:
            sample_text_for_lang_detect = full_text_content[:1000] if full_text_content else "default"
            lang = detect(sample_text_for_lang_detect) if sample_text_for_lang_detect else 'en'
        except Exception as lang_e: logger.warning(f"Lang detect failed: {lang_e}. Defaulting to 'en'."); lang = 'en'
        logger.info(f"Detected language for {uploaded_file.name}: {lang}")
        # model_trans, tokenizer_trans = load_translation_model() # Translation logic (can be kept if needed)
        template_config.setdefault('toc_title', "Table of Contents"); template_config.setdefault('figure_caption_prefix', "Figure"); template_config.setdefault('table_caption_prefix', "Table")
        logo_path = None
        if logo_file:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(logo_file.name)[1]) as tmp_logo:
                tmp_logo.write(logo_file.getvalue()); logo_path = tmp_logo.name
        if not input_docx_obj:
             st.error(f"Could not create document object for {uploaded_file.name}.")
             if temp_input_file_path and os.path.exists(temp_input_file_path): os.remove(temp_input_file_path)
             if logo_path and os.path.exists(logo_path): os.remove(logo_path)
             return None, None
        output_doc = apply_template_to_document(input_docx_obj, template_config, lang, logo_path)
        temp_dir_for_output = os.path.join(tempfile.gettempdir(), "documorph_outputs"); os.makedirs(temp_dir_for_output, exist_ok=True)
        base_name = os.path.splitext(uploaded_file.name)[0]; docx_output_filename = f"formatted_{base_name}_{uuid.uuid4().hex[:8]}.docx"
        docx_output_path = os.path.join(temp_dir_for_output, docx_output_filename); output_doc.save(docx_output_path)
        logger.info(f"Formatted DOCX saved to: {docx_output_path}")
        if logo_path and os.path.exists(logo_path): os.remove(logo_path)
        if temp_input_file_path and os.path.exists(temp_input_file_path): os.remove(temp_input_file_path)
        return docx_output_path, full_text_content
    except Exception as e:
        logger.error(f"Error in process_uploaded_document for {uploaded_file.name}: {e}", exc_info=True)
        st.error(f"Failed to process {uploaded_file.name}: {e}")
        if 'temp_input_file_path' in locals() and temp_input_file_path and os.path.exists(temp_input_file_path): os.remove(temp_input_file_path) # type: ignore
        if 'logo_path' in locals() and logo_path and os.path.exists(logo_path): os.remove(logo_path) # type: ignore
        return None, None


# ReportLab PDF Generation (same as before)
def generate_pdf_with_reportlab(docx_path, template_config, pdf_output_path, logo_path=None):
    # ... (same as before)
    if not REPORTLAB_AVAILABLE: st.error("ReportLab not installed. PDF generation disabled."); return None
    try:
        doc_input = PythonDocxDocument(docx_path) # type: ignore
        pdf_doc = SimpleDocTemplate(pdf_output_path, pagesize=(float(template_config.get('page_width', 8.5)) * inch, float(template_config.get('page_height', 11)) * inch), leftMargin=float(template_config.get('margin_left', 1)) * inch, rightMargin=float(template_config.get('margin_right', 1)) * inch, topMargin=float(template_config.get('margin_top', 1)) * inch, bottomMargin=float(template_config.get('margin_bottom', 1)) * inch) # type: ignore
        styles = getSampleStyleSheet(); body_style_rl = ParagraphStyle('BodyRL', parent=styles['Normal'], fontName=template_config.get('body_font_rl', 'Helvetica'), fontSize=int(template_config.get('body_size', 11)), leading=int(template_config.get('body_size', 11)) * 1.2, spaceBefore=Pt(int(template_config.get('para_space_before_pt', 0))), spaceAfter=Pt(int(template_config.get('para_space_after_pt', 8)))) # type: ignore
        story = []
        if logo_path and os.path.exists(logo_path) and template_config.get('use_logo', False):
            try: img = RLImage(logo_path, width=float(template_config.get('logo_width_inches', 1.0)) * inch, height=0.5*inch); img.hAlign = 'RIGHT'; story.append(img); story.append(Spacer(1, 0.25 * inch)) # type: ignore
            except Exception as e_logo_pdf: logger.warning(f"Could not add logo to PDF: {e_logo_pdf}")
        for para in doc_input.paragraphs: # type: ignore
            text = para.text.strip()
            if not text: story.append(Spacer(1, body_style_rl.fontSize * 0.5)); continue # type: ignore
            current_style = body_style_rl # type: ignore
            if para.style.name.startswith("Heading 1") or para.style.name.startswith("DocuMorphHeading1"): temp_h1_style = ParagraphStyle('TempH1', parent=body_style_rl, fontSize=body_style_rl.fontSize + 4, fontName=template_config.get('h1_font_rl', 'Helvetica-Bold')); story.append(Paragraph(text, temp_h1_style)) # type: ignore
            elif para.style.name.startswith("Heading 2") or para.style.name.startswith("DocuMorphHeading2"): temp_h2_style = ParagraphStyle('TempH2', parent=body_style_rl, fontSize=body_style_rl.fontSize + 2, fontName=template_config.get('h2_font_rl', 'Helvetica-Bold')); story.append(Paragraph(text, temp_h2_style)) # type: ignore
            else: story.append(Paragraph(text, current_style)) # type: ignore
            story.append(Spacer(1, 0.1 * inch)) # type: ignore
        def my_header_footer(canvas_item, pdf_doc_template_item): # type: ignore
            canvas_item.saveState() # type: ignore
            if template_config.get('use_header', False) and template_config.get('header_text', ''): canvas_item.setFont(template_config.get('header_font_rl', 'Helvetica'), int(template_config.get('header_font_size_pt', 10))); header_text = template_config.get('header_text', ''); canvas_item.drawString(pdf_doc_template_item.leftMargin, pdf_doc_template_item.height + pdf_doc_template_item.topMargin - 0.5*inch, header_text) # type: ignore
            if template_config.get('use_footer', False): canvas_item.setFont(template_config.get('footer_font_rl', 'Helvetica'), int(template_config.get('footer_font_size_pt', 10))); footer_text = template_config.get('footer_text', ''); # type: ignore
            if template_config.get('include_page_numbers', False): footer_text += f" | Page {canvas_item.getPageNumber()}" # type: ignore
            canvas_item.drawString(pdf_doc_template_item.leftMargin, 0.5 * inch, footer_text) # type: ignore
            canvas_item.restoreState() # type: ignore
        pdf_doc.build(story, onFirstPage=my_header_footer, onLaterPages=my_header_footer) # type: ignore
        logger.info(f"PDF generated at: {pdf_output_path}"); return pdf_output_path
    except Exception as e_pdf: logger.error(f"Error generating PDF with ReportLab: {e_pdf}", exc_info=True); st.error(f"PDF Generation Error: {e_pdf}"); return None


# --- Template Form (same as before) ---
def create_template_form(existing_template=None):
    # Ensure default_template is always a dict
    default_template = existing_template if existing_template else {}
    template = {}
    st.subheader("Template Configuration")
    template_name = st.text_input("Template Name", default_template.get("name", "My Template"))
    template["shared"] = st.checkbox("Share Template with Team", value=default_template.get("shared", False))
    template["name"] = template_name

    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "Page Setup",
        "Typography (DOCX)",
        "Content Structure",
        "Tables & Figures",
        "Headers & Footers",
        "Advanced"
    ])  # type: ignore

    with tab1:
        st.markdown("#### Page Dimensions")
        template["page_width"] = st.number_input(
            "Page Width (inches)", min_value=5.0, max_value=14.0,
            value=float(default_template.get("page_width", 8.5)), step=0.1
        )
        template["page_height"] = st.number_input(
            "Page Height (inches)", min_value=5.0, max_value=20.0,
            value=float(default_template.get("page_height", 11.0)), step=0.1
        )
        st.markdown("#### Margins")
        template["margin_left"] = st.number_input(
            "Left Margin (inches)", min_value=0.2, max_value=3.0,
            value=float(default_template.get("margin_left", 1.0)), step=0.1
        )
        template["margin_right"] = st.number_input(
            "Right Margin (inches)", min_value=0.2, max_value=3.0,
            value=float(default_template.get("margin_right", 1.0)), step=0.1
        )
        template["margin_top"] = st.number_input(
            "Top Margin (inches)", min_value=0.2, max_value=3.0,
            value=float(default_template.get("margin_top", 1.0)), step=0.1
        )
        template["margin_bottom"] = st.number_input(
            "Bottom Margin (inches)", min_value=0.2, max_value=3.0,
            value=float(default_template.get("margin_bottom", 1.0)), step=0.1
        )

    with tab2:
        st.markdown("#### DOCX Font Settings")
        docx_fonts = ["Calibri", "Arial", "Times New Roman", "Helvetica", "Georgia", "Verdana", "Tahoma"]
        template["body_font"] = st.selectbox(
            "Body Text Font (DOCX)", docx_fonts,
            index=docx_fonts.index(default_template.get("body_font", "Calibri"))
        )
        template["body_size"] = st.number_input(
            "Body Text Size (pt)", min_value=8, max_value=20,
            value=int(default_template.get("body_size", 11))
        )
        for i in range(1, 4):
            st.markdown(f"--- \n ##### Heading {i} (DOCX)")
            template[f"h{i}_font"] = st.selectbox(
                f"H{i} Font", docx_fonts,
                index=docx_fonts.index(default_template.get(f"h{i}_font", "Calibri")),
                key=f"h{i}_font_docx"
            )
            template[f"h{i}_size"] = st.number_input(
                f"H{i} Size (pt)", min_value=10, max_value=30,
                value=int(default_template.get(f"h{i}_size", 16-2*i)),
                key=f"h{i}_size_docx"
            )
            template[f"h{i}_bold"] = st.checkbox(
                f"H{i} Bold", value=default_template.get(f"h{i}_bold", True),
                key=f"h{i}_bold_docx"
            )
            template[f"h{i}_space_before"] = st.number_input(
                f"H{i} Space Before (pt)", min_value=0, max_value=48,
                value=int(default_template.get(f"h{i}_space_before", 12 if i == 1 else 8)),
                key=f"h{i}_sb_docx"
            )
            template[f"h{i}_space_after"] = st.number_input(
                f"H{i} Space After (pt)", min_value=0, max_value=48,
                value=int(default_template.get(f"h{i}_space_after", 6 if i == 1 else 4)),
                key=f"h{i}_sa_docx"
            )
        st.markdown("--- \n #### Paragraph Spacing (DOCX)")
        template["line_spacing"] = st.number_input(
            "Line Spacing (e.g., 1.0, 1.15, 1.5)", min_value=1.0, max_value=3.0,
            value=float(default_template.get("line_spacing", 1.15)), step=0.05
        )
        template["para_space_before_pt"] = st.number_input(
            "Paragraph Space Before (pt)", min_value=0, max_value=24,
            value=int(default_template.get("para_space_before_pt", 0))
        )
        template["para_space_after_pt"] = st.number_input(
            "Paragraph Space After (pt)", min_value=0, max_value=24,
            value=int(default_template.get("para_space_after_pt", 8))
        )
        template["bullet_style"] = st.selectbox(
            "Bullet Style (DOCX - limited effect, primarily for new bullets)",
            ["default", "circle", "square"],
            index=["default", "circle", "square"].index(default_template.get("bullet_style", "default"))
        )
        st.markdown("#### ReportLab PDF Font Settings (subset, for basic PDF generation)")
        rl_fonts = ["Helvetica", "Times-Roman", "Courier", "Helvetica-Bold", "Times-Bold"]
        template["body_font_rl"] = st.selectbox(
            "Body Text Font (PDF)", rl_fonts,
            index=rl_fonts.index(default_template.get("body_font_rl", "Helvetica"))
        )
        template["h1_font_rl"] = st.selectbox(
            "H1 Font (PDF)", rl_fonts,
            index=rl_fonts.index(default_template.get("h1_font_rl", "Helvetica-Bold"))
        )

    with tab3:
        template["use_title_page"] = st.checkbox(
            "Include Title Page", value=default_template.get("use_title_page", True)
        )
        if template["use_title_page"]:
            template["document_title"] = st.text_input(
                "Default Document Title", default_template.get("document_title", "Document Title")
            )
            template["title_font"] = st.selectbox(
                "Title Font (DOCX)", docx_fonts,
                index=docx_fonts.index(default_template.get("title_font", "Calibri")),
                key="title_font_docx"
            )
            template["title_size"] = st.number_input(
                "Title Size (pt)", min_value=16, max_value=48,
                value=int(default_template.get("title_size", 24)),
                key="title_size_docx"
            )
            template["document_author"] = st.text_input(
                "Default Author", default_template.get("document_author", "Author Name")
            )
            template["include_date_on_title_page"] = st.checkbox(
                "Include Date on Title Page", value=default_template.get("include_date_on_title_page", True)
            )
        template["include_toc"] = st.checkbox(
            "Include Table of Contents", value=default_template.get("include_toc", True)
        )
        if template["include_toc"]:
            template["toc_title"] = st.text_input(
                "TOC Title", default_template.get("toc_title", "Table of Contents")
            )

    with tab4:
        table_styles_docx = ["TableGrid", "LightShading-Accent1", "MediumGrid1-Accent1", "TableNormal"]
        template["table_style"] = st.selectbox(
            "Table Style (DOCX)", table_styles_docx,
            index=table_styles_docx.index(default_template.get("table_style", "TableGrid"))
        )
        template["caption_position"] = st.radio(
            "Caption Position", ["Below", "Above"],
            index=["Below", "Above"].index(default_template.get("caption_position", "Below")),
            horizontal=True
        )
        template["figure_caption_prefix"] = st.text_input(
            "Figure Caption Prefix", default_template.get("figure_caption_prefix", "Figure")
        )
        template["table_caption_prefix"] = st.text_input(
            "Table Caption Prefix", default_template.get("table_caption_prefix", "Table")
        )

    with tab5:
        st.markdown("#### Header Settings")
        template["use_header"] = st.checkbox(
            "Use Header", value=default_template.get("use_header", False)
        )
        if template["use_header"]:
            template["header_text"] = st.text_input(
                "Header Text", default_template.get("header_text", "Document Header")
            )

        st.markdown("#### Footer Settings")
        template["use_footer"] = st.checkbox(
            "Use Footer", value=default_template.get("use_footer", True)
        )
        if template["use_footer"]:
            template["footer_text"] = st.text_input(
                "Footer Text", default_template.get("footer_text", "")
            )
            template["include_page_numbers"] = st.checkbox(
                "Include Page Numbers in Footer", value=default_template.get("include_page_numbers", True)
            )

        st.markdown("#### Logo Settings")
        template["use_logo"] = st.checkbox(
            "Use Logo in Header", value=default_template.get("use_logo", False)
        )
        if template["use_logo"]:
            template["logo_width_inches"] = st.number_input(
                "Logo Width in Header (inches)", min_value=0.5, max_value=3.0,
                value=float(default_template.get("logo_width_inches", 1.0)), step=0.1
            )

    with tab6:
        st.markdown("#### NLP Section Detection")
        template["section_keywords"] = st.text_area(
            "Section Keywords (comma-separated, for English NLP detection)",
            default_template.get("section_keywords", "abstract,introduction,references,conclusion,methodology,results,discussion,appendix"),
            help="Keywords used by NLP to identify potential section headings in English text."
        )
        st.markdown("#### ReportLab PDF Specifics")
        template["header_font_rl"] = st.selectbox(
            "Header Font (PDF)", rl_fonts,
            index=rl_fonts.index(default_template.get("header_font_rl", "Helvetica"))
        )
        template["header_font_size_pt"] = st.number_input(
            "Header Font Size (PDF, pt)", min_value=6, max_value=14,
            value=int(default_template.get("header_font_size_pt", 10))
        )
        template["footer_font_rl"] = st.selectbox(
            "Footer Font (PDF)", rl_fonts,
            index=rl_fonts.index(default_template.get("footer_font_rl", "Helvetica"))
        )
        template["footer_font_size_pt"] = st.number_input(
            "Footer Font Size (PDF, pt)", min_value=6, max_value=14,
            value=int(default_template.get("footer_font_size_pt", 10))
        )
    return template

# --- LLM Helper Functions ---
def generate_abstract_with_llm(llm, document_text, user_instructions=""): # Added user_instructions
    if not llm or not document_text.strip():
        return "LLM not available or document is empty."

    max_chars_for_abstract = 30000
    truncated_text = document_text[:max_chars_for_abstract]
    if len(document_text) > max_chars_for_abstract:
        st.warning(f"Document too long for full abstract generation. Using first {max_chars_for_abstract} characters.")

    # Base system prompt
    system_message = "You are an expert in scientific and technical writing, skilled at creating concise and informative abstracts."
    
    # Construct the human message, incorporating user instructions
    human_message_template = """Please generate a well-structured abstract for the following document content.
The abstract should summarize the main objectives, methods, key findings, and principal conclusions.
Ensure the language is professional and clear.

{user_directives}

Document Content:
{document_content}"""

    # Prepare user directives part
    user_directives_text = ""
    if user_instructions and user_instructions.strip():
        user_directives_text = f"Follow these specific instructions carefully:\n- {user_instructions.replace('.', '. ' if not user_instructions.endswith('.') else '')}" # Basic formatting for instructions

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_message),
        ("human", human_message_template)
    ])
    
    chain = prompt | llm | StrOutputParser()
    
    try:
        with st.spinner("🤖 LLM is generating abstract with your instructions..."):
            response = chain.invoke({
                "document_content": truncated_text,
                "user_directives": user_directives_text # Pass the formatted instructions
            })
        return response
    except Exception as e:
        logger.error(f"Error generating abstract with LLM: {e}", exc_info=True)
        return f"Error during abstract generation: {e}"
    
def suggest_section_title_with_llm(llm, section_content, current_title=""):
    if not llm or not section_content.strip():
        return "LLM not available or section content is empty."
    
    max_chars_for_title_suggestion = 10000
    truncated_content = section_content[:max_chars_for_title_suggestion]

    # Construct prompt parts separately to avoid f-string with backslash
    base_prompt = "You are an expert academic editor. Based on the following section content, suggest a concise, descriptive, and engaging new title for this section. "
    current_title_part = f'The current title is "{current_title}". ' if current_title else ""
    instruction_part = "Provide 2-3 alternative titles, each on a new line. Do not add any other commentary. Just the titles.\n\n"
    content_part = "Section Content:\n{section_content}"

    prompt_text = " ".join(part for part in [base_prompt, current_title_part, instruction_part, content_part] if part)

    prompt = ChatPromptTemplate.from_template(prompt_text)
    chain = prompt | llm | StrOutputParser()
    try:
        with st.spinner(f"🤖 LLM is suggesting titles..."):
            response = chain.invoke({"section_content": truncated_content})
        return response
    except Exception as e:
        logger.error(f"Error suggesting section title with LLM: {e}", exc_info=True)
        return f"Error during title suggestion: {e}"

def suggest_document_structure_llm(llm, document_text):
    if not llm or not document_text.strip():
        return "LLM not available or document is empty."
    max_chars_for_structure = 30000
    truncated_text = document_text[:max_chars_for_structure]
    if len(document_text) > max_chars_for_structure:
        st.warning(f"Document too long for full structure analysis. Using first {max_chars_for_structure} characters.")

    prompt_template = PromptTemplate.from_template(
        """Analyze the following document text and propose a hierarchical structure.
        Identify main sections and potential sub-sections.
        For each identified section, suggest a title and a heading level (e.g., H1, H2, H3).
        You can also provide the first few words of the section to help identify it.
        Present your output as a list of suggestions, like:
        - H1: Suggested Title 1 (Starts with: "First few words...")
        - H2: Suggested Sub-Title 1.1 (Starts with: "First few words...")
        - H1: Suggested Title 2 (Starts with: "First few words...")

        Do not try to rewrite the content, only suggest the structure.

        Document Text:
        {document_content}

        Suggested Structure:
        """
    )
    chain = prompt_template | llm | StrOutputParser()
    try:
        with st.spinner("🤖 LLM is analyzing document structure..."):
            response = chain.invoke({"document_content": truncated_text})
        return response
    except Exception as e:
        logger.error(f"Error suggesting document structure with LLM: {e}", exc_info=True)
        return f"Error during structure suggestion: {e}"


# --- Main App Logic ---
def main():
    st.title("📄 DocuMorph AI")
    st.subheader("The Intelligent Document Transformation Engine")

    # Initialize session state variables
    if 'user_id' not in st.session_state: st.session_state.user_id = str(uuid.uuid4())
    if 'current_template' not in st.session_state: st.session_state.current_template = None
    if 'processed_docx_files' not in st.session_state: st.session_state.processed_docx_files = {}
    if 'processed_pdf_files' not in st.session_state: st.session_state.processed_pdf_files = {}
    if 'last_uploaded_files_count' not in st.session_state: st.session_state.last_uploaded_files_count = 0
    if 'document_full_text' not in st.session_state: st.session_state.document_full_text = {} # {filename: text}
    if 'document_sections' not in st.session_state: st.session_state.document_sections = {} # {filename: [sections]}
    if 'groq_api_key' not in st.session_state: st.session_state.groq_api_key = ""
    

    templates_collection, documents_collection = init_mongodb()
    if templates_collection is None:
        st.error("Core Functionality Disabled: Failed to connect to Template Database. Please check MongoDB.")
        return # Stop further execution if DB is not available

    # LLM Initialization - requires API Key
    llm = None
    if LANGCHAIN_AVAILABLE:
        if st.session_state.groq_api_key:
            llm = init_llm(st.session_state.groq_api_key)
            if not llm:
                st.sidebar.error("LLM failed to initialize. Check API key and console logs.")
        # else: # Moved API key input to sidebar section
            # st.sidebar.warning("Groq API Key not set. LLM features disabled.")
    # embeddings_model = init_embeddings_model() # Initialize if needed for other features    # Sidebar for template management AND API Key
    with st.sidebar:
        st.header("📋 Template Management")
        user_and_shared_templates = []
        if templates_collection is not None:
            try:
                user_and_shared_templates = list(templates_collection.find({"$or": [{"user_id": st.session_state.user_id}, {"shared": True}]}).sort("name", 1))
            except Exception as db_e:
                st.error(f"Error fetching templates: {db_e}")
                logger.error(f"MongoDB find error: {db_e}")
        template_names = [t["name"] for t in user_and_shared_templates]
        if st.session_state.current_template and st.session_state.current_template["name"] not in template_names:
                st.session_state.current_template = None
            
        action_options = ["Select Template", "Create New Template", "Edit Selected Template"] if template_names else ["Create New Template"]
        default_action_index = 0
        if template_names and st.session_state.current_template:
            try:
                default_action_index = action_options.index("Edit Selected Template")
            except ValueError:
                pass
                
        template_action = st.radio("Actions", action_options, index=default_action_index)
        selected_template_for_display = None
        
        if template_action == "Select Template":
            current_selection_idx = 0
            if st.session_state.current_template:
                try:
                    current_selection_idx = template_names.index(st.session_state.current_template["name"])
                except ValueError:
                    st.session_state.current_template = None
                    
            selected_template_name = st.selectbox("Choose a template", template_names, index=current_selection_idx)
            if selected_template_name:
                newly_selected_template = next(
                    (t for t in user_and_shared_templates if t["name"] == selected_template_name),
                    None
                )
                if newly_selected_template and newly_selected_template != st.session_state.current_template:
                    st.session_state.current_template = newly_selected_template
                    st.session_state.processed_docx_files.clear()
                    st.session_state.processed_pdf_files.clear()
                    st.session_state.document_full_text.clear()
                    st.session_state.document_sections.clear()
                    st.rerun()
                selected_template_for_display = st.session_state.current_template
                
        elif template_action == "Create New Template":            new_template_data, submit_clicked = create_template_form()
            if submit_clicked:
                if templates_collection is not None:
                        if templates_collection.find_one({
                            "name": new_template_data["name"],
                            "user_id": st.session_state.user_id
                        }):
                            st.error(f"Template '{new_template_data['name']}' already exists.")
                        else:
                            new_template_data["user_id"] = st.session_state.user_id
                            new_template_data["created_at"] = datetime.now()
                            new_template_data["_id"] = str(uuid.uuid4())
                            try:
                                templates_collection.insert_one(new_template_data)
                                st.session_state.current_template = new_template_data
                                st.success(f"Template '{new_template_data['name']}' saved!")
                                st.rerun()
                            except Exception as db_e:
                                st.error(f"DB error saving template: {db_e}")
                                logger.error(f"MongoDB insert_one error: {db_e}")
                    else:
                        st.error("DB not connected.")
            selected_template_for_display = st.session_state.current_template
        elif template_action == "Edit Selected Template":
            if st.session_state.current_template:
                st.write(f"Editing: **{st.session_state.current_template['name']}**")
                is_owner = st.session_state.current_template.get("user_id") == st.session_state.user_id                if is_owner:
                    edited_template_data, submit_clicked = create_template_form(existing_template=st.session_state.current_template)
                    if submit_clicked:
                        if templates_collection is not None:
                                edited_template_data["user_id"] = st.session_state.user_id; edited_template_data["updated_at"] = datetime.now()
                                old_name = st.session_state.current_template["name"]; new_name = edited_template_data["name"]
                                if old_name != new_name and templates_collection.find_one({"name": new_name, "user_id": st.session_state.user_id}): st.error(f"Another template named '{new_name}' already exists.")
                                else:
                                    try: templates_collection.update_one({"_id": st.session_state.current_template["_id"], "user_id": st.session_state.user_id}, {"$set": edited_template_data}); st.session_state.current_template = edited_template_data; st.success(f"Template '{edited_template_data['name']}' updated!"); st.rerun()
                                    except Exception as db_e: st.error(f"DB error updating: {db_e}"); logger.error(f"MongoDB update_one error: {db_e}")
                            else: st.error("DB not connected.")
                else: st.warning("Can only edit owned templates.")
                selected_template_for_display = st.session_state.current_template
            else: st.warning("No template selected to edit.")
        if selected_template_for_display:
            with st.expander("Current Template Details", expanded=False): st.json(json.loads(json.dumps(selected_template_for_display, default=str)), expanded=False)

        st.markdown("---")
        st.header("🔑 API Keys & Settings")
        elif template_action == "Create New Template":
            new_template_data, submit_clicked = create_template_form()
            if submit_clicked:
                if templates_collection is not None:
                    if templates_collection.find_one({
                        "name": new_template_data["name"],
                        "user_id": st.session_state.user_id
                    }):
                        st.error(f"Template '{new_template_data['name']}' already exists.")
                    else:
                        new_template_data["user_id"] = st.session_state.user_id
                        new_template_data["created_at"] = datetime.now()
                        new_template_data["_id"] = str(uuid.uuid4())
                        try:
                            templates_collection.insert_one(new_template_data)
                            st.session_state.current_template = new_template_data
                            st.success(f"Template '{new_template_data['name']}' saved!")
                            st.rerun()
                        except Exception as db_e:
                            st.error(f"DB error saving template: {db_e}")
                            logger.error(f"MongoDB insert_one error: {db_e}")
                else:
                    st.error("DB not connected.")
            selected_template_for_display = st.session_state.current_template
        elif template_action == "Edit Selected Template":
            if st.session_state.current_template:
                st.write(f"Editing: **{st.session_state.current_template['name']}**")
                is_owner = st.session_state.current_template.get("user_id") == st.session_state.user_id
                if is_owner:
                    edited_template_data, submit_clicked = create_template_form(existing_template=st.session_state.current_template)
                    if submit_clicked:
                        if templates_collection is not None:
                            edited_template_data["user_id"] = st.session_state.user_id
                            edited_template_data["updated_at"] = datetime.now()
                            old_name = st.session_state.current_template["name"]
                            new_name = edited_template_data["name"]
                            if old_name != new_name and templates_collection.find_one({"name": new_name, "user_id": st.session_state.user_id}):
                                st.error(f"Another template named '{new_name}' already exists.")
                            else:
                                try:
                                    templates_collection.update_one(
                                        {"_id": st.session_state.current_template["_id"], "user_id": st.session_state.user_id},
                                        {"$set": edited_template_data}
                                    )
                                    st.session_state.current_template = edited_template_data
                                    st.success(f"Template '{edited_template_data['name']}' updated!")
                                    st.rerun()
                                except Exception as db_e:
                                    st.error(f"DB error updating: {db_e}")
                                    logger.error(f"MongoDB update_one error: {db_e}")
                        else:
                            st.error("DB not connected.")
                else:
                    st.warning("Can only edit owned templates.")
                selected_template_for_display = st.session_state.current_template
            else:
                st.warning("No template selected to edit.")
                            # Also try to detect sections here if NLP is available
                            if NLP_AVAILABLE and nlp and extracted_text:
                                lang = 'en' # Assuming English for spaCy for now, or detect lang
                                try: lang = detect(extracted_text[:500])
                                except: pass
                                if lang == 'en':
                                     st.session_state.document_sections[uploaded_file_item.name] = detect_sections(extracted_text, lang, st.session_state.current_template)


            # Create tabs for different features
            llm_tab, transform_tab = st.tabs(["🤖 LLM Assistance", "✨ Transform Documents"])

            # --- Advanced LLM Features Tab ---
            with llm_tab:
                st.subheader("Advanced Content Assistance (Powered by LLM)")
                if not llm:
                    st.warning("LLM not available. Please set a valid Groq API Key in the sidebar and ensure Langchain is installed.")
                else:
                    if not st.session_state.document_full_text:
                        st.info("Upload a document to enable LLM features. Text extraction might take a moment.")
                    else:
                        doc_names_for_llm = list(st.session_state.document_full_text.keys())
                        selected_doc_for_llm = st.selectbox("Select Document for LLM Analysis:", doc_names_for_llm, key="llm_doc_select")

                        if selected_doc_for_llm and selected_doc_for_llm in st.session_state.document_full_text:
                            current_doc_text = st.session_state.document_full_text[selected_doc_for_llm]
                            current_doc_sections = st.session_state.document_sections.get(selected_doc_for_llm, [])

                            st.markdown("---")
                            st.markdown("#### 📝 Generate Abstract with Specific Instructions")
                            
                            # User input for instructions
                            custom_abstract_instructions = st.text_area(
                                "Specific instructions for abstract generation (optional):",
                                placeholder="e.g., Keep it under 100 words. Focus on the impact. Target a general audience.",
                                key=f"custom_abs_instr_{selected_doc_for_llm}"
                            )
                            
                            if st.button("Generate Abstract", key=f"gen_abs_instr_{selected_doc_for_llm}"):
                                abstract = generate_abstract_with_llm(llm, current_doc_text, custom_abstract_instructions)
                                st.text_area("Suggested Abstract:", value=abstract, height=200, key=f"abs_area_instr_{selected_doc_for_llm}")
                            

                            st.markdown("---")
                            st.markdown("#### 🧐 Suggest Document Structure with Specific Instructions")
                            custom_structure_instructions = st.text_area(
                                "Specific instructions for structure analysis (optional):",
                                placeholder="e.g., Identify sections typical of a research proposal. Suggest H1 for main parts and H2 for sub-parts.",
                                key=f"custom_struct_instr_{selected_doc_for_llm}"
                            )
                            if st.button("Analyze Document Structure", key=f"gen_struct_instr_{selected_doc_for_llm}"):
                                structure_suggestion = suggest_document_structure_llm(llm, current_doc_text) # Update this call
                                st.text_area("Suggested Document Structure:", value=structure_suggestion, height=300, key=f"struct_area_instr_{selected_doc_for_llm}")


                            st.markdown("---")
                            st.markdown("#### ✨ Suggest Improved Section Titles with Specific Instructions")
                            if not current_doc_sections:
                                st.info("No sections detected by initial NLP pass for title suggestion, or document still processing.")
                            else:
                                section_titles_for_llm = [s["title"] for s in current_doc_sections if s.get("content")]
                                if not section_titles_for_llm:
                                    st.info("Sections detected, but no content found within them for title suggestion.")
                                else:
                                    selected_section_title_for_llm = st.selectbox("Select section to improve title:", section_titles_for_llm, key=f"sel_sec_title_instr_{selected_doc_for_llm}")
                                    
                                    custom_title_instructions = st.text_area(
                                        f"Specific instructions for title suggestions for '{selected_section_title_for_llm}' (optional):",
                                        placeholder="e.g., Make titles more engaging. Ensure titles are questions. Keep titles under 7 words.",
                                        key=f"custom_title_instr_{selected_doc_for_llm}_{selected_section_title_for_llm.replace(' ','_')}"
                                    )

                                    if selected_section_title_for_llm:
                                        selected_section_content = ""
                                        for sec in current_doc_sections:
                                            if sec["title"] == selected_section_title_for_llm:
                                                selected_section_content = sec.get("content","")
                                                break
                                        
                                        if st.button(f"Suggest Titles for '{selected_section_title_for_llm}'", key=f"sug_title_btn_instr_{selected_doc_for_llm}_{selected_section_title_for_llm.replace(' ','_')}"):
                                            if selected_section_content:
                                                title_suggestions = suggest_section_title_with_llm(llm, selected_section_content, selected_section_title_for_llm, custom_title_instructions)
                                                st.text_area(f"Suggested Titles for '{selected_section_title_for_llm}':", value=title_suggestions, height=100, key=f"title_sug_area_instr_{selected_doc_for_llm}")
                                            else:
                                                st.warning("Selected section has no content.")
                        else:
                            st.info("Select a processed document above.")


            with transform_tab:
                st.subheader("Apply Template and Download")
                # Simplified Preview Area
                if uploaded_files:
                    with st.expander("👀 Simplified Preview of First Uploaded Document", expanded=False):
                        # ... (preview logic, same as before)
                        first_file = uploaded_files[0]; st.markdown(f"**File:** `{first_file.name}`"); st.markdown(f"**Type:** `{first_file.type}`"); st.markdown("**Template Settings (Summary):**"); st.json({"Page Size": f"{st.session_state.current_template.get('page_width')}x{st.session_state.current_template.get('page_height')} inches", "Body Font (DOCX)": f"{st.session_state.current_template.get('body_font')} {st.session_state.current_template.get('body_size')}pt", "Header/Footer": f"Header: {'Yes' if st.session_state.current_template.get('use_header') else 'No'}, Footer: {'Yes' if st.session_state.current_template.get('use_footer') else 'No'}", "Logo": "Yes" if st.session_state.current_template.get('use_logo') and logo_file_upload else "No"}, expanded=False)


                if st.button(f"✨ Transform {len(uploaded_files)} Document(s) (Apply Template)"):
                    # ... (transformation loop and download buttons - same as before)
                    st.session_state.processed_docx_files.clear(); st.session_state.processed_pdf_files.clear()
                    with st.spinner(f"Transforming {len(uploaded_files)} document(s)... Please wait."):
                        for uploaded_file_item in uploaded_files:
                            st.markdown(f"--- \n Processing **{uploaded_file_item.name}**..."); progress_bar = st.progress(0); status_text = st.empty(); status_text.text("Starting transformation...") # type: ignore
                            docx_output_path, _ = process_uploaded_document(uploaded_file_item, st.session_state.current_template.copy(), logo_file_upload)
                            progress_bar.progress(50); status_text.text("DOCX generated...") # type: ignore
                            if docx_output_path:
                                st.session_state.processed_docx_files[uploaded_file_item.name] = docx_output_path; st.success(f"DOCX for '{uploaded_file_item.name}' transformed!") # type: ignore
                                if REPORTLAB_AVAILABLE:
                                    status_text.text("Generating PDF with ReportLab..."); pdf_output_dir = os.path.join(tempfile.gettempdir(), "documorph_pdfs"); os.makedirs(pdf_output_dir, exist_ok=True); pdf_output_filename = f"formatted_{os.path.splitext(uploaded_file_item.name)[0]}_{uuid.uuid4().hex[:8]}.pdf"; pdf_output_path_val = os.path.join(pdf_output_dir, pdf_output_filename); logo_path_for_pdf = None # type: ignore
                                    if logo_file_upload:
                                        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(logo_file_upload.name)[1]) as tmp_logo_rl: tmp_logo_rl.write(logo_file_upload.getvalue()); logo_path_for_pdf = tmp_logo_rl.name
                                    generated_pdf_path = generate_pdf_with_reportlab(docx_output_path, st.session_state.current_template.copy(), pdf_output_path_val, logo_path_for_pdf)
                                    if logo_path_for_pdf and os.path.exists(logo_path_for_pdf): os.remove(logo_path_for_pdf)
                                    if generated_pdf_path: st.session_state.processed_pdf_files[uploaded_file_item.name] = generated_pdf_path; st.info(f"PDF for '{uploaded_file_item.name}' generated (basic format).") # type: ignore
                                    else: st.warning(f"Could not generate PDF for '{uploaded_file_item.name}'.")
                                else: st.info("ReportLab not available, skipping PDF generation.")
                                progress_bar.progress(100); status_text.empty() # type: ignore
                            else: st.error(f"Failed to transform '{uploaded_file_item.name}'."); progress_bar.empty(); status_text.empty() # type: ignore
                        st.balloons()
                if st.session_state.processed_docx_files:
                    st.markdown("--- \n ### 📥 Download Transformed Documents")
                    if len(st.session_state.processed_docx_files) > 1 or len(st.session_state.processed_pdf_files) > 1:
                        if st.button("📦 Download All as ZIP"):
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_f:
                                for original_name, path_val in st.session_state.processed_docx_files.items():
                                    if os.path.exists(path_val): zip_f.write(path_val, arcname=os.path.basename(path_val))
                                for original_name, path_val in st.session_state.processed_pdf_files.items():
                                    if os.path.exists(path_val): zip_f.write(path_val, arcname=os.path.basename(path_val))
                            st.download_button(label="Click to Download ZIP", data=zip_buffer.getvalue(), file_name="transformed_documents.zip", mime="application/zip")
                    for original_name in st.session_state.processed_docx_files.keys():
                        st.markdown(f"#### {original_name}"); cols = st.columns(3) # type: ignore
                        with cols[0]:
                            if original_name in st.session_state.processed_docx_files:
                                docx_path = st.session_state.processed_docx_files[original_name]
                                if os.path.exists(docx_path):
                                    with open(docx_path, "rb") as fp_docx: st.download_button(label=f"Download DOCX", data=fp_docx, file_name=os.path.basename(docx_path), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx_dl_{original_name.replace('.', '_')}")
                        with cols[1]:
                            if original_name in st.session_state.processed_pdf_files:
                                pdf_path_val = st.session_state.processed_pdf_files[original_name]
                                if os.path.exists(pdf_path_val):
                                    with open(pdf_path_val, "rb") as fp_pdf: st.download_button(label=f"Download PDF", data=fp_pdf, file_name=os.path.basename(pdf_path_val), mime="application/pdf", key=f"pdf_dl_{original_name.replace('.', '_')}")
                    if len(st.session_state.processed_docx_files) == 1:
                        first_original_name = list(st.session_state.processed_docx_files.keys())[0]
                        first_docx_path = st.session_state.processed_docx_files[first_original_name]
                        if st.button(f"↗️ Export '{first_original_name}' to Google Docs", key="gdocs_export_single"):
                            if os.path.exists(first_docx_path):
                                export_status = export_to_google_docs(first_docx_path, f"Formatted_{first_original_name}")
                                st.info(export_status) # type: ignore
                            else:
                                st.error("Formatted DOCX file not found for export.")

    st.markdown("---"); st.caption("DocuMorph AI | Advanced Document Transformation")
    if not LANGCHAIN_AVAILABLE: st.sidebar.error("Langchain libraries not found. LLM features disabled. Install requirements.")
    if not REPORTLAB_AVAILABLE: st.sidebar.warning("ReportLab not installed. PDF generation disabled.")
    if not NLP_AVAILABLE or not nlp: st.sidebar.warning("Spacy model 'en_core_web_sm' not found. NLP section detection limited.")

if __name__ == "__main__":
    main()
